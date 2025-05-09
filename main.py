import ttkbootstrap as tk
from ttkbootstrap import style
import tkinter.filedialog
import tkinter.messagebox as msgbox
import tkinter.ttk
import matplotlib.pyplot as plt
from PIL import Image, ImageTk
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from matplotlib.pylab import mpl
mpl.rcParams['font.sans-serif']= ['SimHei'] # 显示中文
mpl.rcParams['axes.unicode_minus'] = False  # 显示负号
import time
import copy
import os
import numpy as np
import pandas as pd
#mid_fre = [1, 1.25, 1.6, 2, 2.5, 3.15, 4, 5, 6.3, 8, 10, 12.5, 16, 20, 25, 31.5, 40, 50, 63, 80, 100, 125, 160, 200,250, 315, 400, 500, 630, 800, 1000, 1200]
#fre_board = [0.89, 1.12, 1.41, 1.78, 2.24, 2.82, 3.55, 4.47, 5.62, 7.08, 8.91, 11.2, 14.1, 17.8, 22.4, 28.2, 35.5,44.7, 56.2, 70.8, 89.1, 112, 141, 178, 224, 282, 355, 447, 562, 708, 891, 1120, 1410]
"文件夹读取Auto————读取文件列表————输入path文件夹路径，输出fpath_name文件列表"
def read_list(path, list_name):  # 传入存储的list
    for file in os.listdir(path):
        file_path = os.path.join(path, file)
        if os.path.isdir(file_path):
            read_list(file_path, list_name)
        else:
            list_name.append(file_path)
fpath_name = []
"文件读取————读取各个通道原始数据————输入fpath_name文件列表，输出所有通道数据"
def read_data(fpath_name):#读取文件列表，生成对应dataframe
    Pamount = len(fpath_name)
    data0 = list(np.arange(Pamount))
    for i in range(Pamount):
        data0[i] = pd.read_csv(
            fpath_name[i],
            sep="   ",
            header=None,
            names=['time', 'a']
        )
    return data0
"预处理————清洗所有通道数据，截取————输入所有通道数据dataframe，输出清洗后所有通道数据"
def clean_data(dataframe):
    data = list(np.arange(len(dataframe)))
    qushixiang=list(np.arange(len(dataframe)))
    a_mean=list(np.arange(len(dataframe)))
    for i in range(len(dataframe)):
    #去空值
        dataframe[i].dropna(inplace=True)
    #消除趋势项
        param = np.polyfit(dataframe[i].time, dataframe[i].a, 5)  # 用5次多项式拟合x，y数组,返回多项式系数
        y_poly_func = np.poly1d(param)  # 拟合完之后，用生成的多项式系数用来生成多项式函数
        y_poly = y_poly_func(dataframe[i].time)  # 生成多项式函数之后，就是获取x在这个多项式处的值
        dataframe[i].a = dataframe[i].a - y_poly  # 原始信号减去拟合函数的值即为去除趋势项后的信号
        qushixiang[i]=np.mean(y_poly)
    #消除零点偏移
        a_mean[i] = np.mean(dataframe[i].a)
        dataframe[i].a = dataframe[i].a - a_mean[i]
    #提取数据
        sum = 0
        n = 0
        time2 = 0
        for j in dataframe[i].a:
            time1 = dataframe[i].time[n]
            n = n + 1
            sum = sum + j * j * (time1 - time2)
            time2 = time1
        sum = sum / time1
        Arms = np.sqrt(sum)
        n = len(dataframe[i].a)
        A = list(range(n))
        for j in A:
           if (abs(dataframe[i].a[j]) >= Arms):
               Indexmax = j
        A.reverse()
        for j in A:
           if (abs(dataframe[i].a[j]) >= Arms):
               Indexmin = j
        data[i] = dataframe[i][Indexmin:Indexmax]
        data[i] = data[i].reset_index(drop=True)
        Startime = data[i].time[0]
        data[i].time = data[i].time - Startime
    return data,qushixiang,a_mean
"预处理————清洗所有通道数据,不截取，留作与DASP对比分析————输入所有通道数据dataframe，输出清洗后所有通道数据"
def clean_data_0(dataframe):
    qushixiang=list(np.arange(len(dataframe)))
    a_mean=list(np.arange(len(dataframe)))
    for i in range(len(dataframe)):
    #去空值
        dataframe[i].dropna(inplace=True)
    #消除趋势项
        param = np.polyfit(dataframe[i].time, dataframe[i].a, 5)  # 用5次多项式拟合x，y数组,返回多项式系数
        y_poly_func = np.poly1d(param)  # 拟合完之后，用生成的多项式系数用来生成多项式函数
        y_poly = y_poly_func(dataframe[i].time)  # 生成多项式函数之后，就是获取x在这个多项式处的值
        dataframe[i].a = dataframe[i].a - y_poly  # 原始信号减去拟合函数的值即为去除趋势项后的信号
        qushixiang[i] = np.mean(y_poly)
    #消除零点偏移
        a_mean[i] = np.mean(dataframe[i].a)
        dataframe[i].a = dataframe[i].a - a_mean[i]
    return dataframe,qushixiang,a_mean
"最值————计算各个通道最大值————输入所有通道数据data，输出各个通道加速度最大值"
def max_data(data):
    MAX = list(np.arange(len(data)))
    for i in range(len(data)):
        max1=max(data[i].a)
        max2=min(data[i].a)
        MAX[i] = max(max1,abs(max2))
    return MAX
"有效值————计算各个通道数据有效值————输入所有通道数据data，输出各个通道加速度有效值"
def arms_data(data):
    Arms = list(np.arange(len(data)))
    for i in range(len(data)):
        sum = 0
        n = 0
        time1 = 0
        time2 = 0
        for j in data[i].a:
            time1 = data[i].loc[n, 'time']
            n = n + 1
            sum = sum + j * j * (time1 - time2)
            time2 = time1
        sum = sum / time1
        Arms[i] = np.sqrt(sum)
    return Arms
"VAL（振动加速度级）————计算各个通道振动加速度级————输入所有通道有效值arms，输出各个通道振动加速度级"
def VAL_data(Arms):
    VAL = list(np.arange(len(Arms)))
    for i in range(len(Arms)):
        VAL[i] = 20 * np.log10(Arms[i] / pow(10, -6))
    return VAL
"FFT————快速傅里叶变换————输入单个通道数据dataframe，输出fre频率,amplitude振幅,power功率"#0，1，2
def fft_data(dataframe):
    time1 = dataframe.loc[0, 'time']
    time2 = dataframe.loc[1, 'time']
    dt = time2 - time1
    FN = int(1 / dt)
    n = len(dataframe.a)
    if n<=65536/2:
        n=65536/2
    c = 0
    while 2 ** c < 2*n:
        c += 1
        N = 2 ** c
    fr = np.arange(0, FN, FN / N)
    fre = fr[range(int(N / 2))]  # 频率
    Zfa = np.fft.fft(dataframe.a, N)
    abs_Ag = np.abs(Zfa)
    amplitude = abs_Ag[range(int(N / 2))] * 2 / N  # 振幅
    power = abs_Ag[range(int(N / 2))] ** 2 / N  # 功率
    return fre,amplitude,power
"FFT——显著频率"
def fft_vis(dataframe):
    time1 = dataframe.loc[0, 'time']
    time2 = dataframe.loc[1, 'time']
    dt = time2 - time1
    FN = int(1 / dt)
    n = len(dataframe.a)
    c = 0
    while 2 ** c < 2 * n:
        c += 1
        N = 2 ** c
    fr = np.arange(0, FN, FN / N)
    fre = fr[range(int(N / 2))]  # 频率
    Dt = fre[2] - fre[1]
    Zfa = np.fft.fft(dataframe.a, N)
    abs_Ag = np.abs(Zfa)
    amplitude = abs_Ag[range(int(N / 2))] * 2 / N  # 振幅
    power = abs_Ag[range(int(N / 2))] ** 2 / N  # 功率
    clean_noise(fre,amplitude,power,4)
    obs_fre = [0, 0, 0, 0]  # 返回显著频率
    obs_item = list(range(4))  # 对应索引
    max_am = [0, 0, 0, 0]  # 返回最大幅值
    max_power = max(power)  # 返回最大功率
    pre_data = abs_Ag[range(int(N / 2))] * 2 / N
    clean_noise(fre,pre_data,power,4)
    for i in range(4):
        max_am[i] = max(pre_data)
        for j in range(len(pre_data)):
            if pre_data[j] == max_am[i]:
                obs_fre[i] = int(fre[j])
                obs_item[i] = j
        if obs_item[i] - int(100 / Dt) < 0:
            for j in range(obs_item[i] + int(100 / Dt)):
                pre_data[j] = 0
        elif obs_item[i] + int(100 / Dt) > len(pre_data):
            for j in range(obs_item[0] + int(100 / Dt)):
                pre_data[len(pre_data) - j - 1] = 0
        else:
            for j in range(2 * int(100 / Dt)):
                pre_data[obs_item[i]-int(100/Dt) + j] = 0
        max_am[i]=round(max_am[i],4)
    pre_obs_fre=copy.deepcopy(obs_fre)
    pre_max_am=copy.deepcopy(max_am)
    obs_fre.sort()
    for i in range(4):
        for j in range(4):
            if pre_obs_fre[j]==obs_fre[i]:
                max_am[i]=pre_max_am[j]
    return fre, amplitude, power, obs_fre, max_am, max_power
"计权最大Z振级VLzmax————输入单个通道数据dataframe，输出VL_time时点,VL_zw85计权曲线,VL_zwk97计权曲线"#0，1，2
def fft_zmax_w(dataframe):
    #返回时间序列及VL_w、VL_wk序列
    dt = dataframe.time[1]-dataframe.time[0]
    fs = int(1/dt)
    wave_num = int((len(dataframe.time)-fs)/(0.25*fs))+1
    length = int(len(dataframe.time)/wave_num)
    VL_time = list(np.arange(wave_num))#返回时间序列
    VL_z = list(np.arange(wave_num))#返回不计权曲线
    VL_zw = list(np.arange(wave_num))#返回85计权曲线
    VL_zwk = list(np.arange(wave_num))#返回97计权曲线
    mid_fre = [1,1.25,1.6,2,2.5,3.15,4,5,6.3,8,10,12.5,16,20,25,31.5,40,50,63,80]
    fre_board = [0.89, 1.12, 1.41, 1.78, 2.24, 2.82, 3.55, 4.47, 5.62, 7.08, 8.91, 11.2, 14.1, 17.8, 22.4, 28.2, 35.5,
                 44.7, 56.2, 70.8, 89.1]
    alpha_w = [6,5,4,3,2,1,0,0,0,0,2,4,6,8,10,12,14,16,18,20]
    alpha_wk = [6.33,6.29,6.12,5.49,4.01,1.9,0.29,-0.33,-0.46,-0.31,0.1,0.89,2.28,3.93,5.8,7.86,10.05,12.19,14.61,17.56]
    for i in range(wave_num):#求每一幅的Z振级
        VL_dt = dataframe[i*length:(i+1)*length]
        VL_dt = VL_dt.reset_index(drop=True)
        VL_time[i] = VL_dt.time[int(length/2)]#每一幅的中心时间
        VL_fre = fft_data(VL_dt)[0]
        VL_am = fft_data(VL_dt)[1]
        item = list(np.arange(len(fre_board)))
        sum_0 = list(np.arange(20))
        sum_w = list(np.arange(20))
        sum_wk = list(np.arange(20))
        datafre = list(np.arange(20))
        dataag = list(np.arange(20))
        for j in range(20):
            sum_0[j]=0
            sum_w[j]=0
            sum_wk[j]=0
        for j in range(len(VL_fre)):
            for k in range(len(fre_board)):
                if (VL_fre[j] <= fre_board[k]):
                    item[k] = j
        sum0 = 0
        sum1 = 0
        sum2 = 0
        for j in range(len(mid_fre)):
            datafre[j] = VL_fre[item[j]:item[j+1]]
            dataag[j] = VL_am[item[j]:item[j+1]]
            sum11=0
            if (len(dataag[j]) == 0):
                sum11 = 0
            else:
                for k in dataag[j]:
                    sum11 = sum11 + k * k / 2
            Arms = np.sqrt(sum11)
            if (Arms == 0):
                VL = 0
            else:
                VL = 20 * np.log10(Arms / pow(10, -6))
            sum_0[j] = 10 ** ((VL) / 10)
            sum_w[j] = 10 ** ((VL - alpha_w[j]) / 10)
            sum_wk[j] = 10 ** ((VL - alpha_wk[j]) / 10)
            sum2 += sum_0[j]
            sum0 += sum_w[j]
            sum1 += sum_wk[j]
        VL_z[i] = 10 * np.log10(sum2)
        VL_zw[i] = 10 * np.log10(sum0)
        VL_zwk[i] = 10 * np.log10(sum1)
    return VL_time,VL_zw,VL_zwk,VL_z
"不计权最大Z振级VLzmax————输入单个通道数据dataframe，输出VL_time时点,VL_z"
def fft_zmax_0(dataframe):
    # 返回时间序列及VL_z序列
    dt = dataframe.time[1] - dataframe.time[0]
    fs = int(1 / dt)
    wave_num = int((len(dataframe.time) - fs) / (0.25 * fs)) + 1
    length = int(len(dataframe.time) / wave_num)
    VL_time = list(np.arange(wave_num))
    VL_z = list(np.arange(wave_num))
    mid_fre = [1, 1.25, 1.6, 2, 2.5, 3.15, 4, 5, 6.3, 8, 10, 12.5, 16, 20, 25, 31.5, 40, 50, 63, 80, 100, 125, 160, 200,
               250, 315, 400, 500, 630, 800, 1000, 1200]
    fre_board = [0.89, 1.12, 1.41, 1.78, 2.24, 2.82, 3.55, 4.47, 5.62, 7.08, 8.91, 11.2, 14.1, 17.8, 22.4, 28.2, 35.5,
                 44.7, 56.2, 70.8, 89.1, 112, 141, 178, 224, 282, 355, 447, 562, 708, 891, 1120, 1410]
    for i in range(wave_num):
        VL_dt = dataframe[i * length:(i + 1) * length]
        VL_dt = VL_dt.reset_index(drop=True)
        VL_time[i] = VL_dt.time[int(length / 2)]
        VL_fre = fft_data(VL_dt)[0]
        VL_am = fft_data(VL_dt)[1]
        item = list(np.arange(len(fre_board)))
        sum_w = list(np.arange(len(mid_fre)))
        datafre = list(np.arange(len(mid_fre)))
        dataag = list(np.arange(len(mid_fre)))
        for j in range(len(VL_fre)):
            for k in range(len(fre_board)):
                if (VL_fre[j] <= fre_board[k]):
                    item[k] = j
        sum0 = 0
        for j in range(len(mid_fre)):
            datafre[j] = VL_fre[item[j]:item[j + 1]]
            dataag[j] = VL_am[item[j]:item[j + 1]]
            if (len(dataag[j]) == 0):
                sum = 0
            else:
                for k in dataag[j]:
                    sum = sum + k * k / 2
            Arms = np.sqrt(sum)
            if (Arms == 0):
                VL = 0
            else:
                VL = 20 * np.log10(Arms / pow(10, -6))
            sum_w[j] = 10 ** (VL / 10)
            sum0 += sum_w[j]
        VL_z[i] = 10 * np.log10(sum0)
    return VL_time, VL_z
"1/3倍频程变换分频最大振级————加窗计权分频最大振级VLmax————输入单个通道数据dataframe，输出VLmax各个中心频率对应的振级"
def fft_max_w(dataframe):
    #VL_max（18）序列
    dt = dataframe.time[1] - dataframe.time[0]
    fs = int(1 / dt)
    wave_num = int((len(dataframe.time) - fs) / (0.25 * fs)) + 1
    length = int(len(dataframe.time) / wave_num)
    mid_fre = [4, 5, 6.3, 8, 10, 12.5, 16, 20, 25, 31.5, 40, 50, 63, 80,100,125,160,200]#输出
    fre_board = [3.55, 4.47, 5.62, 7.08, 8.91, 11.2, 14.1, 17.8, 22.4, 28.2, 35.5,
                 44.7, 56.2, 70.8, 89.1, 112, 141, 178, 224]
    alpha = [0, 0, 0, 0, 0,-1,-2,-4,-6,-8,-10,-12,-14,-17,-21,-25,-30,-36]
    VL_max = list(np.arange(len(mid_fre)))
    for i in range(len(mid_fre)):
        VL_max[i]=list(np.arange(wave_num))
    VLmax = list(np.arange(len(mid_fre)))#输出
    for i in range(wave_num):
        VL_dt = dataframe[i * length:(i + 1) * length]
        VL_dt = VL_dt.reset_index(drop=True)
        VL_fre = fft_data(VL_dt)[0]
        VL_am = fft_data(VL_dt)[1]
        item = list(np.arange(len(fre_board)))
        datafre = list(np.arange(len(mid_fre)))
        dataag = list(np.arange(len(mid_fre)))
        for j in range(len(VL_fre)):
            for k in range(len(fre_board)):
                if (VL_fre[j] <= fre_board[k]):
                    item[k] = j
        for j in range(len(mid_fre)):
            datafre[j] = VL_fre[item[j]:item[j + 1]]
            dataag[j] = VL_am[item[j]:item[j + 1]]
            sum = 0
            if (len(dataag[j]) == 0):
                sum = 0
            else:
                for k in dataag[j]:
                    sum = sum + k * k/2
            Arms = np.sqrt(sum)
            if (Arms == 0):
                VL_max[j][i] = 0
            else:
                VL_max[j][i] = 20 * np.log10(Arms / pow(10, -6))+alpha[j]
    for i in range(len(mid_fre)):
        VLmax[i]=int(max(VL_max[i]))
    return VLmax
"1/3倍频程变换分频最大振级————不加窗不计权分频最大振级VLmax————输入单个通道数据dataframe，输出VLmax各个中心频率对应的振级"
def fft_max_0(dataframe):
    #返回VL_fre频率序列（4-200Hz）以及VL_max（18）序列
    mid_fre = [1, 1.25, 1.6, 2, 2.5, 3.15, 4, 5, 6.3, 8, 10, 12.5, 16, 20, 25, 31.5, 40, 50, 63, 80, 100, 125, 160, 200,
               250, 315, 400, 500, 630, 800, 1000, 1200]
    fre_board = [0.89, 1.12, 1.41, 1.78, 2.24, 2.82, 3.55, 4.47, 5.62, 7.08, 8.91, 11.2, 14.1, 17.8, 22.4, 28.2, 35.5,
                 44.7, 56.2, 70.8, 89.1, 112, 141, 178, 224, 282, 355, 447, 562, 708, 891, 1120, 1410]
    VLmax = list(np.arange(len(mid_fre)))
    VL_fre = fft_data(dataframe)[0]
    VL_am = fft_data(dataframe)[1]
    item = list(np.arange(len(fre_board)))
    datafre = list(np.arange(len(mid_fre)))
    dataag = list(np.arange(len(mid_fre)))
    for i in range(len(VL_fre)):
        for j in range(len(fre_board)):
            if (VL_fre[i] <= fre_board[j]):
                item[j] = i
    for i in range(len(mid_fre)):
        datafre[i] = VL_fre[item[i]:item[i + 1]]
        dataag[i] = VL_am[item[i]:item[i + 1]]
    for i in range(len(mid_fre)):
        sum = 0
        if (len(dataag[i]) == 0):
            sum = 0
        else:
            for j in dataag[i]:
                sum = sum + j * j / 2
        Arms = np.sqrt(sum)
        if (Arms == 0):
            VLmax[i] = 0
        else:
            VLmax[i] = int(20 * np.log10(Arms / pow(10, -6)))
    return VLmax
"1/3倍频程变换分频最大振级————加窗不计权分频最大振级VLmax————输入单个通道数据dataframe，输出VLmax各个中心频率对应的振级"
def fft_max_1(dataframe):
    dt = dataframe.time[1] - dataframe.time[0]
    fs = int(1 / dt)
    wave_num = int((len(dataframe.time) - fs) / (0.25 * fs)) + 1
    length = int(len(dataframe.time) / wave_num)
    mid_fre = [1, 1.25, 1.6, 2, 2.5, 3.15, 4, 5, 6.3, 8, 10, 12.5, 16, 20, 25, 31.5, 40, 50, 63, 80, 100, 125, 160, 200,
               250, 315, 400, 500, 630, 800, 1000, 1200]
    fre_board = [0.89, 1.12, 1.41, 1.78, 2.24, 2.82, 3.55, 4.47, 5.62, 7.08, 8.91, 11.2, 14.1, 17.8, 22.4, 28.2, 35.5,
                 44.7, 56.2, 70.8, 89.1, 112, 141, 178, 224, 282, 355, 447, 562, 708, 891, 1120, 1410]
    VL_max = list(np.arange(len(mid_fre)))
    for i in range(len(mid_fre)):
        VL_max[i] = list(np.arange(wave_num))
    VLmax = list(np.arange(len(mid_fre)))  # 输出
    for i in range(wave_num):
        VL_dt = dataframe[i * length:(i + 1) * length]
        VL_dt = VL_dt.reset_index(drop=True)
        VL_fre = fft_data(VL_dt)[0]
        VL_am = fft_data(VL_dt)[1]
        item = list(np.arange(len(fre_board)))
        datafre = list(np.arange(len(mid_fre)))
        dataag = list(np.arange(len(mid_fre)))
        for j in range(len(VL_fre)):
            for k in range(len(fre_board)):
                if (VL_fre[j] <= fre_board[k]):
                    item[k] = j
        for j in range(len(mid_fre)):
            datafre[j] = VL_fre[item[j]:item[j + 1]]
            dataag[j] = VL_am[item[j]:item[j + 1]]
            sum = 0
            if (len(dataag[j]) == 0):
                sum = 0
            else:
                for k in dataag[j]:
                    sum = sum + k * k / 2
            Arms = np.sqrt(sum)
            if (Arms == 0):
                VL_max[j][i] = 0
            else:
                VL_max[j][i] = 20 * np.log10(Arms / pow(10, -6))
    for i in range(len(mid_fre)):
        VLmax[i] = int(max(VL_max[i]))
    return VLmax
"去噪——输入频率、振幅、功率、去噪阈值，输出去噪后结果"
def clean_noise(fre,amplitude,power,Clean_fre):
    j=0
    while fre[j]<=Clean_fre:#滤除3Hz以下波形
        amplitude[j]=0#振幅
        power[j]=0#功率
        j+=1
    return
def mode1():
    global mode
    mode=1
    msgbox.showinfo('分析模式选择', '已切换至单独分析模式')
def mode2():
    global mode
    mode=2
    msgbox.showinfo('分析模式选择', '已切换至批量分析模式')
def mode3():
    global mode
    mode=3
    msgbox.showinfo('分析模式选择', '已切换至对比分析模式')
def readfile():
    global dirname, numname,Pamount, pass_num, time_last, num_save, fr_save
    global fpath_name
    dirname = time.strftime('%Y%m%d%H%M%S', time.localtime(time.time()))#结果保存路径
    os.makedirs("result\\" + dirname)
    os.makedirs("cache\\" + dirname)
    if mode==1:
        filename = tkinter.filedialog.askdirectory()
        fpath_name = []
        read_list(filename, fpath_name)  # 读取文件列表
        Pamount = len(fpath_name)  # 读取通道数目
        Initialdata = read_data(fpath_name)  # 读取所有通道原始文件
        pass_num = list(range(Pamount))#通道数
        time_last = list(range(Pamount))#采样时间
        num_save = list(range(Pamount))#采样点
        fr_save = list(range(Pamount))#采样频率
        for i in range(Pamount):
            pass_num[i]=i+1
            time_last[i]=int(Initialdata[i].time[len(Initialdata[i].time) - 1])
            num_save[i]=len(Initialdata[i].time)
            fr_save[i]=(1 / (Initialdata[i].time[1] - Initialdata[i].time[0]))
        for i in range(Pamount):
            plt.figure(figsize=(8,0.9))
            plt.plot(Initialdata[i].time, Initialdata[i].a)
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
            plt.savefig("cache\\"+dirname+"\\0时程图-通道{}".format(i+1)+".png")
            plt.close()
            plt.cla()
        for i in range(Pamount):
            plt.figure()
            plt.plot(Initialdata[i].time, Initialdata[i].a)
            plt.xlabel("时间", fontsize=14)
            plt.ylabel("加速度/ms-2", fontsize=14)
            plt.xticks(fontsize=14)
            plt.yticks(fontsize=14)
            plt.savefig("result\\" + dirname + "\\0时程图-通道{}".format(i + 1) + ".png")
            plt.close()
            plt.cla()
        frame=tk.Frame(window,width=1920,height=1080)
        frame.grid(row=0,column=0)
        frame.grid_propagate(0)
        def show(event):
            pre_num=var.get()
        def show_image11():
            plt.figure()
            for i in range(Pamount):
                plt.subplot(Pamount,1,i+1)
                plt.plot(Initialdata[i].time, Initialdata[i].a)
                plt.ylabel("加速度/ms-2", fontsize=14)
                if i ==Pamount-1:
                    plt.xlabel("时间", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
            plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var,width=9)
        combobox['value'] = ['Group1']
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0,column=0)
        label0=tk.Label(frame,text='采样信息')
        label0.grid(row=0,column=1,columnspan=2)
        label1=tk.Label(frame,text='时程图')
        label1.grid(row=0,column=3)
        buttun=tk.Button(frame,text='显示详图',command=show_image11)
        buttun.grid(row=0,column=4)
        label = list(range(Pamount*3))
        img=list(range(Pamount))
        photo=list(range(Pamount))
        for i in range(Pamount*3):
            label[i] = list(range(4))
        text1=['采样时间/s','采样点数','采样频率/Hz']
        for i in range(Pamount):
            label[i][0]=tk.Label(frame,text='通道{}'.format(i+1))
            label[i][0].grid(row=1+i*3,column=0,rowspan=3)
            label[i][1]=list(range(3))
            label[i][2]=list(range(3))
            for j in range(3):
                label[i][1][j]=tk.Label(frame,text=text1[j])
                label[i][1][j].grid(row=1+i*3+j,column=1)
            label[i][2][0]=tk.Label(frame,text=time_last[j])
            label[i][2][0].grid(row=1+i*3+0,column=2)
            label[i][2][1]=tk.Label(frame,text=num_save[j])
            label[i][2][1].grid(row=1+i*3+1,column=2)
            label[i][2][2]=tk.Label(frame,text=fr_save[j])
            label[i][2][2].grid(row=1+i*3+2,column=2)
            label[i][3]=tk.Label(frame)
            img[i] = Image.open("cache\\"+dirname+"\\0时程图-通道{}".format(i+1)+".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            label[i][3].config(image=photo[i])
            label[i][3].image=photo[i]
            label[i][3].grid(row=1+i*3,column=3,rowspan=3,columnspan=2)
    if mode==2:
        filename0 = tkinter.filedialog.askdirectory()
        numname = os.listdir(filename0)
        filename1 = list(range(len(numname)))
        for i in range(len(numname)):
            filename1[i] = os.path.join(filename0, numname[i])
        fpath_name=list(range(len(numname)))
        Pamount=list(range(len(numname)))
        pass_num=list(range(len(numname)))
        time_last=list(range(len(numname)))
        num_save=list(range(len(numname)))
        fr_save=list(range(len(numname)))
        for i in range(len(numname)):
            fpath_name[i] = []
            read_list(filename1[i], fpath_name[i])  # 读取文件列表
            Pamount[i] = len(fpath_name[i])  # 读取通道数目
            Initialdata = read_data(fpath_name[i])  # 读取所有通道原始文件
            pass_num[i] = list(range(Pamount[i]))  # 通道数
            time_last[i] = list(range(Pamount[i]))  # 采样时间
            num_save[i] = list(range(Pamount[i]))  # 采样点
            fr_save[i] = list(range(Pamount[i]))  # 采样频率
            for j in range(Pamount[i]):
                pass_num[i][j] = j + 1
                time_last[i][j] = int(Initialdata[j].time[len(Initialdata[j].time) - 1])
                num_save[i][j] = len(Initialdata[j].time)
                fr_save[i][j] = (1 / (Initialdata[j].time[1] - Initialdata[j].time[0]))
            # 图片可视化文件保存至cache文件夹中
            os.makedirs("cache\\" + dirname + "\\" + numname[i])
            for j in range(Pamount[i]):
                plt.figure(figsize=(8, 0.9))
                plt.plot(Initialdata[j].time, Initialdata[j].a)
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname +"\\"+numname[i]+ "\\0时程图-通道{}".format(j + 1) + ".png")
                plt.close()
                plt.cla()
        for i in range(len(numname)):
            os.makedirs("result\\" + dirname + "\\" + numname[i])
            Initialdata = read_data(fpath_name[i])
            for j in range(Pamount[i]):
                plt.figure()
                plt.plot(Initialdata[j].time, Initialdata[j].a)
                plt.xlabel("时间", fontsize=14)
                plt.ylabel("加速度/ms-2", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname +"\\"+numname[i]+ "\\0时程图-通道{}".format(j + 1) + ".png")
                plt.close()
                plt.cla()
        frame=tk.Frame(window,width=1920,height=1080)
        frame.grid(row=0,column=0)
        frame.grid_propagate(0)
        label = list(range(Pamount[0] * 3))
        img = list(range(Pamount[0]))
        photo = list(range(Pamount[0]))
        for i in range(Pamount[0] * 3):
            label[i] = list(range(4))
        text1 = ['采样时间/s', '采样点数', '采样频率/Hz']
        for i in range(Pamount[0]):
            label[i][0] = tk.Label(frame, text='通道{}'.format(i + 1))
            label[i][0].grid(row=1 + i * 3, column=0, rowspan=3)
            label[i][1] = list(range(3))
            label[i][2] = list(range(3))
            for j in range(3):
                label[i][1][j] = tk.Label(frame, text=text1[j])
                label[i][1][j].grid(row=1 + i * 3 + j, column=1)
            label[i][2][0] = tk.Label(frame, text=time_last[0][j])
            label[i][2][0].grid(row=1 + i * 3 + 0, column=2)
            label[i][2][1] = tk.Label(frame, text=num_save[0][j])
            label[i][2][1].grid(row=1 + i * 3 + 1, column=2)
            label[i][2][2] = tk.Label(frame, text=fr_save[0][j])
            label[i][2][2].grid(row=1 + i * 3 + 2, column=2)
            label[i][3] = tk.Label(frame)
            img[i] = Image.open("cache\\" + dirname + "\\" + numname[0] + "\\0时程图-通道{}".format(i + 1) + ".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            label[i][3].config(image=photo[i])
            label[i][3].image = photo[i]
            label[i][3].grid(row=1 + i * 3, column=3, rowspan=3, columnspan=2)
        pre_num=0
        def show(event):
            a=var.get()
            for i in range(len(numname)):
                if a==numname[i]:
                    pre_num=i
            for i in range(Pamount[pre_num]):
                label[i][2][0].config(text=time_last[pre_num][i])
                label[i][2][1].config(text=num_save[pre_num][i])
                label[i][2][2].config(text=fr_save[pre_num][i])
                img[i] = Image.open("cache\\" + dirname + "\\" + numname[pre_num]+ "\\0时程图-通道{}".format(i + 1) + ".png")  # 打开图片
                photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                label[i][3].config(image=photo[i])
                label[i][3].image = photo[i]
                label[i][3].grid(row=1 + i * 3, column=3, rowspan=3, columnspan=2)
        def show_image12():
            a=var.get()
            for i in range(len(numname)):
                if a==numname[i]:
                    pre_num=i
            plt.figure(num=numname[pre_num])
            Initialdata = read_data(fpath_name[pre_num])
            for i in range(Pamount[pre_num]):
                plt.subplot(Pamount[pre_num],1,i+1)
                plt.plot(Initialdata[i].time, Initialdata[i].a)
                if j ==Pamount[pre_num]-1:
                    plt.xlabel("时间", fontsize=14)
                plt.ylabel("加速度/ms-2", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
            plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var,width=9)
        combobox['value'] = numname
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0,column=0)
        label0=tk.Label(frame,text='采样信息')
        label0.grid(row=0,column=1,columnspan=2)
        label1=tk.Label(frame,text='时程图')
        label1.grid(row=0,column=3)
        buttun=tk.Button(frame,text='显示详图',command=show_image12)
        buttun.grid(row=0,column=4)
    if mode==3:
        filename0 = tkinter.filedialog.askdirectory()
        numname = os.listdir(filename0)
        numname.append('汇总对比')
        filename1 = list(range(len(numname)-1))
        for i in range(len(numname)-1):
            filename1[i] = os.path.join(filename0, numname[i])
        fpath_name = list(range(len(numname)-1))
        Pamount = list(range(len(numname)-1))
        pass_num = list(range(len(numname)-1))
        time_last = list(range(len(numname)-1))
        num_save = list(range(len(numname)-1))
        fr_save = list(range(len(numname)-1))
        for i in range(len(numname)-1):
            fpath_name[i] = []
            read_list(filename1[i], fpath_name[i])  # 读取文件列表
            Pamount[i] = len(fpath_name[i])  # 读取通道数目
            Initialdata = read_data(fpath_name[i])  # 读取所有通道原始文件
            pass_num[i] = list(range(Pamount[i]))  # 通道数
            time_last[i] = list(range(Pamount[i]))  # 采样时间
            num_save[i] = list(range(Pamount[i]))  # 采样点
            fr_save[i] = list(range(Pamount[i]))  # 采样频率
            for j in range(Pamount[i]):
                pass_num[i][j] = j + 1
                time_last[i][j] = int(Initialdata[j].time[len(Initialdata[j].time) - 1])
                num_save[i][j] = len(Initialdata[j].time)
                fr_save[i][j] = (1 / (Initialdata[j].time[1] - Initialdata[j].time[0]))
            # 图片可视化文件保存至cache文件夹中
            os.makedirs("cache\\" + dirname + "\\" + numname[i])
            for j in range(Pamount[i]):
                plt.figure(figsize=(8, 0.9))
                plt.plot(Initialdata[j].time, Initialdata[j].a)
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname + "\\" + numname[i] + "\\0时程图-通道{}".format(j + 1) + ".png")
                plt.close()
                plt.cla()
        for i in range(len(numname)-1):
            os.makedirs("result\\" + dirname + "\\" + numname[i])
            Initialdata = read_data(fpath_name[i])
            for j in range(Pamount[i]):
                plt.figure()
                plt.plot(Initialdata[j].time, Initialdata[j].a)
                plt.xlabel("时间", fontsize=14)
                plt.ylabel("加速度/ms-2", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname +"\\"+numname[i]+ "\\0时程图-通道{}".format(j+ 1) + ".png")
                plt.close()
                plt.cla()
        frame = tk.Frame(window,width=1920,height=1080)
        frame.grid(row=0, column=0)
        frame.grid_propagate(0)
        label = list(range(Pamount[0] * 3))
        img = list(range(Pamount[0]))
        photo = list(range(Pamount[0]))
        for i in range(Pamount[0] * 3):
            label[i] = list(range(4))
        text1 = ['采样时间/s', '采样点数', '采样频率/Hz']
        for i in range(Pamount[0]):
            label[i][0] = tk.Label(frame, text='通道{}'.format(i + 1))
            label[i][0].grid(row=1 + i * 3, column=0, rowspan=3)
            label[i][1] = list(range(3))
            label[i][2] = list(range(3))
            for j in range(3):
                label[i][1][j] = tk.Label(frame, text=text1[j])
                label[i][1][j].grid(row=1 + i * 3 + j, column=1)
            label[i][2][0] = tk.Label(frame, text=time_last[0][j])
            label[i][2][0].grid(row=1 + i * 3 + 0, column=2)
            label[i][2][1] = tk.Label(frame, text=num_save[0][j])
            label[i][2][1].grid(row=1 + i * 3 + 1, column=2)
            label[i][2][2] = tk.Label(frame, text=fr_save[0][j])
            label[i][2][2].grid(row=1 + i * 3 + 2, column=2)
            label[i][3] = tk.Label(frame)
            img[i] = Image.open("cache\\" + dirname + "\\" + numname[0] + "\\0时程图-通道{}".format(i + 1) + ".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            label[i][3].config(image=photo[i])
            label[i][3].image = photo[i]
            label[i][3].grid(row=1 + i * 3, column=3, rowspan=3, columnspan=2)
        pre_num = 0
        def show(event):
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            if pre_num<=len(numname)-2:
                label1.grid(row=0, column=3)
                buttun.grid(row=0, column=4)
                for i in range(Pamount[pre_num]):
                    label[i][2][0].config(text=time_last[pre_num][i])
                    label[i][2][1].config(text=num_save[pre_num][i])
                    label[i][2][2].config(text=fr_save[pre_num][i])
                    img[i] = Image.open(
                        "cache\\" + dirname + "\\" + numname[pre_num] + "\\0时程图-通道{}".format(i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    label[i][3].config(image=photo[i])
                    label[i][3].image = photo[i]
                    label[i][3].grid(row=1 + i * 3, column=3, rowspan=3, columnspan=2)
            if pre_num==len(numname)-1:
                for i in range(Pamount[pre_num-1]):
                    timelast=[]
                    numsave=[]
                    frsave=[]
                    for j in range(len(numname)-1):
                            timelast.append(time_last[j][i])
                            numsave.append(num_save[j][i])
                            frsave.append(fr_save[j][i])
                    label[i][2][0].config(text=timelast)
                    label[i][2][1].config(text=numsave)
                    label[i][2][2].config(text=frsave)
                    label[i][3].grid_forget()
                    label1.grid_forget()
                    buttun.grid_forget()
        def show_image13():
            a=var.get()
            for i in range(len(numname)):
                if a==numname[i]:
                    pre_num=i
            if pre_num==len(numname)-1:
                pass
            else:
                plt.figure(num=numname[pre_num])
                Initialdata = read_data(fpath_name[pre_num])
                for i in range(Pamount[pre_num]):
                    plt.subplot(Pamount[pre_num], 1, i + 1)
                    plt.plot(Initialdata[i].time, Initialdata[i].a)
                    if j == Pamount[pre_num] - 1:
                        plt.xlabel("时间", fontsize=14)
                    plt.ylabel("加速度/ms-2", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = numname
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
        label0 = tk.Label(frame, text='采样信息')
        label0.grid(row=0, column=1, columnspan=2)
        label1 = tk.Label(frame, text='时程图')
        label1.grid(row=0, column=3)
        buttun = tk.Button(frame, text='显示详图', command=show_image13)
        buttun.grid(row=0, column=4)
def prepare_0():
    global data,x,y
    if mode==1:
        Initialdata = read_data(fpath_name)
        pre_data=clean_data(Initialdata)
        data = pre_data[0]
        x = pre_data[1]
        y = pre_data[2]
        text2=['零点偏移量','时间趋势项大小','异常值个数']
        frame=tk.Frame(window,width=1920,height=1080)
        frame.grid(row=0,column=0)
        frame.grid_propagate(0)
        def show(event):
            pre_num=var.get()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var,width=9)
        combobox['value'] = ['Group1']
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0,column=0)
        label=list(range(Pamount+1))
        for i in range(Pamount+1):
            label[i]=list(range(4))
        for i in range(3):
            label[0][i+1]=tk.Label(frame,text=text2[i])
            label[0][i+1].grid(row=0,column=i+1)
        for i in range(Pamount):
            label[i + 1][0]=tk.Label(frame,text='通道{}'.format(i+1))
            label[i+1][0].grid(row=i+1,column=0)
            label[i + 1][1]=tk.Label(frame,text=round(x[i],3))
            label[i+1][1].grid(row=i+1,column=1)
            label[i + 1][2]=tk.Label(frame,text=round(y[i],3))
            label[i+1][2].grid(row=i+1,column=2)
            label[i + 1][3]=tk.Label(frame,text=0)
            label[i+1][3].grid(row=i+1,column=3)
    if mode==2:
        data=list(range(len(numname)))
        x=list(range(len(numname)))
        y=list(range(len(numname)))
        for i in range(len(numname)):
            Initialdata=read_data(fpath_name[i])
            pre_data = clean_data(Initialdata)
            data[i] = pre_data[0]
            x[i] = pre_data[1]
            y[i] = pre_data[2]
        text2 = [ '零点偏移量', '时间趋势项大小', '异常值个数']
        frame = tk.Frame(window,width=1920,height=1080)
        frame.grid(row=0, column=0)
        frame.grid_propagate(0)
        label = list(range(Pamount[0] + 1))
        for i in range(Pamount[0] + 1):
            label[i] = list(range(4))
        for i in range(3):
            label[0][i+1]=tk.Label(frame,text=text2[i])
            label[0][i+1].grid(row=0,column=i+1)
        for i in range(Pamount[0]):
            label[i + 1][0]=tk.Label(frame,text='通道{}'.format(i+1))
            label[i+1][0].grid(row=i+1,column=0)
            label[i + 1][1]=tk.Label(frame,text=round(x[0][i],3))
            label[i+1][1].grid(row=i+1,column=1)
            label[i + 1][2]=tk.Label(frame,text=round(y[0][i],3))
            label[i+1][2].grid(row=i+1,column=2)
            label[i + 1][3]=tk.Label(frame,text=0)
            label[i+1][3].grid(row=i+1,column=3)
        pre_num=0
        def show(event):
            a=var.get()
            for i in range(len(numname)):
                if a==numname[i]:
                    pre_num=i
            for i in range(Pamount[pre_num]):
                label[i + 1][1].config(text=round(x[pre_num][i], 3))
                label[i + 1][2].config(text=round(y[pre_num][i], 3))
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = numname
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
    if mode==3:
        global x_tj,y_tj
        data=list(range(len(numname)-1))
        x=list(range(len(numname)-1))
        y=list(range(len(numname)-1))
        for i in range(len(numname)-1):
            Initialdata=read_data(fpath_name[i])
            pre_data = clean_data(Initialdata)
            data[i] = pre_data[0]
            x[i] = pre_data[1]
            y[i] = pre_data[2]
        x_tj=list(range(Pamount[0]))
        y_tj=list(range(Pamount[0]))
        for i in range(Pamount[0]):
            x_tj[i] = list(range(3))
            y_tj[i] = list(range(3))
            x_tj[i][0]=round(min(x[0][i],x[1][i],x[2][i],x[3][i],x[4][i]),3)
            sum1=0
            for j in range(len(numname)-1):
                sum1+=x[j][i]
            x_tj[i][1]=round(sum1/(len(numname)-1),3)
            x_tj[i][2]=round(max(x[0][i],x[1][i],x[2][i],x[3][i],x[4][i]),3)
            y_tj[i][0]=round(min(y[0][i],y[1][i],y[2][i],y[3][i],y[4][i]),3)
            sum2=0
            for j in range(len(numname)-1):
                sum2+=y[j][i]
            y_tj[i][1]=round(sum2/(len(numname)-1),3)
            y_tj[i][2]=round(max(y[0][i],y[1][i],y[2][i],y[3][i],y[4][i]),3)
        text2 = [ '零点偏移量', '时间趋势项大小', '异常值个数']
        frame = tk.Frame(window,width=1920,height=1080)
        frame.grid(row=0, column=0)
        frame.grid_propagate(0)
        label = list(range(Pamount[0] + 1))
        for i in range(Pamount[0] + 1):
            label[i] = list(range(4))
        for i in range(3):
            label[0][i+1]=tk.Label(frame,text=text2[i])
            label[0][i+1].grid(row=0,column=i+1)
        for i in range(Pamount[0]):
            label[i + 1][0]=tk.Label(frame,text='通道{}'.format(i+1))
            label[i+1][0].grid(row=i+1,column=0)
            label[i + 1][1]=tk.Label(frame,text=round(x[0][i],3))
            label[i+1][1].grid(row=i+1,column=1)
            label[i + 1][2]=tk.Label(frame,text=round(y[0][i],3))
            label[i+1][2].grid(row=i+1,column=2)
            label[i + 1][3]=tk.Label(frame,text=0)
            label[i+1][3].grid(row=i+1,column=3)
        pre_num=0
        def show(event):
            a=var.get()
            for i in range(len(numname)):
                if a==numname[i]:
                    pre_num=i
            if pre_num<=len(numname)-2:
                for i in range(3):
                    label[0][i + 1].config(text=text2[i])
                for i in range(3):
                    label[0][i+1].config(text=text2[i])
                for i in range(Pamount[pre_num]):
                    label[i + 1][1].config(text=round(x[pre_num][i], 3))
                    label[i + 1][2].config(text=round(y[pre_num][i], 3))
            if pre_num == len(numname) - 1:
                label[0][1].config(text='零点偏移量(最小值、均值、最大值）')
                label[0][2].config(text='时间趋势项大小(最小值、均值、最大值）')
                label[0][3].config(text='异常值个数(最小值、均值、最大值）')
                for i in range(Pamount[pre_num-1]):
                    label[i + 1][1].config(text=x_tj[i])
                    label[i + 1][2].config(text=y_tj[i])
                    label[i + 1][3].config(text=[0,0,0])
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = numname
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
def prepare_1():
    global data,x,y
    if mode==1:
        Initialdata = read_data(fpath_name)
        pre_data=clean_data_0(Initialdata)
        data = pre_data[0]
        x = pre_data[1]
        y = pre_data[2]
        text2=['零点偏移量','时间趋势项大小','异常值个数']
        frame=tk.Frame(window,width=1920,height=1080)
        frame.grid(row=0,column=0)
        frame.grid_propagate(0)
        def show(event):
            pre_num=var.get()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var,width=9)
        combobox['value'] = ['Group1']
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0,column=0)
        label=list(range(Pamount+1))
        for i in range(Pamount+1):
            label[i]=list(range(4))
        for i in range(3):
            label[0][i+1]=tk.Label(frame,text=text2[i])
            label[0][i+1].grid(row=0,column=i+1)
        for i in range(Pamount):
            label[i + 1][0]=tk.Label(frame,text='通道{}'.format(i+1))
            label[i+1][0].grid(row=i+1,column=0)
            label[i + 1][1]=tk.Label(frame,text=round(x[i],3))
            label[i+1][1].grid(row=i+1,column=1)
            label[i + 1][2]=tk.Label(frame,text=round(y[i],3))
            label[i+1][2].grid(row=i+1,column=2)
            label[i + 1][3]=tk.Label(frame,text=0)
            label[i+1][3].grid(row=i+1,column=3)
    if mode==2:
        data=list(range(len(numname)))
        x=list(range(len(numname)))
        y=list(range(len(numname)))
        for i in range(len(numname)):
            Initialdata=read_data(fpath_name[i])
            pre_data = clean_data_0(Initialdata)
            data[i] = pre_data[0]
            x[i] = pre_data[1]
            y[i] = pre_data[2]
        text2 = [ '零点偏移量', '时间趋势项大小', '异常值个数']
        frame = tk.Frame(window,width=1920,height=1080)
        frame.grid(row=0, column=0)
        frame.grid_propagate(0)
        label = list(range(Pamount[0] + 1))
        for i in range(Pamount[0] + 1):
            label[i] = list(range(4))
        for i in range(3):
            label[0][i+1]=tk.Label(frame,text=text2[i])
            label[0][i+1].grid(row=0,column=i+1)
        for i in range(Pamount[0]):
            label[i + 1][0]=tk.Label(frame,text='通道{}'.format(i+1))
            label[i+1][0].grid(row=i+1,column=0)
            label[i + 1][1]=tk.Label(frame,text=round(x[0][i],3))
            label[i+1][1].grid(row=i+1,column=1)
            label[i + 1][2]=tk.Label(frame,text=round(y[0][i],3))
            label[i+1][2].grid(row=i+1,column=2)
            label[i + 1][3]=tk.Label(frame,text=0)
            label[i+1][3].grid(row=i+1,column=3)
        pre_num=0
        def show(event):
            a=var.get()
            for i in range(len(numname)):
                if a==numname[i]:
                    pre_num=i
            for i in range(Pamount[pre_num]):
                label[i + 1][1].config(text=round(x[pre_num][i], 3))
                label[i + 1][2].config(text=round(y[pre_num][i], 3))
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = numname
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
    if mode==3:
        global x_tj,y_tj
        data=list(range(len(numname)-1))
        x=list(range(len(numname)-1))
        y=list(range(len(numname)-1))
        for i in range(len(numname)-1):
            Initialdata=read_data(fpath_name[i])
            pre_data = clean_data_0(Initialdata)
            data[i] = pre_data[0]
            x[i] = pre_data[1]
            y[i] = pre_data[2]
        x_tj=list(range(Pamount[0]))
        y_tj=list(range(Pamount[0]))
        for i in range(Pamount[0]):
            x_tj[i] = list(range(3))
            y_tj[i] = list(range(3))
            x_tj[i][0]=round(min(x[0][i],x[1][i],x[2][i],x[3][i],x[4][i]),3)
            sum1=0
            for j in range(len(numname)-1):
                sum1+=x[j][i]
            x_tj[i][1]=round(sum1/(len(numname)-1),3)
            x_tj[i][2]=round(max(x[0][i],x[1][i],x[2][i],x[3][i],x[4][i]),3)
            y_tj[i][0]=round(min(y[0][i],y[1][i],y[2][i],y[3][i],y[4][i]),3)
            sum2=0
            for j in range(len(numname)-1):
                sum2+=y[j][i]
            y_tj[i][1]=round(sum2/(len(numname)-1),3)
            y_tj[i][2]=round(max(y[0][i],y[1][i],y[2][i],y[3][i],y[4][i]),3)
        text2 = [ '零点偏移量', '时间趋势项大小', '异常值个数']
        frame = tk.Frame(window,width=1920,height=1080)
        frame.grid(row=0, column=0)
        frame.grid_propagate(0)
        label = list(range(Pamount[0] + 1))
        for i in range(Pamount[0] + 1):
            label[i] = list(range(4))
        for i in range(3):
            label[0][i+1]=tk.Label(frame,text=text2[i])
            label[0][i+1].grid(row=0,column=i+1)
        for i in range(Pamount[0]):
            label[i + 1][0]=tk.Label(frame,text='通道{}'.format(i+1))
            label[i+1][0].grid(row=i+1,column=0)
            label[i + 1][1]=tk.Label(frame,text=round(x[0][i],3))
            label[i+1][1].grid(row=i+1,column=1)
            label[i + 1][2]=tk.Label(frame,text=round(y[0][i],3))
            label[i+1][2].grid(row=i+1,column=2)
            label[i + 1][3]=tk.Label(frame,text=0)
            label[i+1][3].grid(row=i+1,column=3)
        pre_num=0
        def show(event):
            a=var.get()
            for i in range(len(numname)):
                if a==numname[i]:
                    pre_num=i
            if pre_num<=len(numname)-2:
                for i in range(3):
                    label[0][i + 1].config(text=text2[i])
                for i in range(3):
                    label[0][i+1].config(text=text2[i])
                for i in range(Pamount[pre_num]):
                    label[i + 1][1].config(text=round(x[pre_num][i], 3))
                    label[i + 1][2].config(text=round(y[pre_num][i], 3))
            if pre_num == len(numname) - 1:
                label[0][1].config(text='零点偏移量(最小值、均值、最大值）')
                label[0][2].config(text='时间趋势项大小(最小值、均值、最大值）')
                label[0][3].config(text='异常值个数(最小值、均值、最大值）')
                for i in range(Pamount[pre_num-1]):
                    label[i + 1][1].config(text=x_tj[i])
                    label[i + 1][2].config(text=y_tj[i])
                    label[i + 1][3].config(text=[0,0,0])
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = numname
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
def vltime():
    global vlmax,vlarms,val
    if mode==1:
        vlmax=max_data(data)
        vlarms=arms_data(data)
        val=VAL_data(vlarms)
        for i in range(Pamount):
            plt.figure(figsize=(8,0.9))
            plt.plot(data[i].time,data[i].a)
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
            plt.savefig("cache\\"+dirname+"\\1处理后时程图-通道{}".format(i+1)+".png")
            plt.close()
        for i in range(Pamount):
            plt.figure()
            plt.plot(data[i].time, data[i].a)
            plt.xlabel("时间", fontsize=14)
            plt.ylabel("加速度/ms-2", fontsize=14)
            plt.xticks(fontsize=14)
            plt.yticks(fontsize=14)
            plt.savefig("result\\" + dirname + "\\1处理后时程图-通道{}".format(i + 1) + ".png")
            plt.close()
        frame=tk.Frame(window,width=1920,height=1080)
        frame.grid(row=0,column=0)
        frame.grid_propagate(0)
        def show(event):
            pre_num=var.get()
        def show_image31():
            plt.figure()
            for i in range(Pamount):
                plt.subplot(Pamount,1,i+1)
                plt.plot(data[i].time,data[i].a)
                plt.ylabel("加速度/ms-2", fontsize=14)
                if i ==Pamount-1:
                    plt.xlabel("时间", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
            plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var,width=9)
        combobox['value'] = ['Group1']
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0,column=0)
        label0=tk.Label(frame,text='时域分析结果')
        label0.grid(row=0,column=1,columnspan=2)
        label1=tk.Label(frame,text='时程图')
        label1.grid(row=0,column=3)
        buttun=tk.Button(frame,text='显示详图',command=show_image31)
        buttun.grid(row=0,column=4)
        text3=['加速度最值/ms-2','加速度有效值/ms-2','VAL/dB']
        label = list(range(Pamount*3))
        img=list(range(Pamount))
        photo=list(range(Pamount))
        for i in range(Pamount*3):
            label[i] = list(range(4))
        for i in range(Pamount):
            label[i][0]=tk.Label(frame,text='通道{}'.format(i+1))
            label[i][0].grid(row=1+i*3,column=0,rowspan=3)
            label[i][1]=list(range(3))
            label[i][2]=list(range(3))
            for j in range(3):
                label[i][1][j]=tk.Label(frame,text=text3[j])
                label[i][1][j].grid(row=1+i*3+j,column=1)
            label[i][2][0]=tk.Label(frame,text=round(vlmax[i],3))
            label[i][2][0].grid(row=1+i*3+0,column=2)
            label[i][2][1]=tk.Label(frame,text=round(vlarms[i],3))
            label[i][2][1].grid(row=1+i*3+1,column=2)
            label[i][2][2]=tk.Label(frame,text=round(val[i],0))
            label[i][2][2].grid(row=1+i*3+2,column=2)
            label[i][3]=tk.Label(frame)
            img[i] = Image.open("cache\\"+dirname+"\\1处理后时程图-通道{}".format(i+1)+".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            label[i][3].config(image=photo[i])
            label[i][3].image=photo[i]
            label[i][3].grid(row=1+i*3,column=3,rowspan=3,columnspan=2)
    if mode==2:
        vlmax = list(range(len(numname)))
        vlarms = list(range(len(numname)))
        val = list(range(len(numname)))
        for i in range(len(numname)):
            vlmax[i] = max_data(data[i])
            vlarms[i] = arms_data(data[i])
            val[i] = VAL_data(vlarms[i])
            for j in range(Pamount[i]):
                plt.figure(figsize=(8, 0.9))
                plt.plot(data[i][j].time, data[i][j].a)
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname + "\\" + numname[i] + "\\1处理后时程图-通道{}".format(j + 1) + ".png")
                plt.close()
        for i in range(len(numname)):
            for j in range(Pamount[i]):
                plt.figure()
                plt.plot(data[i][j].time, data[i][j].a)
                plt.xlabel("时间", fontsize=14)
                plt.ylabel("加速度/ms-2", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname + "\\" + numname[i]+ "\\1处理后时程图-通道{}".format(j + 1) + ".png")
                plt.close()
        frame=tk.Frame(window,width=1920,height=1080)
        frame.grid(row=0,column=0)
        frame.grid_propagate(0)
        label = list(range(Pamount[0] * 3))
        img = list(range(Pamount[0]))
        photo = list(range(Pamount[0]))
        for i in range(Pamount[0] * 3):
            label[i] = list(range(4))
        text3=['加速度最值/ms-2','加速度有效值/ms-2','VAL/dB']
        for i in range(Pamount[0]):
            label[i][0] = tk.Label(frame, text='通道{}'.format(i + 1))
            label[i][0].grid(row=1 + i * 3, column=0, rowspan=3)
            label[i][1] = list(range(3))
            label[i][2] = list(range(3))
            for j in range(3):
                label[i][1][j] = tk.Label(frame, text=text3[j])
                label[i][1][j].grid(row=1 + i * 3 + j, column=1)
            label[i][2][0]=tk.Label(frame,text=round(vlmax[0][i],3))
            label[i][2][0].grid(row=1+i*3+0,column=2)
            label[i][2][1]=tk.Label(frame,text=round(vlarms[0][i],3))
            label[i][2][1].grid(row=1+i*3+1,column=2)
            label[i][2][2]=tk.Label(frame,text=round(val[0][i],0))
            label[i][2][2].grid(row=1+i*3+2,column=2)
            label[i][3] = tk.Label(frame)
            img[i] = Image.open("cache\\" + dirname + "\\" + numname[0] + "\\1处理后时程图-通道{}".format(j + 1) + ".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            label[i][3].config(image=photo[i])
            label[i][3].image = photo[i]
            label[i][3].grid(row=1 + i * 3, column=3, rowspan=3, columnspan=2)
        pre_num=0
        def show(event):
            a=var.get()
            for i in range(len(numname)):
                if a==numname[i]:
                    pre_num=i
            for i in range(Pamount[pre_num]):
                label[i][2][0].config(text=round(vlmax[pre_num][i],3))
                label[i][2][1].config(text=round(vlarms[pre_num][i],3))
                label[i][2][2].config(text=round(val[pre_num][i],0))
                img[i] = Image.open("cache\\" + dirname + "\\" + numname[pre_num] + "\\1处理后时程图-通道{}".format(j + 1) + ".png")  # 打开图片
                photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                label[i][3].config(image=photo[i])
                label[i][3].image = photo[i]
                label[i][3].grid(row=1 + i * 3, column=3, rowspan=3, columnspan=2)
        def show_image32():
            a=var.get()
            for i in range(len(numname)):
                if a==numname[i]:
                    pre_num=i
            plt.figure(num=numname[pre_num])
            for i in range(Pamount[pre_num]):
                plt.subplot(Pamount[pre_num],1,i+1)
                plt.plot(data[pre_num][i].time, data[pre_num][i].a)
                if j ==Pamount[pre_num]-1:
                    plt.xlabel("时间", fontsize=14)
                plt.ylabel("加速度/ms-2", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
            plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var,width=9)
        combobox['value'] = numname
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0,column=0)
        label0=tk.Label(frame,text='时域分析结果')
        label0.grid(row=0,column=1,columnspan=2)
        label1=tk.Label(frame,text='时程图')
        label1.grid(row=0,column=3)
        buttun=tk.Button(frame,text='显示详图',command=show_image32)
        buttun.grid(row=0,column=4)
    if mode==3:
        global vlmax_tj,vlarms_tj,val_tj
        vlmax = list(range(len(numname)-1))
        vlarms = list(range(len(numname)-1))
        val = list(range(len(numname)-1))
        vlmax_tj=list(range(Pamount[0]))
        vlarms_tj=list(range(Pamount[0]))
        val_tj=list(range(Pamount[0]))
        for i in range(len(numname)-1):
            vlmax[i] = max_data(data[i])
            vlarms[i] = arms_data(data[i])
            val[i] = VAL_data(vlarms[i])
            for j in range(Pamount[i]):
                plt.figure(figsize=(8, 0.9))
                plt.plot(data[i][j].time, data[i][j].a)
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname + "\\" + numname[i] + "\\1处理后时程图-通道{}".format(j + 1) + ".png")
                plt.close()
        for i in range(len(numname)-1):
            for j in range(Pamount[i]):
                plt.figure()
                plt.plot(data[i][j].time, data[i][j].a)
                plt.xlabel("时间", fontsize=14)
                plt.ylabel("加速度/ms-2", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname + "\\" + numname[i]+ "\\1处理后时程图-通道{}".format(j + 1) + ".png")
                plt.close()
        for i in range(Pamount[0]):
            vlmax_tj[i] = list(range(3))
            vlarms_tj[i] = list(range(3))
            val_tj[i] = list(range(3))
            vlmax_tj[i][0]=round(min(vlmax[0][i],vlmax[1][i],vlmax[2][i],vlmax[3][i],vlmax[4][i]),3)
            sum1=0
            for j in range(len(numname)-1):
                sum1+=vlmax[j][i]
            vlmax_tj[i][1]=round(sum1/(len(numname)-1),3)
            vlmax_tj[i][2]=round(max(vlmax[0][i],vlmax[1][i],vlmax[2][i],vlmax[3][i],vlmax[4][i]),3)
            vlarms_tj[i][0]=round(min(vlarms[0][i],vlarms[1][i],vlarms[2][i],vlarms[3][i],vlarms[4][i]),3)
            sum2=0
            for j in range(len(numname)-1):
                sum2+=vlarms[j][i]
            vlarms_tj[i][1]=round(sum2/(len(numname)-1),3)
            vlarms_tj[i][2]=round(max(vlarms[0][i],vlarms[1][i],vlarms[2][i],vlarms[3][i],vlarms[4][i]),3)
            val_tj[i][0]=round(min(val[0][i],val[1][i],val[2][i],val[3][i],val[4][i]),3)
            sum3=0
            for j in range(len(numname)-1):
                sum3+=val[j][i]
            val_tj[i][1]=round(sum3/(len(numname)-1),3)
            val_tj[i][2]=round(max(val[0][i],val[1][i],val[2][i],val[3][i],val[4][i]),0)
        a=list(range(Pamount[0]))
        b=list(range(Pamount[0]))
        c=list(range(Pamount[0]))
        for i in range(Pamount[0]):
            a[i]=[]
            b[i]=[]
            c[i]=[]
        for i in range(Pamount[0]):
            for j in range(len(numname)-1):
                a[i].append(vlmax[j][i])
                b[i].append(vlarms[j][i])
                c[i].append(val[j][i])
        for i in range(Pamount[0]):
            plt.figure(figsize=(8, 0.9))
            plt.subplot(1,3,1)
            plt.title('max')
            plt.plot(numname[0:5],a[i])
            plt.subplot(1,3,2)
            plt.title('arms')
            plt.plot(numname[0:5],b[i])
            plt.subplot(1,3,3)
            plt.title('val')
            plt.plot(numname[0:5],c[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.27, top=0.78,wspace=0.25)
            plt.savefig("cache\\" + dirname  + "\\1时域结果对比分析-通道{}".format(i + 1) + ".png")
            plt.close()
            plt.figure(figsize=(8, 0.9))
            plt.subplot(1, 3, 1)
            plt.title('max')
            plt.plot(numname[0:5], a[i])
            plt.subplot(1, 3, 2)
            plt.title('arms')
            plt.plot(numname[0:5], b[i])
            plt.subplot(1, 3, 3)
            plt.title('val')
            plt.plot(numname[0:5], c[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.27, top=0.78, wspace=0.25)
            plt.savefig("result\\" + dirname + "\\1时域结果对比分析-通道{}".format(i + 1) + ".png")
            plt.close()
        frame=tk.Frame(window,width=1920,height=1080)
        frame.grid(row=0,column=0)
        frame.grid_propagate(0)
        label = list(range(Pamount[0] * 3))
        img = list(range(Pamount[0]))
        photo = list(range(Pamount[0]))
        for i in range(Pamount[0] * 3):
            label[i] = list(range(4))
        text3=['加速度最值/ms-2','加速度有效值/ms-2','VAL/dB']
        for i in range(Pamount[0]):
            label[i][0] = tk.Label(frame, text='通道{}'.format(i + 1))
            label[i][0].grid(row=1 + i * 3, column=0, rowspan=3)
            label[i][1] = list(range(3))
            label[i][2] = list(range(3))
            for j in range(3):
                label[i][1][j] = tk.Label(frame, text=text3[j])
                label[i][1][j].grid(row=1 + i * 3 + j, column=1)
            label[i][2][0]=tk.Label(frame,text=round(vlmax[0][i],3))
            label[i][2][0].grid(row=1+i*3+0,column=2)
            label[i][2][1]=tk.Label(frame,text=round(vlarms[0][i],3))
            label[i][2][1].grid(row=1+i*3+1,column=2)
            label[i][2][2]=tk.Label(frame,text=round(val[0][i],0))
            label[i][2][2].grid(row=1+i*3+2,column=2)
            label[i][3] = tk.Label(frame)
            img[i] = Image.open("cache\\" + dirname + "\\" + numname[0] + "\\1处理后时程图-通道{}".format(i + 1) + ".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            label[i][3].config(image=photo[i])
            label[i][3].image = photo[i]
            label[i][3].grid(row=1 + i * 3, column=3, rowspan=3, columnspan=2)
        pre_num = 0
        def show(event):
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            if pre_num <= len(numname) - 2:
                label1.config(text='时程图')
                for i in range(Pamount[0]):
                    for j in range(3):
                        label[i][1][j].config(text=text3[j])
                for i in range(Pamount[pre_num]):
                    label[i][2][0].config(text=round(vlmax[pre_num][i],3))
                    label[i][2][1].config(text=round(vlarms[pre_num][i],3))
                    label[i][2][2].config(text=round(val[pre_num][i],3))
                    img[i] = Image.open(
                        "cache\\" + dirname + "\\" + numname[pre_num] + "\\1处理后时程图-通道{}".format(i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    label[i][3].config(image=photo[i])
                    label[i][3].image = photo[i]
            if pre_num == len(numname) - 1:
                label1.config(text='对比分析图')
                for i in range(Pamount[0]):
                    label[i][1][0].config(text='最值/ms-2(最小值、均值、最大值）')
                    label[i][1][1].config(text='有效值/ms-2(最小值、均值、最大值）')
                    label[i][1][2].config(text='val/ms-2(最小值、均值、最大值）')
                for i in range(Pamount[pre_num - 1]):
                    label[i][2][0].config(text=vlmax_tj[i])
                    label[i][2][1].config(text=vlarms_tj[i])
                    label[i][2][2].config(text=val_tj[i])
                    img[i] = Image.open("cache\\" + dirname  + "\\1时域结果对比分析-通道{}".format(i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    label[i][3].config(image=photo[i])
                    label[i][3].image = photo[i]
        def show_image33():
            a=var.get()
            for i in range(len(numname)):
                if a==numname[i]:
                    pre_num=i
            if pre_num == len(numname) - 1:
                plt.figure(num=numname[pre_num])
                for i in range(Pamount[0]):
                    plt.subplot(Pamount[0], 3, 1+i*3)
                    plt.title('max')
                    plt.plot(numname[0:5], a[i])
                    plt.subplot(Pamount[0], 3, 2+i*3)
                    plt.title('arms')
                    plt.plot(numname[0:5], b[i])
                    plt.subplot(Pamount[0], 3, 3+i*3)
                    plt.title('val')
                    plt.plot(numname[0:5], c[i])
                plt.show()
            else:
                plt.figure(num=numname[pre_num])
                for i in range(Pamount[pre_num]):
                    plt.subplot(Pamount[pre_num], 1, i + 1)
                    plt.plot(data[pre_num][i].time, data[pre_num][i].a)
                    if j == Pamount[pre_num] - 1:
                        plt.xlabel("时间", fontsize=14)
                    plt.ylabel("加速度/ms-2", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = numname
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
        label0 = tk.Label(frame, text='时域分析结果')
        label0.grid(row=0, column=1, columnspan=2)
        label1 = tk.Label(frame, text='时程图')
        label1.grid(row=0, column=3)
        buttun = tk.Button(frame, text='显示详图', command=show_image33)
        buttun.grid(row=0, column=4)
def fre():
    global obs_fre,max_am,max_power
    if mode==1:
        fre=list(range(Pamount))
        amplitude=list(range(Pamount))
        power=list(range(Pamount))
        obs_fre=list(range(Pamount))
        max_am=list(range(Pamount))
        max_power=list(range(Pamount))
        for i in range(Pamount):
            fre[i]=fft_vis(data[i])[0]
            amplitude[i]=fft_vis(data[i])[1]
            power[i]=fft_vis(data[i])[2]
            obs_fre[i]=fft_vis(data[i])[3]
            max_am[i]=fft_vis(data[i])[4]
            max_power[i]=fft_vis(data[i])[5]
        for i in range(Pamount):
            plt.figure(figsize=(8,0.9))
            plt.plot(fre[i], amplitude[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
            plt.savefig("cache\\"+dirname+"\\2-1频谱-通道{}".format(i+1)+".png")
            plt.close()
            plt.figure(figsize=(8,0.9))
            plt.plot(fre[i], power[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
            plt.savefig("cache\\"+dirname+"\\2-2功率谱-通道{}".format(i + 1) + ".png")
            plt.close()
        for i in range(Pamount):
            plt.figure()
            plt.plot(fre[i], amplitude[i])
            plt.xlabel("频率/Hz", fontsize=14)
            plt.ylabel("振幅/ms-2", fontsize=14)
            plt.xticks(fontsize=14)
            plt.yticks(fontsize=14)
            plt.savefig("result\\" + dirname + "\\2-1频谱-通道{}".format(i + 1) + ".png")
            plt.close()
            plt.figure()
            plt.plot(fre[i], power[i])
            plt.xlabel("频率/Hz", fontsize=14)
            plt.ylabel("功率/w", fontsize=14)
            plt.xticks(fontsize=14)
            plt.yticks(fontsize=14)
            plt.savefig("result\\" + dirname + "\\2-2功率谱-通道{}".format(i + 1) + ".png")
            plt.close()
        frame=tk.Frame(window,width=1920,height=1080)
        frame.grid(row=0,column=0)
        frame.grid_propagate(0)
        def show(event):
            pre_num = var.get()
        def show_change41():
            if label1["text"]=='频谱':
                label1.config(text='功率谱')
                for i in range(Pamount):
                    img[i] = Image.open("cache\\"+dirname+"\\2-2功率谱-通道{}".format(i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
            else:
                label1.config(text='频谱')
                for i in range(Pamount):
                    img[i] = Image.open("cache\\"+dirname+"\\2-1频谱-通道{}".format(i+1)+".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
        def show_image41():
            plt.figure()
            for i in range(Pamount):
                plt.subplot(Pamount,2,i*2+1)
                plt.plot(fre[i], amplitude[i])
                plt.ylabel("振幅", fontsize=14)
                if i ==Pamount-1:
                    plt.xlabel("频率", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.subplot(Pamount,2,i*2+2)
                plt.plot(fre[i], power[i])
                plt.ylabel("频率", fontsize=14)
                if i ==Pamount-1:
                    plt.xlabel("功率", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
            plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var,width=9)
        combobox['value'] = ['Group1']
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0,column=0)
        label0 = tk.Label(frame, text='频域分析结果')
        label0.grid(row=0, column=1, columnspan=2)
        buttun1 = tk.Button(frame, text='频谱/功率谱切换', command=show_change41)
        buttun1.grid(row=0, column=3)
        label1 = tk.Label(frame, text='频谱')
        label1.grid(row=0, column=4)
        buttun2 = tk.Button(frame, text='显示详图', command=show_image41)
        buttun2.grid(row=0, column=5)
        labelc1=list(range(Pamount))
        for i in range(Pamount):
            labelc1[i]=tk.Label(frame,text='通道{}'.format(i+1))
            labelc1[i].grid(row=1+i*3,column=0,rowspan=3)
        labelc2=list(range(Pamount*3))
        for i in range(Pamount):
            labelc2[0+3*i]=tk.Label(frame,text='振动显著频率/Hz')
            labelc2[0 + 3 * i].grid(row=1+i*3+0,column=1)
            labelc2[1+3*i]=tk.Label(frame,text='显著频率对应振幅值/ms-2')
            labelc2[1 + 3 * i].grid(row=1+i*3+1,column=1)
            labelc2[2+3*i]=tk.Label(frame,text='功率最值/w')
            labelc2[2 + 3 * i].grid(row=1+i*3+2,column=1)
        labelc3=list(range(Pamount*3))
        for i in range(Pamount):
            labelc3[0+3*i]=tk.Label(frame,text=obs_fre[i])
            labelc3[0 + 3 * i].grid(row=1+i*3+0,column=2)
            labelc3[1+3*i]=tk.Label(frame,text=max_am[i])
            labelc3[1 + 3 * i].grid(row=1+i*3+1,column=2)
            labelc3[2+3*i]=tk.Label(frame,text=round(max_power[i],3))
            labelc3[2 + 3 * i].grid(row=1+i*3+2,column=2)
        labelc4=list(range(Pamount))
        img=list(range(Pamount))
        photo=list(range(Pamount))
        for i in range(Pamount):
            labelc4[i]=tk.Label(frame)
            img[i] = Image.open("cache\\"+dirname+"\\2-1频谱-通道{}".format(i+1)+".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            labelc4[i].config(image=photo[i])
            labelc4[i].image=photo[i]
            labelc4[i].grid(row=1 + i * 3, column=3, rowspan=3, columnspan=3)
    if mode==2:
        fre = list(range(len(numname)))
        amplitude = list(range(len(numname)))
        power = list(range(len(numname)))
        obs_fre = list(range(len(numname)))
        max_am = list(range(len(numname)))
        max_power = list(range(len(numname)))
        for i in range(len(numname)):
            fre[i] = list(range(Pamount[i]))
            amplitude[i] = list(range(Pamount[i]))
            power[i] = list(range(Pamount[i]))
            obs_fre[i] = list(range(Pamount[i]))
            max_am[i] = list(range(Pamount[i]))
            max_power[i] = list(range(Pamount[i]))
        for i in range(len(numname)):
            for j in range(Pamount[i]):
                fre[i][j] = fft_vis(data[i][j])[0]
                amplitude[i][j] = fft_vis(data[i][j])[1]
                power[i][j] = fft_vis(data[i][j])[2]
                obs_fre[i][j] = fft_vis(data[i][j])[3]
                max_am[i][j] = fft_vis(data[i][j])[4]
                max_power[i][j] = fft_vis(data[i][j])[5]
        for i in range(len(numname)):
            for j in range(Pamount[i]):
                plt.figure(figsize=(8, 0.9))
                plt.plot(fre[i][j], amplitude[i][j])
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname + "\\"+numname[i]+"\\2-1频谱-通道{}".format(j + 1) + ".png")
                plt.close()
                plt.figure(figsize=(8, 0.9))
                plt.plot(fre[i][j], power[i][j])
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname +"\\"+numname[i]+ "\\2-2功率谱-通道{}".format(j + 1) + ".png")
                plt.close()
        for i in range(len(numname)):
            for j in range(Pamount[i]):
                plt.figure()
                plt.plot(fre[i][j], amplitude[i][j])
                plt.xlabel("频率/Hz", fontsize=14)
                plt.ylabel("振幅/ms-2", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname + "\\" + numname[i] + "\\2-1频谱-通道{}".format(j + 1) + ".png")
                plt.close()
                plt.figure()
                plt.plot(fre[i][j], power[i][j])
                plt.xlabel("频率/Hz", fontsize=14)
                plt.ylabel("功率/w", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname + "\\" + numname[i] + "\\2-2功率谱-通道{}".format(j + 1) + ".png")
                plt.close()
        frame=tk.Frame(window,width=1920,height=1080)
        frame.grid(row=0,column=0)
        frame.grid_propagate(0)
        pre_num=0
        def show(event):
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            label1.config(text='频谱')
            for i in range(Pamount[0]):
                labelc3[0 + 3 * i] .config(text=obs_fre[pre_num][i])
                labelc3[1 + 3 * i] .config(text=max_am[pre_num][i])
                labelc3[2 + 3 * i] .config(text=round(max_power[pre_num][i], 3))
            for i in range(Pamount[0]):
                img[i] = Image.open(
                    "cache\\" + dirname + "\\" + numname[pre_num] + "\\2-1频谱-通道{}".format(i + 1) + ".png")  # 打开图片
                photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                labelc4[i].config(image=photo[i])
                labelc4[i].image = photo[i]
        def show_change42():
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            if label1["text"]=='频谱':
                label1.config(text='功率谱')
                for i in range(Pamount[pre_num]):
                    img[i] = Image.open("cache\\"+dirname+"\\" + numname[pre_num] +"\\2-2功率谱-通道{}".format(i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
            else:
                label1.config(text='频谱')
                for i in range(Pamount[pre_num]):
                    img[i] = Image.open("cache\\"+dirname+"\\" + numname[pre_num] +"\\2-1频谱-通道{}".format(i+1)+".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
        def show_image42():
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            plt.figure(num=numname[pre_num])
            for i in range(Pamount[pre_num]):
                plt.subplot(Pamount[pre_num],2,i*2+1)
                plt.plot(fre[pre_num][i], amplitude[pre_num][i])
                plt.ylabel("振幅", fontsize=14)
                if i ==Pamount[pre_num]-1:
                    plt.xlabel("频率", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.subplot(Pamount[pre_num],2,i*2+2)
                plt.plot(fre[pre_num][i], power[pre_num][i])
                plt.ylabel("频率", fontsize=14)
                if i ==Pamount[pre_num]-1:
                    plt.xlabel("功率", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
            plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = numname
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
        label0 = tk.Label(frame, text='频域分析结果')
        label0.grid(row=0, column=1, columnspan=2)
        buttun1 = tk.Button(frame, text='频谱/功率谱切换', command=show_change42)
        buttun1.grid(row=0, column=3)
        label1 = tk.Label(frame, text='频谱')
        label1.grid(row=0, column=4)
        buttun2 = tk.Button(frame, text='显示详图', command=show_image42)
        buttun2.grid(row=0, column=5)
        labelc1 = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc1[i] = tk.Label(frame, text='通道{}'.format(i + 1))
            labelc1[i].grid(row=1 + i * 3, column=0, rowspan=3)
        labelc2 = list(range(Pamount[0] * 3))
        for i in range(Pamount[0]):
            labelc2[0 + 3 * i] = tk.Label(frame, text='振动显著频率/Hz')
            labelc2[0 + 3 * i].grid(row=1 + i * 3 + 0, column=1)
            labelc2[1 + 3 * i] = tk.Label(frame, text='显著频率对应振幅值/ms-2')
            labelc2[1 + 3 * i].grid(row=1 + i * 3 + 1, column=1)
            labelc2[2 + 3 * i] = tk.Label(frame, text='功率最值/w')
            labelc2[2 + 3 * i].grid(row=1 + i * 3 + 2, column=1)
        labelc3 = list(range(Pamount[0] * 3))
        for i in range(Pamount[0]):
            labelc3[0 + 3 * i] = tk.Label(frame, text=obs_fre[0][i])
            labelc3[0 + 3 * i].grid(row=1 + i * 3 + 0, column=2)
            labelc3[1 + 3 * i] = tk.Label(frame, text=max_am[0][i])
            labelc3[1 + 3 * i].grid(row=1 + i * 3 + 1, column=2)
            labelc3[2 + 3 * i] = tk.Label(frame, text=round(max_power[0][i], 3))
            labelc3[2 + 3 * i].grid(row=1 + i * 3 + 2, column=2)
        labelc4 = list(range(Pamount[0]))
        img = list(range(Pamount[0]))
        photo = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc4[i] = tk.Label(frame)
            img[i] = Image.open("cache\\" + dirname + "\\"+numname[0]+"\\2-1频谱-通道{}".format(i + 1) + ".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            labelc4[i].config(image=photo[i])
            labelc4[i].image = photo[i]
            labelc4[i].grid(row=1 + i * 3, column=3, rowspan=3, columnspan=3)
    if mode==3:
        global am_tj,power_tj
        am_tj=list(range(Pamount[0]))
        power_tj=list(range(Pamount[0]))
        fre = list(range(len(numname)-1))
        amplitude = list(range(len(numname)-1))
        power = list(range(len(numname)-1))
        obs_fre = list(range(len(numname)-1))
        max_am = list(range(len(numname)-1))
        max_power = list(range(len(numname)-1))
        for i in range(len(numname)-1):
            fre[i] = list(range(Pamount[i]))
            amplitude[i] = list(range(Pamount[i]))
            power[i] = list(range(Pamount[i]))
            obs_fre[i] = list(range(Pamount[i]))
            max_am[i] = list(range(Pamount[i]))
            max_power[i] = list(range(Pamount[i]))
        for i in range(len(numname)-1):
            for j in range(Pamount[i]):
                fre[i][j] = fft_vis(data[i][j])[0]
                amplitude[i][j] = fft_vis(data[i][j])[1]
                power[i][j] = fft_vis(data[i][j])[2]
                obs_fre[i][j] = fft_vis(data[i][j])[3]
                max_am[i][j] = fft_vis(data[i][j])[4]
                max_power[i][j] = fft_vis(data[i][j])[5]
        for i in range(len(numname)-1):
            for j in range(Pamount[i]):
                plt.figure(figsize=(8, 0.9))
                plt.plot(fre[i][j], amplitude[i][j])
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname + "\\"+numname[i]+"\\2-1频谱-通道{}".format(j + 1) + ".png")
                plt.close()
                plt.figure(figsize=(8, 0.9))
                plt.plot(fre[i][j], power[i][j])
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname +"\\"+numname[i]+ "\\2-2功率谱-通道{}".format(j + 1) + ".png")
                plt.close()
        for i in range(len(numname)-1):
            for j in range(Pamount[i]):
                plt.figure()
                plt.plot(fre[i][j], amplitude[i][j])
                plt.xlabel("频率/Hz", fontsize=14)
                plt.ylabel("振幅/ms-2", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname + "\\" + numname[i] + "\\2-1频谱-通道{}".format(j + 1) + ".png")
                plt.close()
                plt.figure()
                plt.plot(fre[i][j], power[i][j])
                plt.xlabel("频率/Hz", fontsize=14)
                plt.ylabel("功率/w", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname + "\\" + numname[i] + "\\2-2功率谱-通道{}".format(j + 1) + ".png")
                plt.close()
        amp=list(range(Pamount[0]))
        pow=list(range(Pamount[0]))
        for i in range(Pamount[0]):
            am_tj[i]=list(range(3))
            power_tj[i]=list(range(3))
            amp[i]=[]
            pow[i]=[]
            for j in range(len(numname)-1):
                amp[i].append(max(max_am[j][i]))
                pow[i].append(max_power[j][i])
            am_tj[i][0]=min(max(max_am[0][i]),max(max_am[1][i]),max(max_am[2][i]),max(max_am[3][i]),max(max_am[4][i]))
            sum1=0
            for j in range(len(numname)-1):
                sum1+=max(max_am[j][i])
            am_tj[i][1]=round(sum1/(len(numname)-1),3)
            am_tj[i][2]=max(max(max_am[0][i]),max(max_am[1][i]),max(max_am[2][i]),max(max_am[3][i]),max(max_am[4][i]))
            power_tj[i][0]=min(max_power[0][i],max_power[1][i],max_power[2][i],max_power[3][i],max_power[4][i])
            sum2=0
            for j in range(len(numname)-1):
                sum2+=max_power[j][i]
            power_tj[i][1]=round(sum2/(len(numname)-1),3)
            power_tj[i][2]=max(max_power[0][i],max_power[1][i],max_power[2][i],max_power[3][i],max_power[4][i])
            plt.figure(figsize=(8, 0.9))
            plt.subplot(1,2,1)
            plt.title('振幅最值变化')
            plt.plot(numname[0:5], amp[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplot(1,2,2)
            plt.title('功率最值变化')
            plt.plot(numname[0:5], pow[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.27, top=0.75,wspace=0.25)
            plt.savefig("cache\\" + dirname + "\\2频域结果对比分析-通道{}".format(i + 1) + ".png")
            plt.close()
            plt.figure(figsize=(8, 0.9))
            plt.subplot(1, 2, 1)
            plt.title('振幅最值变化')
            plt.plot(numname[0:5], amp[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplot(1, 2, 2)
            plt.title('功率最值变化')
            plt.plot(numname[0:5], pow[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.27, top=0.75, wspace=0.25)
            plt.savefig("result\\" + dirname + "\\2频域结果对比分析-通道{}".format(i + 1) + ".png")
            plt.close()
        frame = tk.Frame(window, width=1920, height=1080)
        frame.grid(row=0, column=0)
        frame.grid_propagate(0)
        pre_num = 0
        def show(event):
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            if pre_num<=len(numname)-2:
                buttun1.grid(row=0, column=3)
                label1.config(text='频谱')
                for i in range(Pamount[0]):
                    labelc2[0 + 3 * i].config(text='振动显著频率/Hz')
                    labelc2[0 + 3 * i].grid(row=1 + i * 3 + 0, column=1)
                    labelc2[1 + 3 * i].config(text='显著频率对应振幅值/ms-2')
                    labelc2[2 + 3 * i].config(text='功率最值/w')
                    labelc3[0 + 3 * i].config(text=obs_fre[pre_num][i])
                    labelc3[0 + 3 * i].grid(row=1 + i * 3 + 0, column=2)
                    labelc3[1 + 3 * i].config(text=max_am[pre_num][i])
                    labelc3[2 + 3 * i].config(text=round(max_power[pre_num][i], 3))
                for i in range(Pamount[0]):
                    img[i] = Image.open("cache\\" + dirname + "\\"+numname[pre_num]+"\\2-1频谱-通道{}".format(i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
            if pre_num==len(numname)-1:
                buttun1.grid_forget()
                label1.config(text='对比分析结果')
                for i in range(Pamount[0]):
                    labelc3[0 + 3 * i].grid_forget()
                    labelc3[1 + 3 * i].config(text=am_tj[i])
                    labelc3[2 + 3 * i].config(text=power_tj[i])
                    labelc2[0 + 3 * i].grid_forget()
                    labelc2[1 + 3 * i].config(text='幅值最值（最小值、均值、最大值）/ms-2')
                    labelc2[2 + 3 * i].config(text='功率最值（最小值、均值、最大值）/w')
                for i in range(Pamount[0]):
                    img[i] = Image.open("cache\\" + dirname + "\\2频域结果对比分析-通道{}".format(i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
        def show_change43():
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            if pre_num<=len(numname)-2:
                if label1["text"] == '频谱':
                    label1.config(text='功率谱')
                    for i in range(Pamount[pre_num]):
                        img[i] = Image.open(
                            "cache\\" + dirname + "\\" + numname[pre_num] + "\\2-2功率谱-通道{}".format(i + 1) + ".png")  # 打开图片
                        photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                        labelc4[i].config(image=photo[i])
                        labelc4[i].image = photo[i]
                else:
                    label1.config(text='频谱')
                    for i in range(Pamount[pre_num]):
                        img[i] = Image.open(
                            "cache\\" + dirname + "\\" + numname[pre_num] + "\\2-1频谱-通道{}".format(i + 1) + ".png")  # 打开图片
                        photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                        labelc4[i].config(image=photo[i])
                        labelc4[i].image = photo[i]
        def show_image43():
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            if pre_num<=len(numname)-2:
                plt.figure(num=numname[pre_num])
                for i in range(Pamount[pre_num]):
                    plt.subplot(Pamount[pre_num], 2, i * 2 + 1)
                    plt.plot(fre[pre_num][i], amplitude[pre_num][i])
                    plt.ylabel("振幅", fontsize=14)
                    if i == Pamount[pre_num] - 1:
                        plt.xlabel("频率", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                    plt.subplot(Pamount[pre_num], 2, i * 2 + 2)
                    plt.plot(fre[pre_num][i], power[pre_num][i])
                    plt.ylabel("频率", fontsize=14)
                    if i == Pamount[pre_num] - 1:
                        plt.xlabel("功率", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                plt.show()
            else:
                plt.figure(num=numname[pre_num])
                for i in range(Pamount[0]):
                    plt.subplot(Pamount[0], 2, i * 2 + 1)
                    plt.plot(numname[0:5], amp[i])
                    plt.ylabel("振幅", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                    plt.subplot(Pamount[0], 2, i * 2 + 2)
                    plt.plot(numname[0:5], pow[i])
                    plt.ylabel("功率", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = numname
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
        label0 = tk.Label(frame, text='频域分析结果')
        label0.grid(row=0, column=1, columnspan=2)
        buttun1 = tk.Button(frame, text='频谱/功率谱切换', command=show_change43)
        buttun1.grid(row=0, column=3)
        label1 = tk.Label(frame, text='频谱')
        label1.grid(row=0, column=4)
        buttun2 = tk.Button(frame, text='显示详图', command=show_image43)
        buttun2.grid(row=0, column=5)
        labelc1 = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc1[i] = tk.Label(frame, text='通道{}'.format(i + 1))
            labelc1[i].grid(row=1 + i * 3, column=0, rowspan=3)
        labelc2 = list(range(Pamount[0] * 3))
        for i in range(Pamount[0]):
            labelc2[0 + 3 * i] = tk.Label(frame, text='振动显著频率/Hz')
            labelc2[0 + 3 * i].grid(row=1 + i * 3 + 0, column=1)
            labelc2[1 + 3 * i] = tk.Label(frame, text='显著频率对应振幅值/ms-2')
            labelc2[1 + 3 * i].grid(row=1 + i * 3 + 1, column=1)
            labelc2[2 + 3 * i] = tk.Label(frame, text='功率最值/w')
            labelc2[2 + 3 * i].grid(row=1 + i * 3 + 2, column=1)
        labelc3 = list(range(Pamount[0] * 3))
        for i in range(Pamount[0]):
            labelc3[0 + 3 * i] = tk.Label(frame, text=obs_fre[0][i])
            labelc3[0 + 3 * i].grid(row=1 + i * 3 + 0, column=2)
            labelc3[1 + 3 * i] = tk.Label(frame, text=max_am[0][i])
            labelc3[1 + 3 * i].grid(row=1 + i * 3 + 1, column=2)
            labelc3[2 + 3 * i] = tk.Label(frame, text=round(max_power[0][i], 3))
            labelc3[2 + 3 * i].grid(row=1 + i * 3 + 2, column=2)
        labelc4 = list(range(Pamount[0]))
        img = list(range(Pamount[0]))
        photo = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc4[i] = tk.Label(frame)
            img[i] = Image.open(
                "cache\\" + dirname + "\\" + numname[0] + "\\2-1频谱-通道{}".format(i + 1) + ".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            labelc4[i].config(image=photo[i])
            labelc4[i].image = photo[i]
            labelc4[i].grid(row=1 + i * 3, column=3, rowspan=3, columnspan=3)
def vlzmax():
    global vl_maxz,vl_maxzw,vl_maxzwk
    if mode==1:
        vl_maxz = list(range(Pamount))
        vl_maxzw = list(range(Pamount))
        vl_maxzwk = list(range(Pamount))
        vl_time = list(range(Pamount))
        vl_z = list(np.arange(Pamount))
        vl_zw = list(np.arange(Pamount))
        vl_zwk = list(np.arange(Pamount))
        for i in range(Pamount):
            pre_resu=fft_zmax_w(data[i])
            vl_time[i] = pre_resu[0]
            vl_z[i] = pre_resu[3]
            vl_zw[i] = pre_resu[1]
            vl_zwk[i] = pre_resu[2]
            vl_maxz[i] = int(max(vl_z[i]))
            vl_maxzw[i] = int(max(vl_zw[i]))
            vl_maxzwk[i] = int(max(vl_zwk[i]))
            plt.figure(figsize=(8,0.9))
            plt.plot(vl_time[i], vl_z[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
            plt.savefig("cache\\" + dirname + "\\3-1Z振级-不计权-通道{}".format(i + 1) + ".png")
            plt.close()
            plt.figure(figsize=(8,0.9))
            plt.plot(vl_time[i], vl_zw[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
            plt.savefig("cache\\" + dirname + "\\3-2Z振级-85版-通道{}".format(i + 1) + ".png")
            plt.close()
            plt.figure(figsize=(8,0.9))
            plt.plot(vl_time[i], vl_zwk[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
            plt.savefig("cache\\" + dirname + "\\3-3Z振级-97版-通道{}".format(i + 1) + ".png")
            plt.close()
        for i in range(Pamount):
            plt.figure()
            plt.plot(vl_time[i], vl_z[i])
            plt.xlabel("时间", fontsize=14)
            plt.ylabel("Z振级/dB", fontsize=14)
            plt.xticks(fontsize=14)
            plt.yticks(fontsize=14)
            plt.savefig("result\\" + dirname + "\\3-1Z振级-不计权-通道{}".format(i + 1) + ".png")
            plt.close()
            plt.figure()
            plt.plot(vl_time[i], vl_zw[i])
            plt.xlabel("时间", fontsize=14)
            plt.ylabel("Z振级/dB", fontsize=14)
            plt.xticks(fontsize=14)
            plt.yticks(fontsize=14)
            plt.savefig("result\\" + dirname + "\\3-2Z振级-85版-通道{}".format(i + 1) + ".png")
            plt.close()
            plt.figure()
            plt.plot(vl_time[i], vl_zwk[i])
            plt.xlabel("时间", fontsize=14)
            plt.ylabel("Z振级/dB", fontsize=14)
            plt.xticks(fontsize=14)
            plt.yticks(fontsize=14)
            plt.savefig("result\\" + dirname + "\\3-3Z振级-97版-通道{}".format(i + 1) + ".png")
            plt.close()
        frame=tk.Frame(window,width=1920,height=1080)
        frame.grid(row=0,column=0)
        frame.grid_propagate(0)
        def show(event):
            pre_num = var.get()
        def show_change51():
            if label1["text"]=='最大Z振级（不计权）/dB':
                label1.config(text='最大Z振级（85版）/dB')
                for i in range(Pamount):
                    img[i] = Image.open("cache\\" + dirname + "\\3-2Z振级-85版-通道{}".format(i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
            elif label1["text"]=='最大Z振级（85版）/dB':
                label1.config(text='最大Z振级（97版）/dB')
                for i in range(Pamount):
                    img[i] = Image.open("cache\\" + dirname + "\\3-3Z振级-97版-通道{}".format(i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
            else:
                label1.config(text='最大Z振级（不计权）/dB')
                for i in range(Pamount):
                    img[i] = Image.open("cache\\" + dirname + "\\3-1Z振级-不计权-通道{}".format(i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
        def show_image51():
            if label1["text"]=='最大Z振级（不计权）/dB':
                plt.figure()
                for i in range(Pamount):
                    plt.subplot(Pamount, 1, i  + 1)
                    plt.plot(vl_time[i], vl_z[i])
                    plt.ylabel("Z振级/dB", fontsize=14)
                    if i == Pamount - 1:
                        plt.xlabel("时间/s", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                plt.show()
            elif label1["text"]=='最大Z振级（85版）/dB':
                plt.figure()
                for i in range(Pamount):
                    plt.subplot(Pamount, 1, i  + 1)
                    plt.plot(vl_time[i], vl_zw[i])
                    plt.ylabel("Z振级/dB", fontsize=14)
                    if i == Pamount - 1:
                        plt.xlabel("时间/s", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                plt.show()
            else:
                plt.figure()
                for i in range(Pamount):
                    plt.subplot(Pamount, 1, i  + 1)
                    plt.plot(vl_time[i], vl_zwk[i])
                    plt.ylabel("Z振级/dB", fontsize=14)
                    if i == Pamount - 1:
                        plt.xlabel("时间/s", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = ['Group1']
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
        label0 = tk.Label(frame, text='最大Z振级')
        label0.grid(row=0, column=1, columnspan=2)
        buttun1 = tk.Button(frame, text='不计权曲线/85版计权曲线/97版计权曲线切换', command=show_change51)
        buttun1.grid(row=0, column=3)
        label1 = tk.Label(frame, text='最大Z振级（不计权）/dB')
        label1.grid(row=0, column=4)
        buttun2 = tk.Button(frame, text='显示详图', command=show_image51)
        buttun2.grid(row=0, column=5)
        labelc1 = list(range(Pamount))
        for i in range(Pamount):
            labelc1[i] = tk.Label(frame, text='通道{}'.format(i + 1))
            labelc1[i].grid(row=1 + i * 3, column=0, rowspan=3)
        labelc2 = list(range(Pamount * 3))
        for i in range(Pamount):
            labelc2[0 + 3 * i] = tk.Label(frame, text='最大Z振级（不计权）/dB')
            labelc2[0 + 3 * i].grid(row=1 + i * 3 + 0, column=1)
            labelc2[1 + 3 * i] = tk.Label(frame, text='最大Z振级（85版）/dB')
            labelc2[1 + 3 * i].grid(row=1 + i * 3 + 1, column=1)
            labelc2[2 + 3 * i] = tk.Label(frame, text='最大Z振级（97版）/dB')
            labelc2[2 + 3 * i].grid(row=1 + i * 3 + 2, column=1)
        labelc3 = list(range(Pamount * 3))
        for i in range(Pamount):
            labelc3[0 + 3 * i] = tk.Label(frame, text=vl_maxz[i])
            labelc3[0 + 3 * i].grid(row=1 + i * 3 + 0, column=2)
            labelc3[1 + 3 * i] = tk.Label(frame, text=vl_maxzw[i])
            labelc3[1 + 3 * i].grid(row=1 + i * 3 + 1, column=2)
            labelc3[2 + 3 * i] = tk.Label(frame, text=vl_maxzwk[i])
            labelc3[2 + 3 * i].grid(row=1 + i * 3 + 2, column=2)
        labelc4 = list(range(Pamount))
        img = list(range(Pamount))
        photo = list(range(Pamount))
        for i in range(Pamount):
            labelc4[i] = tk.Label(frame)
            img[i] = Image.open("cache\\" + dirname + "\\3-1Z振级-不计权-通道{}".format(i + 1) + ".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            labelc4[i].config(image=photo[i])
            labelc4[i].image = photo[i]
            labelc4[i].grid(row=1 + i * 3, column=3, rowspan=3, columnspan=3)
    if mode==2:
        vl_maxz = list(range(len(numname)))
        vl_maxzw = list(range(len(numname)))
        vl_maxzwk = list(range(len(numname)))
        vl_time = list(range(len(numname)))
        vl_z = list(range(len(numname)))
        vl_zw = list(range(len(numname)))
        vl_zwk = list(range(len(numname)))
        for i in range(len(numname)):
            vl_maxz[i] = list(range(Pamount[i]))
            vl_maxzw[i] = list(range(Pamount[i]))
            vl_maxzwk[i] = list(range(Pamount[i]))
            vl_time[i] = list(range(Pamount[i]))
            vl_z[i] = list(range(Pamount[i]))
            vl_zw[i] = list(range(Pamount[i]))
            vl_zwk[i] = list(range(Pamount[i]))
        for i in range(len(numname)):
            for j in range(Pamount[i]):
                pre_resu = fft_zmax_w(data[i][j])
                vl_time[i][j] = pre_resu[0]
                vl_z[i][j] = pre_resu[3]
                vl_zw[i][j] = pre_resu[1]
                vl_zwk[i][j] = pre_resu[2]
                vl_maxz[i][j] = int(max(vl_z[i][j]))
                vl_maxzw[i][j] = int(max(vl_zw[i][j]))
                vl_maxzwk[i][j] = int(max(vl_zwk[i][j]))
                plt.figure(figsize=(8, 0.9))
                plt.plot(vl_time[i][j], vl_z[i][j])
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname + "\\"+numname[i] + "\\3-1Z振级-不计权-通道{}".format(j + 1) + ".png")
                plt.close()
                plt.figure(figsize=(8, 0.9))
                plt.plot(vl_time[i][j], vl_zw[i][j])
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname +   "\\"+numname[i]+"\\3-2Z振级-85版-通道{}".format(j + 1) + ".png")
                plt.close()
                plt.figure(figsize=(8, 0.9))
                plt.plot(vl_time[i][j], vl_zwk[i][j])
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname +  "\\"+numname[i]+ "\\3-3Z振级-97版-通道{}".format(j + 1) + ".png")
                plt.close()
        for i in range(len(numname)):
            for j in range(Pamount[i]):
                plt.figure()
                plt.plot(vl_time[i][j], vl_z[i][j])
                plt.xlabel("时间", fontsize=14)
                plt.ylabel("Z振级/dB", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname +"\\"+numname[i]+ "\\3-1Z振级-不计权-通道{}".format(j + 1) + ".png")
                plt.close()
                plt.figure()
                plt.plot(vl_time[i][j], vl_zw[i][j])
                plt.xlabel("时间", fontsize=14)
                plt.ylabel("Z振级/dB", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname +"\\"+numname[i]+"\\3-2Z振级-85版-通道{}".format(j + 1) + ".png")
                plt.close()
                plt.figure()
                plt.plot(vl_time[i][j], vl_zwk[i][j])
                plt.xlabel("时间", fontsize=14)
                plt.ylabel("Z振级/dB", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname +"\\"+numname[i]+ "\\3-3Z振级-97版-通道{}".format(j + 1) + ".png")
                plt.close()
        frame = tk.Frame(window, width=1920, height=1080)
        frame.grid(row=0, column=0)
        frame.grid_propagate(0)
        pre_num = 0
        def show(event):
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            label1.config(text='最大Z振级（不计权）/dB')
            for i in range(Pamount[pre_num]):
                labelc3[0 + 3 * i] .config(text=vl_maxz[pre_num][i])
                labelc3[1 + 3 * i] .config(text=vl_maxzw[pre_num][i])
                labelc3[2 + 3 * i] .config(text=vl_maxzwk[pre_num][i])
            for i in range(Pamount[0]):
                img[i] = Image.open("cache\\" + dirname+ "\\"+numname[pre_num] + "\\3-1Z振级-不计权-通道{}".format(i + 1) + ".png")  # 打开图片
                photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                labelc4[i].config(image=photo[i])
                labelc4[i].image = photo[i]
        def show_change52():
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            if label1["text"]=='最大Z振级（不计权）/dB':
                label1.config(text='最大Z振级（85版）/dB')
                for i in range(Pamount[pre_num]):
                    img[i] = Image.open("cache\\" + dirname+ "\\"+numname[pre_num] + "\\3-2Z振级-85版-通道{}".format(i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
            elif label1["text"]=='最大Z振级（85版）/dB':
                label1.config(text='最大Z振级（97版）/dB')
                for i in range(Pamount[pre_num]):
                    img[i] = Image.open("cache\\" + dirname + "\\"+numname[pre_num]+ "\\3-3Z振级-97版-通道{}".format(i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
            else:
                label1.config(text='最大Z振级（不计权）/dB')
                for i in range(Pamount[pre_num]):
                    img[i] = Image.open("cache\\" + dirname + "\\"+numname[pre_num]+ "\\3-1Z振级-不计权-通道{}".format(i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
        def show_image52():
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            if label1["text"]=='最大Z振级（不计权）/dB':
                plt.figure()
                for i in range(Pamount[pre_num]):
                    plt.subplot(Pamount[pre_num], 1, i  + 1)
                    plt.plot(vl_time[pre_num][i], vl_z[pre_num][i])
                    plt.ylabel("Z振级/dB", fontsize=14)
                    if i == Pamount[pre_num] - 1:
                        plt.xlabel("时间/s", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                plt.show()
            elif label1["text"]=='最大Z振级（85版）/dB':
                plt.figure()
                for i in range(Pamount[pre_num]):
                    plt.subplot(Pamount[pre_num], 1, i  + 1)
                    plt.plot(vl_time[pre_num][i], vl_zw[pre_num][i])
                    plt.ylabel("Z振级/dB", fontsize=14)
                    if i == Pamount[pre_num] - 1:
                        plt.xlabel("时间/s", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                plt.show()
            else:
                plt.figure()
                for i in range(Pamount[pre_num]):
                    plt.subplot(Pamount[pre_num], 1, i  + 1)
                    plt.plot(vl_time[pre_num][i], vl_zwk[pre_num][i])
                    plt.ylabel("Z振级/dB", fontsize=14)
                    if i == Pamount[pre_num] - 1:
                        plt.xlabel("时间/s", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = numname
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
        label0 = tk.Label(frame, text='最大Z振级')
        label0.grid(row=0, column=1, columnspan=2)
        buttun1 = tk.Button(frame, text='不计权曲线/85版计权曲线/97版计权曲线切换', command=show_change52)
        buttun1.grid(row=0, column=3)
        label1 = tk.Label(frame, text='最大Z振级（不计权）/dB')
        label1.grid(row=0, column=4)
        buttun2 = tk.Button(frame, text='显示详图', command=show_image52)
        buttun2.grid(row=0, column=5)
        labelc1 = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc1[i] = tk.Label(frame, text='通道{}'.format(i + 1))
            labelc1[i].grid(row=1 + i * 3, column=0, rowspan=3)
        labelc2 = list(range(Pamount[0] * 3))
        for i in range(Pamount[0]):
            labelc2[0 + 3 * i] = tk.Label(frame, text='最大Z振级（不计权）/dB')
            labelc2[0 + 3 * i].grid(row=1 + i * 3 + 0, column=1)
            labelc2[1 + 3 * i] = tk.Label(frame, text='最大Z振级（85版）/dB')
            labelc2[1 + 3 * i].grid(row=1 + i * 3 + 1, column=1)
            labelc2[2 + 3 * i] = tk.Label(frame, text='最大Z振级（97版）/dB')
            labelc2[2 + 3 * i].grid(row=1 + i * 3 + 2, column=1)
        labelc3 = list(range(Pamount[0] * 3))
        for i in range(Pamount[0]):
            labelc3[0 + 3 * i] = tk.Label(frame, text=vl_maxz[0][i])
            labelc3[0 + 3 * i].grid(row=1 + i * 3 + 0, column=2)
            labelc3[1 + 3 * i] = tk.Label(frame, text=vl_maxzw[0][i])
            labelc3[1 + 3 * i].grid(row=1 + i * 3 + 1, column=2)
            labelc3[2 + 3 * i] = tk.Label(frame, text=vl_maxzwk[0][i])
            labelc3[2 + 3 * i].grid(row=1 + i * 3 + 2, column=2)
        labelc4 = list(range(Pamount[0]))
        img = list(range(Pamount[0]))
        photo = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc4[i] = tk.Label(frame)
            img[i] = Image.open("cache\\" + dirname+ "\\"+numname[0] + "\\3-1Z振级-不计权-通道{}".format(i + 1) + ".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            labelc4[i].config(image=photo[i])
            labelc4[i].image = photo[i]
            labelc4[i].grid(row=1 + i * 3, column=3, rowspan=3, columnspan=3)
    if mode==3:
        global vl_maxz_tj,vl_maxzw_tj,vl_maxzwk_tj
        vl_maxz = list(range(len(numname)-1))
        vl_maxzw = list(range(len(numname)-1))
        vl_maxzwk = list(range(len(numname)-1))
        vl_time = list(range(len(numname)-1))
        vl_z = list(range(len(numname)-1))
        vl_zw = list(range(len(numname)-1))
        vl_zwk = list(range(len(numname)-1))
        vl_maxz_tj = list(range(Pamount[0]))
        vl_maxzw_tj = list(range(Pamount[0]))
        vl_maxzwk_tj = list(range(Pamount[0]))
        zp=list(range(Pamount[0]))
        zwp=list(range(Pamount[0]))
        zwkp=list(range(Pamount[0]))
        for i in range(len(numname)-1):
            vl_maxz[i] = list(range(Pamount[i]))
            vl_maxzw[i] = list(range(Pamount[i]))
            vl_maxzwk[i] = list(range(Pamount[i]))
            vl_time[i] = list(range(Pamount[i]))
            vl_z[i] = list(range(Pamount[i]))
            vl_zw[i] = list(range(Pamount[i]))
            vl_zwk[i] = list(range(Pamount[i]))
        for i in range(Pamount[0]):
            vl_maxz_tj[i] = list(range(3))
            vl_maxzw_tj[i] = list(range(3))
            vl_maxzwk_tj[i] = list(range(3))
            zp[i] = []
            zwp[i] = []
            zwkp[i] = []
        for i in range(len(numname)-1):
            for j in range(Pamount[i]):
                pre_resu = fft_zmax_w(data[i][j])
                vl_time[i][j] = pre_resu[0]
                vl_z[i][j] = pre_resu[3]
                vl_zw[i][j] = pre_resu[1]
                vl_zwk[i][j] = pre_resu[2]
                vl_maxz[i][j] = int(max(vl_z[i][j]))
                vl_maxzw[i][j] = int(max(vl_zw[i][j]))
                vl_maxzwk[i][j] = int(max(vl_zwk[i][j]))
                plt.figure(figsize=(8, 0.9))
                plt.plot(vl_time[i][j], vl_z[i][j])
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname + "\\" + numname[i] + "\\3-1Z振级-不计权-通道{}".format(j + 1) + ".png")
                plt.close()
                plt.figure(figsize=(8, 0.9))
                plt.plot(vl_time[i][j], vl_zw[i][j])
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname + "\\" + numname[i] + "\\3-2Z振级-85版-通道{}".format(j + 1) + ".png")
                plt.close()
                plt.figure(figsize=(8, 0.9))
                plt.plot(vl_time[i][j], vl_zwk[i][j])
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname + "\\" + numname[i] + "\\3-3Z振级-97版-通道{}".format(j + 1) + ".png")
                plt.close()
        for i in range(len(numname)-1):
            for j in range(Pamount[i]):
                plt.figure()
                plt.plot(vl_time[i][j], vl_z[i][j])
                plt.xlabel("时间", fontsize=14)
                plt.ylabel("Z振级/dB", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname +"\\"+numname[i]+ "\\3-1Z振级-不计权-通道{}".format(j + 1) + ".png")
                plt.close()
                plt.figure()
                plt.plot(vl_time[i][j], vl_zw[i][j])
                plt.xlabel("时间", fontsize=14)
                plt.ylabel("Z振级/dB", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname +"\\"+numname[i]+"\\3-2Z振级-85版-通道{}".format(j + 1) + ".png")
                plt.close()
                plt.figure()
                plt.plot(vl_time[i][j], vl_zwk[i][j])
                plt.xlabel("时间", fontsize=14)
                plt.ylabel("Z振级/dB", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname +"\\"+numname[i]+ "\\3-3Z振级-97版-通道{}".format(j + 1) + ".png")
                plt.close()
        for i in range(Pamount[0]):
            for j in range(len(numname) - 1):
                zp[i].append(vl_maxz[j][i])
                zwp[i].append(vl_maxzw[j][i])
                zwkp[i].append(vl_maxzwk[j][i])
            vl_maxz_tj[i][0]=min(vl_maxz[0][i],vl_maxz[1][i],vl_maxz[2][i],vl_maxz[3][i],vl_maxz[4][i])
            sum1=0
            for j in range(len(numname)-1):
                sum1+=vl_maxz[j][i]
            vl_maxz_tj[i][1] =round(sum1/(len(numname)-1),0)
            vl_maxz_tj[i][2] =max(vl_maxz[0][i],vl_maxz[1][i],vl_maxz[2][i],vl_maxz[3][i],vl_maxz[4][i])
            vl_maxzw_tj[i][0]=min(vl_maxzw[0][i],vl_maxzw[1][i],vl_maxzw[2][i],vl_maxzw[3][i],vl_maxzw[4][i])
            sum2=0
            for j in range(len(numname)-1):
                sum2+=vl_maxzw[j][i]
            vl_maxzw_tj[i][1] =round(sum2/(len(numname)-1),0)
            vl_maxzw_tj[i][2] =max(vl_maxzw[0][i],vl_maxzw[1][i],vl_maxzw[2][i],vl_maxzw[3][i],vl_maxzw[4][i])
            vl_maxzwk_tj[i][0]=min(vl_maxzwk[0][i],vl_maxzwk[1][i],vl_maxzwk[2][i],vl_maxzwk[3][i],vl_maxzwk[4][i])
            sum3=0
            for j in range(len(numname)-1):
                sum3+=vl_maxzwk[j][i]
            vl_maxzwk_tj[i][1] =round(sum3/(len(numname)-1),0)
            vl_maxzwk_tj[i][2] =max(vl_maxzwk[0][i],vl_maxzwk[1][i],vl_maxzwk[2][i],vl_maxzwk[3][i],vl_maxzwk[4][i])
            plt.figure(figsize=(8, 0.9))
            plt.plot(numname[0:5], zp[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
            plt.savefig("cache\\" + dirname +  "\\3-1Z振级-不计权对比分析-通道{}".format(i + 1) + ".png")
            plt.close()
            plt.figure(figsize=(8, 0.9))
            plt.plot(numname[0:5], zwp[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
            plt.savefig("cache\\" + dirname +  "\\3-2Z振级-85版对比分析-通道{}".format(i + 1) + ".png")
            plt.close()
            plt.figure(figsize=(8, 0.9))
            plt.plot(numname[0:5], zwkp[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
            plt.savefig("cache\\" + dirname +  "\\3-3Z振级-97版对比分析-通道{}".format(i + 1) + ".png")
            plt.close()
            plt.figure(figsize=(8, 0.9))
            plt.plot(numname[0:5], zp[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
            plt.savefig("result\\" + dirname + "\\3-1Z振级-不计权对比分析-通道{}".format(i + 1) + ".png")
            plt.close()
            plt.figure(figsize=(8, 0.9))
            plt.plot(numname[0:5], zwp[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
            plt.savefig("result\\" + dirname + "\\3-2Z振级-85版对比分析-通道{}".format(i + 1) + ".png")
            plt.close()
            plt.figure(figsize=(8, 0.9))
            plt.plot(numname[0:5], zwkp[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.25, top=0.9)
            plt.savefig("result\\" + dirname + "\\3-3Z振级-97版对比分析-通道{}".format(i + 1) + ".png")
            plt.close()
        frame = tk.Frame(window, width=1920, height=1080)
        frame.grid(row=0, column=0)
        frame.grid_propagate(0)
        pre_num = 0
        def show(event):
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            if pre_num <= len(numname) - 2:
                label1.config(text='最大Z振级（不计权）/dB')
                for i in range(Pamount[pre_num]):
                    labelc2[0 + 3 * i].config(text='最大Z振级（不计权）/dB')
                    labelc2[1 + 3 * i].config(text='最大Z振级（85版）/dB')
                    labelc2[2 + 3 * i].config(text='最大Z振级（97版）/dB')
                    labelc3[0 + 3 * i].config(text=vl_maxz[pre_num][i])
                    labelc3[1 + 3 * i].config(text=vl_maxzw[pre_num][i])
                    labelc3[2 + 3 * i].config(text=vl_maxzwk[pre_num][i])
                for i in range(Pamount[0]):
                    img[i] = Image.open("cache\\" + dirname + "\\" + numname[pre_num] + "\\3-1Z振级-不计权-通道{}".format(
                        i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
            if pre_num == len(numname) - 1:
                label1.config(text='最大Z振级（不计权）对比分析/dB')
                for i in range(Pamount[pre_num-1]):
                    labelc2[0 + 3 * i].config(text='最大Z振级（不计权）（最小值、均值、最大值）/dB')
                    labelc2[1 + 3 * i].config(text='最大Z振级（85版）（最小值、均值、最大值）/dB')
                    labelc2[2 + 3 * i].config(text='最大Z振级（97版）（最小值、均值、最大值）/dB')
                    labelc3[0 + 3 * i].config(text=vl_maxz_tj[i])
                    labelc3[1 + 3 * i].config(text=vl_maxzw_tj[i])
                    labelc3[2 + 3 * i].config(text=vl_maxzwk_tj[i])
                for i in range(Pamount[0]):
                    img[i] = Image.open("cache\\" + dirname +  "\\3-1Z振级-不计权对比分析-通道{}".format(i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
        def show_change53():
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            if pre_num <= len(numname) - 2:
                if label1["text"] == '最大Z振级（不计权）/dB':
                    label1.config(text='最大Z振级（85版）/dB')
                    for i in range(Pamount[pre_num]):
                        img[i] = Image.open("cache\\" + dirname + "\\" + numname[pre_num] + "\\3-2Z振级-85版-通道{}".format(
                            i + 1) + ".png")  # 打开图片
                        photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                        labelc4[i].config(image=photo[i])
                        labelc4[i].image = photo[i]
                elif label1["text"] == '最大Z振级（85版）/dB':
                    label1.config(text='最大Z振级（97版）/dB')
                    for i in range(Pamount[pre_num]):
                        img[i] = Image.open("cache\\" + dirname + "\\" + numname[pre_num] + "\\3-3Z振级-97版-通道{}".format(
                            i + 1) + ".png")  # 打开图片
                        photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                        labelc4[i].config(image=photo[i])
                        labelc4[i].image = photo[i]
                else:
                    label1.config(text='最大Z振级（不计权）/dB')
                    for i in range(Pamount[pre_num]):
                        img[i] = Image.open("cache\\" + dirname + "\\" + numname[pre_num] + "\\3-1Z振级-不计权-通道{}".format(
                            i + 1) + ".png")  # 打开图片
                        photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                        labelc4[i].config(image=photo[i])
                        labelc4[i].image = photo[i]
            if pre_num == len(numname) - 1:
                if label1["text"] == '最大Z振级（不计权）对比分析/dB':
                    label1.config(text='最大Z振级（85版）对比分析/dB')
                    for i in range(Pamount[pre_num-1]):
                        img[i] = Image.open("cache\\" + dirname +  "\\3-1Z振级-不计权对比分析-通道{}".format(i + 1) + ".png")  # 打开图片
                        photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                        labelc4[i].config(image=photo[i])
                        labelc4[i].image = photo[i]
                elif label1["text"] == '最大Z振级（85版）对比分析/dB':
                    label1.config(text='最大Z振级（97版）对比分析/dB')
                    for i in range(Pamount[pre_num-1]):
                        img[i] = Image.open("cache\\" + dirname +  "\\3-2Z振级-85版对比分析-通道{}".format(i + 1) + ".png")  # 打开图片
                        photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                        labelc4[i].config(image=photo[i])
                        labelc4[i].image = photo[i]
                else:
                    label1.config(text='最大Z振级（不计权）对比分析/dB')
                    for i in range(Pamount[pre_num-1]):
                        img[i] = Image.open("cache\\" + dirname +  "\\3-3Z振级-97版对比分析-通道{}".format(i + 1) + ".png")  # 打开图片
                        photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                        labelc4[i].config(image=photo[i])
                        labelc4[i].image = photo[i]
        def show_image53():
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            if pre_num <= len(numname) - 2:
                if label1["text"] == '最大Z振级（不计权）/dB':
                    plt.figure()
                    for i in range(Pamount[pre_num]):
                        plt.subplot(Pamount[pre_num], 1, i + 1)
                        plt.plot(vl_time[pre_num][i], vl_z[pre_num][i])
                        plt.ylabel("Z振级/dB", fontsize=14)
                        if i == Pamount[pre_num] - 1:
                            plt.xlabel("时间/s", fontsize=14)
                        plt.xticks(fontsize=14)
                        plt.yticks(fontsize=14)
                    plt.show()
                elif label1["text"] == '最大Z振级（85版）/dB':
                    plt.figure()
                    for i in range(Pamount[pre_num]):
                        plt.subplot(Pamount[pre_num], 1, i + 1)
                        plt.plot(vl_time[pre_num][i], vl_zw[pre_num][i])
                        plt.ylabel("Z振级/dB", fontsize=14)
                        if i == Pamount[pre_num] - 1:
                            plt.xlabel("时间/s", fontsize=14)
                        plt.xticks(fontsize=14)
                        plt.yticks(fontsize=14)
                    plt.show()
                else:
                    plt.figure()
                    for i in range(Pamount[pre_num]):
                        plt.subplot(Pamount[pre_num], 1, i + 1)
                        plt.plot(vl_time[pre_num][i], vl_zwk[pre_num][i])
                        plt.ylabel("Z振级/dB", fontsize=14)
                        if i == Pamount[pre_num] - 1:
                            plt.xlabel("时间/s", fontsize=14)
                        plt.xticks(fontsize=14)
                        plt.yticks(fontsize=14)
                    plt.show()
            if pre_num == len(numname) - 1:
                if label1["text"] == '最大Z振级（不计权）对比分析/dB':
                    plt.figure()
                    for i in range(Pamount[pre_num-1]):
                        plt.subplot(Pamount[pre_num-1], 1, i + 1)
                        plt.plot(numname[0:5], zp[i])
                        plt.ylabel("Z振级/dB", fontsize=14)
                        plt.xticks(fontsize=14)
                        plt.yticks(fontsize=14)
                    plt.show()
                elif label1["text"] == '最大Z振级（85版）对比分析/dB':
                    plt.figure()
                    for i in range(Pamount[pre_num-1]):
                        plt.subplot(Pamount[pre_num-1], 1, i + 1)
                        plt.plot(numname[0:5], zwp[i])
                        plt.ylabel("Z振级/dB", fontsize=14)
                        plt.xticks(fontsize=14)
                        plt.yticks(fontsize=14)
                    plt.show()
                else:
                    plt.figure()
                    for i in range(Pamount[pre_num-1]):
                        plt.subplot(Pamount[pre_num-1], 1, i + 1)
                        plt.plot(numname[0:5], zwkp[i])
                        plt.ylabel("Z振级/dB", fontsize=14)
                        plt.xticks(fontsize=14)
                        plt.yticks(fontsize=14)
                    plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = numname
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
        label0 = tk.Label(frame, text='最大Z振级分析结果')
        label0.grid(row=0, column=1, columnspan=2)
        buttun1 = tk.Button(frame, text='不计权曲线/85版计权曲线/97版计权曲线切换', command=show_change53)
        buttun1.grid(row=0, column=3)
        label1 = tk.Label(frame, text='最大Z振级（不计权）/dB')
        label1.grid(row=0, column=4)
        buttun2 = tk.Button(frame, text='显示详图', command=show_image53)
        buttun2.grid(row=0, column=5)
        labelc1 = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc1[i] = tk.Label(frame, text='通道{}'.format(i + 1))
            labelc1[i].grid(row=1 + i * 3, column=0, rowspan=3)
        labelc2 = list(range(Pamount[0] * 3))
        for i in range(Pamount[0]):
            labelc2[0 + 3 * i] = tk.Label(frame, text='最大Z振级（不计权）/dB')
            labelc2[0 + 3 * i].grid(row=1 + i * 3 + 0, column=1)
            labelc2[1 + 3 * i] = tk.Label(frame, text='最大Z振级（85版）/dB')
            labelc2[1 + 3 * i].grid(row=1 + i * 3 + 1, column=1)
            labelc2[2 + 3 * i] = tk.Label(frame, text='最大Z振级（97版）/dB')
            labelc2[2 + 3 * i].grid(row=1 + i * 3 + 2, column=1)
        labelc3 = list(range(Pamount[0] * 3))
        for i in range(Pamount[0]):
            labelc3[0 + 3 * i] = tk.Label(frame, text=vl_maxz[0][i])
            labelc3[0 + 3 * i].grid(row=1 + i * 3 + 0, column=2)
            labelc3[1 + 3 * i] = tk.Label(frame, text=vl_maxzw[0][i])
            labelc3[1 + 3 * i].grid(row=1 + i * 3 + 1, column=2)
            labelc3[2 + 3 * i] = tk.Label(frame, text=vl_maxzwk[0][i])
            labelc3[2 + 3 * i].grid(row=1 + i * 3 + 2, column=2)
        labelc4 = list(range(Pamount[0]))
        img = list(range(Pamount[0]))
        photo = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc4[i] = tk.Label(frame)
            img[i] = Image.open("cache\\" + dirname+ "\\"+numname[0] + "\\3-1Z振级-不计权-通道{}".format(i + 1) + ".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            labelc4[i].config(image=photo[i])
            labelc4[i].image = photo[i]
            labelc4[i].grid(row=1 + i * 3, column=3, rowspan=3, columnspan=3)
def vlmax0():
    global vl_max_0,vl_max_0_am,vl_max_0_fre
    if mode==1:
        mid_fre = [1, 1.25, 1.6, 2, 2.5, 3.15, 4, 5, 6.3, 8, 10, 12.5, 16, 20, 25, 31.5, 40, 50, 63, 80, 100, 125, 160,
                   200, 250, 315, 400, 500, 630, 800, 1000, 1200]
        mid_frep = ['1', '1.25', '1.6', '2', '2.5', '3.15', '4', '5', '6.3', '8', '10', '12.5',
                    '16', '20', '25', '31.5', '40', '50', '63', '80', '100', '125', '160',
                    '200', '250', '315', '400', '500', '630', '800', '1000', '1200']
        vl_max_0 = list(range(Pamount))
        vl_max_0_am=list(range(Pamount))
        vl_max_0_fre=list(range(Pamount))
        for i in range(Pamount):
            vl_max_0[i] = fft_max_0(data[i])
            vl_max_0_am[i]=max(vl_max_0[i])
            for j in range(len(mid_fre)):
                if vl_max_0_am[i]==vl_max_0[i][j]:
                    vl_max_0_fre[i]=mid_fre[j]
        for i in range(Pamount):
            plt.figure(figsize=(8,0.9))
            plt.plot(mid_frep, vl_max_0[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.05, right=0.95, bottom=0.25, top=0.9)
            plt.savefig("cache\\" + dirname + "\\4-1未分幅未计权分频振级-通道{}".format(i + 1) + ".png")
            plt.close()
        for i in range(Pamount):
            plt.figure()
            plt.plot(mid_fre, vl_max_0[i])
            plt.xlabel("中心频率/Hz", fontsize=14)
            plt.ylabel("分频振级/dB", fontsize=14)
            plt.xticks(fontsize=14)
            plt.yticks(fontsize=14)
            plt.savefig("result\\" + dirname + "\\4-1未分幅未计权分频振级-通道{}".format(i + 1) + ".png")
            plt.close()
        frame = tk.Frame(window, width=1920, height=1080)
        frame.grid(row=0, column=0)
        frame.grid_propagate(0)
        def show(event):
            pre_num = var.get()
        def show_image51():
            plt.figure()
            for i in range(Pamount):
                plt.subplot(Pamount, 1, i  + 1)
                plt.plot(mid_frep, vl_max_0[i])
                plt.ylabel("分频振级/dB", fontsize=14)
                if i == Pamount - 1:
                    plt.xlabel("中心频率", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
            plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = ['Group1']
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
        label0 = tk.Label(frame, text='分频振级（未加窗未计权）分析结果')
        label0.grid(row=0, column=1, columnspan=2)
        label1 = tk.Label(frame, text='分频振级图')
        label1.grid(row=0, column=4)
        buttun2 = tk.Button(frame, text='显示详图', command=show_image51)
        buttun2.grid(row=0, column=5)
        labelc1 = list(range(Pamount))
        for i in range(Pamount):
            labelc1[i] = tk.Label(frame, text='通道{}'.format(i + 1))
            labelc1[i].grid(row=1 + i * 2, column=0, rowspan=2)
        labelc2 = list(range(Pamount * 2))
        for i in range(Pamount):
            labelc2[0 + 2 * i] = tk.Label(frame, text='分频最大振级/dB')
            labelc2[0 + 2 * i].grid(row=1 + i * 2 + 0, column=1)
            labelc2[1 + 2 * i] = tk.Label(frame, text='分频最大振级对应中心频率/Hz')
            labelc2[1 + 2 * i].grid(row=1 + i * 2 + 1, column=1)
        labelc3 = list(range(Pamount * 2))
        for i in range(Pamount):
            labelc3[0 + 2 * i] = tk.Label(frame, text=vl_max_0_am[i])
            labelc3[0 + 2 * i].grid(row=1 + i * 2 + 0, column=2)
            labelc3[1 + 2 * i] = tk.Label(frame, text=vl_max_0_fre[i])
            labelc3[1 + 2 * i].grid(row=1 + i * 2 + 1, column=2)
        labelc4 = list(range(Pamount))
        img = list(range(Pamount))
        photo = list(range(Pamount))
        for i in range(Pamount):
            labelc4[i] = tk.Label(frame)
            img[i] = Image.open("cache\\" + dirname + "\\4-1未分幅未计权分频振级-通道{}".format(i + 1) + ".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            labelc4[i].config(image=photo[i])
            labelc4[i].image = photo[i]
            labelc4[i].grid(row=1 + i * 2, column=3, rowspan=2, columnspan=3)
    if mode==2:
        mid_fre = [1, 1.25, 1.6, 2, 2.5, 3.15, 4, 5, 6.3, 8, 10, 12.5, 16, 20, 25, 31.5, 40, 50, 63, 80, 100, 125, 160,
                   200, 250, 315, 400, 500, 630, 800, 1000, 1200]
        mid_frep = ['1', '1.25', '1.6', '2', '2.5', '3.15', '4', '5', '6.3', '8', '10', '12.5',
                    '16', '20', '25', '31.5', '40', '50', '63', '80', '100', '125', '160',
                    '200', '250', '315', '400', '500', '630', '800', '1000', '1200']
        vl_max_0 = list(range(len(numname)))
        vl_max_0_am = list(range(len(numname)))
        vl_max_0_fre = list(range(len(numname)))
        for i in range(len(numname)):
            vl_max_0[i] = list(range(Pamount[i]))
            vl_max_0_am[i] = list(range(Pamount[i]))
            vl_max_0_fre[i] = list(range(Pamount[i]))
        for i in range(len(numname)):
            for j in range(Pamount[i]):
                vl_max_0[i][j] = fft_max_0(data[i][j])
                vl_max_0_am[i][j] = max(vl_max_0[i][j])
                for k in range(len(mid_fre)):
                    if vl_max_0_am[i][j] == vl_max_0[i][j][k]:
                        vl_max_0_fre[i][j] = mid_fre[k]
        for i in range(len(numname)):
            for j in range(Pamount[i]):
                plt.figure(figsize=(8, 0.9))
                plt.plot(mid_frep, vl_max_0[i][j])
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.05, right=0.95, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname + "\\" + numname[i] + "\\4-1未分幅未计权分频振级-通道{}".format(j + 1) + ".png")
                plt.close()
        for i in range(len(numname)):
            for j in range(Pamount[i]):
                plt.figure()
                plt.plot(mid_fre, vl_max_0[i][j])
                plt.xlabel("中心频率/Hz", fontsize=14)
                plt.ylabel("分频振级/dB", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname + "\\" + numname[i] +"\\4-1未分幅未计权分频振级-通道{}".format(j + 1) + ".png")
                plt.close()
        frame = tk.Frame(window, width=1920, height=1080)
        frame.grid(row=0, column=0)
        frame.grid_propagate(0)
        def show(event):
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            labelc3[0 + 2 * i].config( text=vl_max_0_am[pre_num][i])
            labelc3[1 + 2 * i].config(text=vl_max_0_fre[pre_num][i])
            for i in range(Pamount[0]):
                labelc4[i] = tk.Label(frame)
                img[i] = Image.open(
                    "cache\\" + dirname + "\\" + numname[pre_num] + "\\4-1未分幅未计权分频振级-通道{}".format(i + 1) + ".png")  # 打开图片
                photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                labelc4[i].config(image=photo[i])
                labelc4[i].image = photo[i]
        def show_image52():
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            plt.figure(num=numname[pre_num])
            for i in range(Pamount[pre_num]):
                plt.subplot(Pamount[pre_num], 1, i + 1)
                plt.plot(mid_frep, vl_max_0[pre_num][i])
                plt.ylabel("分频振级/dB", fontsize=14)
                if i == Pamount[pre_num] - 1:
                    plt.xlabel("中心频率", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
            plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = numname
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
        label0 = tk.Label(frame, text='分频振级（未加窗未计权）分析结果')
        label0.grid(row=0, column=1, columnspan=2)
        label1 = tk.Label(frame, text='分频振级图')
        label1.grid(row=0, column=4)
        buttun2 = tk.Button(frame, text='显示详图', command=show_image52)
        buttun2.grid(row=0, column=5)
        labelc1 = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc1[i] = tk.Label(frame, text='通道{}'.format(i + 1))
            labelc1[i].grid(row=1 + i * 2, column=0, rowspan=2)
        labelc2 = list(range(Pamount[0] * 2))
        for i in range(Pamount[0]):
            labelc2[0 + 2 * i] = tk.Label(frame, text='分频最大振级/dB')
            labelc2[0 + 2 * i].grid(row=1 + i * 2 + 0, column=1)
            labelc2[1 + 2 * i] = tk.Label(frame, text='分频最大振级对应中心频率/Hz')
            labelc2[1 + 2 * i].grid(row=1 + i * 2 + 1, column=1)
        labelc3 = list(range(Pamount[0] * 2))
        for i in range(Pamount[0]):
            labelc3[0 + 2 * i] = tk.Label(frame, text=vl_max_0_am[0][i])
            labelc3[0 + 2 * i].grid(row=1 + i * 2 + 0, column=2)
            labelc3[1 + 2 * i] = tk.Label(frame, text=vl_max_0_fre[0][i])
            labelc3[1 + 2 * i].grid(row=1 + i * 2 + 1, column=2)
        labelc4 = list(range(Pamount[0]))
        img = list(range(Pamount[0]))
        photo = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc4[i] = tk.Label(frame)
            img[i] = Image.open("cache\\" + dirname + "\\" + numname[0] + "\\4-1未分幅未计权分频振级-通道{}".format(i + 1) + ".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            labelc4[i].config(image=photo[i])
            labelc4[i].image = photo[i]
            labelc4[i].grid(row=1 + i * 2, column=3, rowspan=2, columnspan=3)
    if mode==3:
        global  vl_max_0_am_tj, vl_max_0_fre_tj
        mid_fre = [1, 1.25, 1.6, 2, 2.5, 3.15, 4, 5, 6.3, 8, 10, 12.5, 16, 20, 25, 31.5, 40, 50, 63, 80, 100, 125, 160,
                   200, 250, 315, 400, 500, 630, 800, 1000, 1200]
        mid_frep = ['1', '1.25', '1.6', '2', '2.5', '3.15', '4', '5', '6.3', '8', '10', '12.5',
                    '16', '20', '25', '31.5', '40', '50', '63', '80', '100', '125', '160',
                    '200', '250', '315', '400', '500', '630', '800', '1000', '1200']
        vl_max_0 = list(range(len(numname)-1))
        vl_max_0_am = list(range(len(numname)-1))
        vl_max_0_fre = list(range(len(numname)-1))
        vl_max_0_am_tj=list(range(Pamount[0]))
        vl_max_0_fre_tj=list(range(Pamount[0]))
        amp=list(range(Pamount[0]))
        frep=list(range(Pamount[0]))
        for i in range(len(numname)-1):
            vl_max_0[i] = list(range(Pamount[i]))
            vl_max_0_am[i] = list(range(Pamount[i]))
            vl_max_0_fre[i] = list(range(Pamount[i]))
        for i in range(len(numname)-1):
            for j in range(Pamount[i]):
                vl_max_0[i][j] = fft_max_0(data[i][j])
                vl_max_0_am[i][j] = max(vl_max_0[i][j])
                for k in range(len(mid_fre)):
                    if vl_max_0_am[i][j] == vl_max_0[i][j][k]:
                        vl_max_0_fre[i][j] = mid_fre[k]
        for i in range(len(numname)-1):
            for j in range(Pamount[i]):
                plt.figure(figsize=(8, 0.9))
                plt.plot(mid_frep, vl_max_0[i][j])
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.05, right=0.95, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname + "\\" + numname[i] + "\\4-1未分幅未计权分频振级-通道{}".format(j + 1) + ".png")
                plt.close()
        for i in range(len(numname)-1):
            for j in range(Pamount[i]):
                plt.figure()
                plt.plot(mid_fre, vl_max_0[i][j])
                plt.xlabel("中心频率/Hz", fontsize=14)
                plt.ylabel("分频振级/dB", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname + "\\" + numname[i] +"\\4-1未分幅未计权分频振级-通道{}".format(j + 1) + ".png")
                plt.close()
        for i in range(Pamount[0]):
            vl_max_0_am_tj[i]=list(range(3))
            vl_max_0_fre_tj[i]=list(range(3))
            amp[i]=[]
            frep[i]=[]
            for j in range(len(numname) - 1):
                amp[i].append(vl_max_0_am[j][i])
                frep[i].append(vl_max_0_fre[j][i])
            vl_max_0_am_tj[i][0]=min(vl_max_0_am[0][i],vl_max_0_am[1][i],vl_max_0_am[2][i],vl_max_0_am[3][i],vl_max_0_am[4][i])
            sum1 = 0
            for j in range(len(numname) - 1):
                sum1 += vl_max_0_am[j][i]
            vl_max_0_am_tj[i][1]=round(sum1/(len(numname)-1),0)
            vl_max_0_am_tj[i][2]=max(vl_max_0_am[0][i],vl_max_0_am[1][i],vl_max_0_am[2][i],vl_max_0_am[3][i],vl_max_0_am[4][i])
            vl_max_0_fre_tj[i][0]=min(vl_max_0_fre[0][i],vl_max_0_fre[1][i],vl_max_0_fre[2][i],vl_max_0_fre[3][i],vl_max_0_fre[4][i])
            sum2 = 0
            for j in range(len(numname) - 1):
                sum2 += vl_max_0_fre[j][i]
            vl_max_0_fre_tj[i][1]=round(sum2/(len(numname)-1),0)
            vl_max_0_fre_tj[i][2]=max(vl_max_0_fre[0][i],vl_max_0_fre[1][i],vl_max_0_fre[2][i],vl_max_0_fre[3][i],vl_max_0_fre[4][i])
            plt.figure(figsize=(8, 0.9))
            plt.subplot(1, 2, 1)
            plt.title('分频最大振级变化')
            plt.plot(numname[0:5], amp[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplot(1, 2, 2)
            plt.title('分频振级对应中心频率变化')
            plt.plot(numname[0:5], frep[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.27, top=0.75, wspace=0.25)
            plt.savefig("cache\\" + dirname + "\\4-1未分幅未计权分频振级对比分析-通道{}".format(i + 1) + ".png")
            plt.close()
            plt.figure()
            plt.subplot(1, 2, 1)
            plt.title('分频最大振级变化')
            plt.plot(numname[0:5], amp[i])
            plt.ylabel("分频振级/dB", fontsize=14)
            plt.xticks(fontsize=14)
            plt.yticks(fontsize=14)
            plt.subplot(1, 2, 2)
            plt.title('分频振级对应中心频率变化')
            plt.plot(numname[0:5], frep[i])
            plt.ylabel("对应中心频率/Hz", fontsize=14)
            plt.xticks(fontsize=14)
            plt.yticks(fontsize=14)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.27, top=0.75, wspace=0.25)
            plt.savefig("result\\" + dirname + "\\4-1未分幅未计权分频振级对比分析-通道{}".format(i + 1) + ".png")
            plt.close()
        frame = tk.Frame(window, width=1920, height=1080)
        frame.grid(row=0, column=0)
        frame.grid_propagate(0)
        def show(event):
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            if pre_num<=len(numname)-2:
                label1.config(text='分频振级图')
                labelc2[0 + 2 * i].config(text='分频最大振级/dB')
                labelc2[1 + 2 * i].config(text='分频最大振级对应中心频率/Hz')
                labelc3[0 + 2 * i].config(text=vl_max_0_am[pre_num][i])
                labelc3[1 + 2 * i].config(text=vl_max_0_fre[pre_num][i])
                for i in range(Pamount[0]):
                    img[i] = Image.open(
                        "cache\\" + dirname + "\\" + numname[pre_num] + "\\4-1未分幅未计权分频振级-通道{}".format(
                            i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
            else:
                label1.config(text='对比分析结果')
                for i in range(Pamount[0]):
                    labelc2[0 + 2 * i].config(text='分频最大振级(最小值、均值、最大值）/dB')
                    labelc2[1 + 2 * i].config(text='分频最大振级对应中心频率(最小值、均值、最大值）/Hz')
                    labelc3[0 + 2 * i].config(text=vl_max_0_am_tj[i])
                    labelc3[1 + 2 * i].config(text=vl_max_0_fre_tj[i])
                    img[i] = Image.open("cache\\" + dirname + "\\4-1未分幅未计权分频振级对比分析-通道{}".format(i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
        def show_image53():
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            if pre_num <= len(numname) - 2:
                plt.figure(num=numname[pre_num])
                for i in range(Pamount[pre_num]):
                    plt.subplot(Pamount[pre_num], 1, i + 1)
                    plt.plot(mid_frep, vl_max_0[pre_num][i])
                    plt.ylabel("分频振级/dB", fontsize=14)
                    if i == Pamount[pre_num] - 1:
                        plt.xlabel("中心频率/Hz", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                plt.show()
            else:
                plt.figure(num=numname[pre_num])
                for i in range(Pamount[0]):
                    plt.subplot(Pamount[0], 2, i * 2 + 1)
                    plt.plot(numname[0:5], amp[i])
                    plt.ylabel("分频振级/dB", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                    plt.subplot(Pamount[0], 2, i * 2 + 2)
                    plt.plot(numname[0:5], frep[i])
                    plt.ylabel("中心频率/Hz", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = numname
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
        label0 = tk.Label(frame, text='分频振级（未加窗未计权）分析结果')
        label0.grid(row=0, column=1, columnspan=2)
        label1 = tk.Label(frame, text='分频振级图')
        label1.grid(row=0, column=4)
        buttun2 = tk.Button(frame, text='显示详图', command=show_image53)
        buttun2.grid(row=0, column=5)
        labelc1 = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc1[i] = tk.Label(frame, text='通道{}'.format(i + 1))
            labelc1[i].grid(row=1 + i * 2, column=0, rowspan=2)
        labelc2 = list(range(Pamount[0] * 2))
        for i in range(Pamount[0]):
            labelc2[0 + 2 * i] = tk.Label(frame, text='分频最大振级/dB')
            labelc2[0 + 2 * i].grid(row=1 + i * 2 + 0, column=1)
            labelc2[1 + 2 * i] = tk.Label(frame, text='分频最大振级对应中心频率/Hz')
            labelc2[1 + 2 * i].grid(row=1 + i * 2 + 1, column=1)
        labelc3 = list(range(Pamount[0] * 2))
        for i in range(Pamount[0]):
            labelc3[0 + 2 * i] = tk.Label(frame, text=vl_max_0_am[0][i])
            labelc3[0 + 2 * i].grid(row=1 + i * 2 + 0, column=2)
            labelc3[1 + 2 * i] = tk.Label(frame, text=vl_max_0_fre[0][i])
            labelc3[1 + 2 * i].grid(row=1 + i * 2 + 1, column=2)
        labelc4 = list(range(Pamount[0]))
        img = list(range(Pamount[0]))
        photo = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc4[i] = tk.Label(frame)
            img[i] = Image.open(
                "cache\\" + dirname + "\\" + numname[0] + "\\4-1未分幅未计权分频振级-通道{}".format(i + 1) + ".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            labelc4[i].config(image=photo[i])
            labelc4[i].image = photo[i]
            labelc4[i].grid(row=1 + i * 2, column=3, rowspan=2, columnspan=3)
def vlmax1():
    global vl_max_1,vl_max_1_am,vl_max_1_fre
    if mode==1:
        mid_fre = [1, 1.25, 1.6, 2, 2.5, 3.15, 4, 5, 6.3, 8, 10, 12.5, 16, 20, 25, 31.5, 40, 50, 63, 80, 100, 125, 160,
                   200, 250, 315, 400, 500, 630, 800, 1000, 1200]
        mid_frep = ['1', '1.25', '1.6', '2', '2.5', '3.15', '4', '5', '6.3', '8', '10', '12.5',
                    '16', '20', '25', '31.5', '40', '50', '63', '80', '100', '125', '160',
                    '200', '250', '315', '400', '500', '630', '800', '1000', '1200']
        vl_max_1 = list(range(Pamount))
        vl_max_1_am= list(range(Pamount))
        vl_max_1_fre= list(range(Pamount))
        for i in range(Pamount):
            vl_max_1[i] = fft_max_1(data[i])
            vl_max_1_am[i]=max(vl_max_1[i])
            for j in range(len(mid_fre)):
                if vl_max_1_am[i]==vl_max_1[i][j]:
                    vl_max_1_fre[i]=mid_fre[j]
        for i in range(Pamount):
            plt.figure(figsize=(8,0.9))
            plt.plot(mid_frep, vl_max_1[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.05, right=0.95, bottom=0.25, top=0.9)
            plt.savefig("cache\\" + dirname + "\\4-2分幅未计权分频振级-通道{}".format(i + 1) + ".png")
            plt.close()
        for i in range(Pamount):
            plt.figure()
            plt.plot(mid_fre, vl_max_1[i])
            plt.xlabel("中心频率/Hz", fontsize=14)
            plt.ylabel("分频振级/dB", fontsize=14)
            plt.xticks(fontsize=14)
            plt.yticks(fontsize=14)
            plt.savefig("result\\" + dirname + "\\4-2分幅未计权分频振级-通道{}".format(i + 1) + ".png")
            plt.close()
        frame = tk.Frame(window, width=1920, height=1080)
        frame.grid(row=0, column=0)
        frame.grid_propagate(0)
        def show(event):
            pre_num = var.get()
        def show_image61():
            plt.figure()
            for i in range(Pamount):
                plt.subplot(Pamount, 1, i + 1)
                plt.plot(mid_frep, vl_max_1[i])
                plt.ylabel("分频振级/dB", fontsize=14)
                if i == Pamount - 1:
                    plt.xlabel("中心频率", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
            plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = ['Group1']
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
        label0 = tk.Label(frame, text='分频振级（加窗未计权）分析结果')
        label0.grid(row=0, column=1, columnspan=2)
        label1 = tk.Label(frame, text='分频振级图')
        label1.grid(row=0, column=4)
        buttun2 = tk.Button(frame, text='显示详图', command=show_image61)
        buttun2.grid(row=0, column=5)
        labelc1 = list(range(Pamount))
        for i in range(Pamount):
            labelc1[i] = tk.Label(frame, text='通道{}'.format(i + 1))
            labelc1[i].grid(row=1 + i * 2, column=0, rowspan=2)
        labelc2 = list(range(Pamount * 2))
        for i in range(Pamount):
            labelc2[0 + 2 * i] = tk.Label(frame, text='分频最大振级/dB')
            labelc2[0 + 2 * i].grid(row=1 + i * 2 + 0, column=1)
            labelc2[1 + 2 * i] = tk.Label(frame, text='分频最大振级对应中心频率/Hz')
            labelc2[1 + 2 * i].grid(row=1 + i * 2 + 1, column=1)
        labelc3 = list(range(Pamount * 2))
        for i in range(Pamount):
            labelc3[0 + 2 * i] = tk.Label(frame, text=vl_max_1_am[i])
            labelc3[0 + 2 * i].grid(row=1 + i * 2 + 0, column=2)
            labelc3[1 + 2 * i] = tk.Label(frame, text=vl_max_1_fre[i])
            labelc3[1 + 2 * i].grid(row=1 + i * 2 + 1, column=2)
        labelc4 = list(range(Pamount))
        img = list(range(Pamount))
        photo = list(range(Pamount))
        for i in range(Pamount):
            labelc4[i] = tk.Label(frame)
            img[i] = Image.open("cache\\" + dirname + "\\4-2分幅未计权分频振级-通道{}".format(i + 1) + ".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            labelc4[i].config(image=photo[i])
            labelc4[i].image = photo[i]
            labelc4[i].grid(row=1 + i * 2, column=3, rowspan=2, columnspan=3)
    if mode==2:
        mid_fre = [1, 1.25, 1.6, 2, 2.5, 3.15, 4, 5, 6.3, 8, 10, 12.5, 16, 20, 25, 31.5, 40, 50, 63, 80, 100, 125, 160,
                   200, 250, 315, 400, 500, 630, 800, 1000, 1200]
        mid_frep = ['1', '1.25', '1.6', '2', '2.5', '3.15', '4', '5', '6.3', '8', '10', '12.5',
                    '16', '20', '25', '31.5', '40', '50', '63', '80', '100', '125', '160',
                    '200', '250', '315', '400', '500', '630', '800', '1000', '1200']
        vl_max_1 = list(range(len(numname)))
        vl_max_1_am = list(range(len(numname)))
        vl_max_1_fre = list(range(len(numname)))
        for i in range(len(numname)):
            vl_max_1[i] = list(range(Pamount[i]))
            vl_max_1_am[i] = list(range(Pamount[i]))
            vl_max_1_fre[i] = list(range(Pamount[i]))
        for i in range(len(numname)):
            for j in range(Pamount[i]):
                vl_max_1[i][j] = fft_max_1(data[i][j])
                vl_max_1_am[i][j] = max(vl_max_1[i][j])
                for k in range(len(mid_fre)):
                    if vl_max_1_am[i][j] == vl_max_1[i][j][k]:
                        vl_max_1_fre[i][j] = mid_fre[k]
        for i in range(len(numname)):
            for j in range(Pamount[i]):
                plt.figure(figsize=(8, 0.9))
                plt.plot(mid_frep, vl_max_1[i][j])
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.05, right=0.95, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname + "\\" + numname[i] + "\\4-2分幅未计权分频振级-通道{}".format(j + 1) + ".png")
                plt.close()
        for i in range(len(numname)):
            for j in range(Pamount[i]):
                plt.figure()
                plt.plot(mid_fre, vl_max_1[i][j])
                plt.xlabel("中心频率/Hz", fontsize=14)
                plt.ylabel("分频振级/dB", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname + "\\" + numname[i] +"\\4-2分幅未计权分频振级-通道{}".format(j + 1) + ".png")
                plt.close()
        frame = tk.Frame(window, width=1920, height=1080)
        frame.grid(row=0, column=0)
        frame.grid_propagate(0)
        def show(event):
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            labelc3[0 + 2 * i].config( text=vl_max_1_am[pre_num][i])
            labelc3[1 + 2 * i].config(text=vl_max_1_fre[pre_num][i])
            for i in range(Pamount[0]):
                labelc4[i] = tk.Label(frame)
                img[i] = Image.open(
                    "cache\\" + dirname + "\\" + numname[pre_num] + "\\4-2分幅未计权分频振级-通道{}".format(i + 1) + ".png")  # 打开图片
                photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                labelc4[i].config(image=photo[i])
                labelc4[i].image = photo[i]
        def show_image62():
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            plt.figure(num=numname[pre_num])
            for i in range(Pamount[pre_num]):
                plt.subplot(Pamount[pre_num], 1, i + 1)
                plt.plot(mid_frep, vl_max_1[pre_num][i])
                plt.ylabel("分频振级/dB", fontsize=14)
                if i == Pamount[pre_num] - 1:
                    plt.xlabel("中心频率", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
            plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = numname
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
        label0 = tk.Label(frame, text='分频振级（加窗未计权）分析结果')
        label0.grid(row=0, column=1, columnspan=2)
        label1 = tk.Label(frame, text='分频振级图')
        label1.grid(row=0, column=4)
        buttun2 = tk.Button(frame, text='显示详图', command=show_image62)
        buttun2.grid(row=0, column=5)
        labelc1 = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc1[i] = tk.Label(frame, text='通道{}'.format(i + 1))
            labelc1[i].grid(row=1 + i * 2, column=0, rowspan=2)
        labelc2 = list(range(Pamount[0] * 2))
        for i in range(Pamount[0]):
            labelc2[0 + 2 * i] = tk.Label(frame, text='分频最大振级/dB')
            labelc2[0 + 2 * i].grid(row=1 + i * 2 + 0, column=1)
            labelc2[1 + 2 * i] = tk.Label(frame, text='分频最大振级对应中心频率/Hz')
            labelc2[1 + 2 * i].grid(row=1 + i * 2 + 1, column=1)
        labelc3 = list(range(Pamount[0] * 2))
        for i in range(Pamount[0]):
            labelc3[0 + 2 * i] = tk.Label(frame, text=vl_max_1_am[0][i])
            labelc3[0 + 2 * i].grid(row=1 + i * 2 + 0, column=2)
            labelc3[1 + 2 * i] = tk.Label(frame, text=vl_max_1_fre[0][i])
            labelc3[1 + 2 * i].grid(row=1 + i * 2 + 1, column=2)
        labelc4 = list(range(Pamount[0]))
        img = list(range(Pamount[0]))
        photo = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc4[i] = tk.Label(frame)
            img[i] = Image.open("cache\\" + dirname + "\\" + numname[0] + "\\4-2分幅未计权分频振级-通道{}".format(i + 1) + ".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            labelc4[i].config(image=photo[i])
            labelc4[i].image = photo[i]
            labelc4[i].grid(row=1 + i * 2, column=3, rowspan=2, columnspan=3)
    if mode==3:
        global  vl_max_1_am_tj, vl_max_1_fre_tj
        mid_fre = [1, 1.25, 1.6, 2, 2.5, 3.15, 4, 5, 6.3, 8, 10, 12.5, 16, 20, 25, 31.5, 40, 50, 63, 80, 100, 125, 160,
                   200, 250, 315, 400, 500, 630, 800, 1000, 1200]
        mid_frep = ['1', '1.25', '1.6', '2', '2.5', '3.15', '4', '5', '6.3', '8', '10', '12.5',
                    '16', '20', '25', '31.5', '40', '50', '63', '80', '100', '125', '160',
                    '200', '250', '315', '400', '500', '630', '800', '1000', '1200']
        vl_max_1 = list(range(len(numname) - 1))
        vl_max_1_am = list(range(len(numname) - 1))
        vl_max_1_fre = list(range(len(numname) - 1))
        vl_max_1_am_tj=list(range(Pamount[0]))
        vl_max_1_fre_tj=list(range(Pamount[0]))
        amp=list(range(Pamount[0]))
        frep=list(range(Pamount[0]))
        for i in range(len(numname)-1):
            vl_max_1[i] = list(range(Pamount[i]))
            vl_max_1_am[i] = list(range(Pamount[i]))
            vl_max_1_fre[i] = list(range(Pamount[i]))
        for i in range(len(numname)-1):
            for j in range(Pamount[i]):
                vl_max_1[i][j] = fft_max_1(data[i][j])
                vl_max_1_am[i][j] = max(vl_max_1[i][j])
                for k in range(len(mid_fre)):
                    if vl_max_1_am[i][j] == vl_max_1[i][j][k]:
                        vl_max_1_fre[i][j] = mid_fre[k]
        for i in range(len(numname)-1):
            for j in range(Pamount[i]):
                plt.figure(figsize=(8, 0.9))
                plt.plot(mid_frep, vl_max_1[i][j])
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.05, right=0.95, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname + "\\" + numname[i] + "\\4-2分幅未计权分频振级-通道{}".format(j + 1) + ".png")
                plt.close()
        for i in range(len(numname)-1):
            for j in range(Pamount[i]):
                plt.figure()
                plt.plot(mid_fre, vl_max_1[i][j])
                plt.xlabel("中心频率/Hz", fontsize=14)
                plt.ylabel("分频振级/dB", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname + "\\" + numname[i] +"\\4-2分幅未计权分频振级-通道{}".format(j + 1) + ".png")
                plt.close()
        for i in range(Pamount[0]):
            vl_max_1_am_tj[i]=list(range(3))
            vl_max_1_fre_tj[i]=list(range(3))
            amp[i]=[]
            frep[i]=[]
            for j in range(len(numname) - 1):
                amp[i].append(vl_max_1_am[j][i])
                frep[i].append(vl_max_1_fre[j][i])
            vl_max_1_am_tj[i][0]=min(vl_max_1_am[0][i], vl_max_1_am[1][i], vl_max_1_am[2][i], vl_max_1_am[3][i], vl_max_1_am[4][i])
            sum1 = 0
            for j in range(len(numname) - 1):
                sum1 += vl_max_1_am[j][i]
            vl_max_1_am_tj[i][1]=round(sum1 / (len(numname) - 1), 0)
            vl_max_1_am_tj[i][2]=max(vl_max_1_am[0][i], vl_max_1_am[1][i], vl_max_1_am[2][i], vl_max_1_am[3][i], vl_max_1_am[4][i])
            vl_max_1_fre_tj[i][0]=min(vl_max_1_fre[0][i], vl_max_1_fre[1][i], vl_max_1_fre[2][i], vl_max_1_fre[3][i], vl_max_1_fre[4][i])
            sum2 = 0
            for j in range(len(numname) - 1):
                sum2 += vl_max_1_fre[j][i]
            vl_max_1_fre_tj[i][1]=round(sum2 / (len(numname) - 1), 0)
            vl_max_1_fre_tj[i][2]=max(vl_max_1_fre[0][i], vl_max_1_fre[1][i], vl_max_1_fre[2][i], vl_max_1_fre[3][i], vl_max_1_fre[4][i])
            plt.figure(figsize=(8, 0.9))
            plt.subplot(1, 2, 1)
            plt.title('分频最大振级变化')
            plt.plot(numname[0:5], amp[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplot(1, 2, 2)
            plt.title('分频振级对应中心频率变化')
            plt.plot(numname[0:5], frep[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.27, top=0.75, wspace=0.25)
            plt.savefig("cache\\" + dirname + "\\4-2分幅未计权分频振级对比分析-通道{}".format(i + 1) + ".png")
            plt.close()
            plt.figure(figsize=(8, 0.9))
            plt.subplot(1, 2, 1)
            plt.title('分频最大振级变化')
            plt.plot(numname[0:5], amp[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplot(1, 2, 2)
            plt.title('分频振级对应中心频率变化')
            plt.plot(numname[0:5], frep[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.27, top=0.75, wspace=0.25)
            plt.savefig("result\\" + dirname + "\\4-2分幅未计权分频振级对比分析-通道{}".format(i + 1) + ".png")
            plt.close()
        frame = tk.Frame(window, width=1920, height=1080)
        frame.grid(row=0, column=0)
        frame.grid_propagate(0)
        def show(event):
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            if pre_num<=len(numname)-2:
                label1.config(text='分频振级图')
                labelc2[0 + 2 * i].config(text='分频最大振级/dB')
                labelc2[1 + 2 * i].config(text='分频最大振级对应中心频率/Hz')
                labelc3[0 + 2 * i].config(text=vl_max_1_am[pre_num][i])
                labelc3[1 + 2 * i].config(text=vl_max_1_fre[pre_num][i])
                for i in range(Pamount[0]):
                    img[i] = Image.open(
                        "cache\\" + dirname + "\\" + numname[pre_num] + "\\4-2分幅未计权分频振级-通道{}".format(
                            i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
            else:
                label1.config(text='对比分析结果')
                for i in range(Pamount[0]):
                    labelc2[0 + 2 * i].config(text='分频最大振级(最小值、均值、最大值）/dB')
                    labelc2[1 + 2 * i].config(text='分频最大振级对应中心频率(最小值、均值、最大值）/Hz')
                    labelc3[0 + 2 * i].config(text=vl_max_1_am_tj[i])
                    labelc3[1 + 2 * i].config(text=vl_max_1_fre_tj[i])
                    img[i] = Image.open("cache\\" + dirname + "\\4-2分幅未计权分频振级对比分析-通道{}".format(i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
        def show_image63():
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            if pre_num <= len(numname) - 2:
                plt.figure(num=numname[pre_num])
                for i in range(Pamount[pre_num]):
                    plt.subplot(Pamount[pre_num], 1, i + 1)
                    plt.plot(mid_frep, vl_max_1[pre_num][i])
                    plt.ylabel("分频振级/dB", fontsize=14)
                    if i == Pamount[pre_num] - 1:
                        plt.xlabel("中心频率/Hz", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                plt.show()
            else:
                plt.figure(num=numname[pre_num])
                for i in range(Pamount[0]):
                    plt.subplot(Pamount[0], 2, i * 2 + 1)
                    plt.plot(numname[0:5], amp[i])
                    plt.ylabel("分频振级/dB", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                    plt.subplot(Pamount[0], 2, i * 2 + 2)
                    plt.plot(numname[0:5], frep[i])
                    plt.ylabel("中心频率/Hz", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = numname
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
        label0 = tk.Label(frame, text='分频振级（加窗未计权）分析结果')
        label0.grid(row=0, column=1, columnspan=2)
        label1 = tk.Label(frame, text='分频振级图')
        label1.grid(row=0, column=4)
        buttun2 = tk.Button(frame, text='显示详图', command=show_image63)
        buttun2.grid(row=0, column=5)
        labelc1 = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc1[i] = tk.Label(frame, text='通道{}'.format(i + 1))
            labelc1[i].grid(row=1 + i * 2, column=0, rowspan=2)
        labelc2 = list(range(Pamount[0] * 2))
        for i in range(Pamount[0]):
            labelc2[0 + 2 * i] = tk.Label(frame, text='分频最大振级/dB')
            labelc2[0 + 2 * i].grid(row=1 + i * 2 + 0, column=1)
            labelc2[1 + 2 * i] = tk.Label(frame, text='分频最大振级对应中心频率/Hz')
            labelc2[1 + 2 * i].grid(row=1 + i * 2 + 1, column=1)
        labelc3 = list(range(Pamount[0] * 2))
        for i in range(Pamount[0]):
            labelc3[0 + 2 * i] = tk.Label(frame, text=vl_max_1_am[0][i])
            labelc3[0 + 2 * i].grid(row=1 + i * 2 + 0, column=2)
            labelc3[1 + 2 * i] = tk.Label(frame, text=vl_max_1_fre[0][i])
            labelc3[1 + 2 * i].grid(row=1 + i * 2 + 1, column=2)
        labelc4 = list(range(Pamount[0]))
        img = list(range(Pamount[0]))
        photo = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc4[i] = tk.Label(frame)
            img[i] = Image.open(
                "cache\\" + dirname + "\\" + numname[0] + "\\4-2分幅未计权分频振级-通道{}".format(i + 1) + ".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            labelc4[i].config(image=photo[i])
            labelc4[i].image = photo[i]
            labelc4[i].grid(row=1 + i * 2, column=3, rowspan=2, columnspan=3)
def vlmaxw():
    global vl_max_w,vl_max_w_am,vl_max_w_fre
    if mode==1:
        mid_fre = [4, 5, 6.3, 8, 10, 12.5, 16, 20, 25, 31.5, 40, 50, 63, 80, 100, 125, 160,200]
        mid_frep = ['4', '5', '6.3', '8', '10', '12.5','16', '20', '25', '31.5', '40', '50', '63', '80', '100', '125', '160','200']
        vl_max_w = list(range(Pamount))
        vl_max_w_am=list(range(Pamount))
        vl_max_w_fre=list(range(Pamount))
        for i in range(Pamount):
            vl_max_w[i] = fft_max_w(data[i])
            vl_max_w_am[i]=max(vl_max_w[i])
            for j in range(len(mid_fre)):
                if vl_max_w_am[i]==vl_max_w[i][j]:
                    vl_max_w_fre[i]=mid_fre[j]
        for i in range(Pamount):
            plt.figure(figsize=(8,0.9))
            plt.plot(mid_frep, vl_max_w[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.05, right=0.95, bottom=0.25, top=0.9)
            plt.savefig("cache\\" + dirname + "\\4-3分幅计权分频振级-通道{}".format(i + 1) + ".png")
            plt.close()
        for i in range(Pamount):
            plt.figure()
            plt.plot(mid_fre, vl_max_w[i])
            plt.xlabel("中心频率/Hz", fontsize=14)
            plt.ylabel("分频振级/dB", fontsize=14)
            plt.xticks(fontsize=14)
            plt.yticks(fontsize=14)
            plt.savefig("result\\" + dirname + "\\4-3分幅计权分频振级-通道{}".format(i + 1) + ".png")
            plt.close()
        frame = tk.Frame(window, width=1920, height=1080)
        frame.grid(row=0, column=0)
        frame.grid_propagate(0)
        def show(event):
            pre_num = var.get()
        def show_image71():
            plt.figure()
            for i in range(Pamount):
                plt.subplot(Pamount, 1, i + 1)
                plt.plot(mid_frep, vl_max_w[i])
                plt.ylabel("分频振级/dB", fontsize=14)
                if i == Pamount - 1:
                    plt.xlabel("中心频率", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
            plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = ['Group1']
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
        label0 = tk.Label(frame, text='分频振级（加窗计权）分析结果')
        label0.grid(row=0, column=1, columnspan=2)
        label1 = tk.Label(frame, text='分频振级图')
        label1.grid(row=0, column=4)
        buttun2 = tk.Button(frame, text='显示详图', command=show_image71)
        buttun2.grid(row=0, column=5)
        labelc1 = list(range(Pamount))
        for i in range(Pamount):
            labelc1[i] = tk.Label(frame, text='通道{}'.format(i + 1))
            labelc1[i].grid(row=1 + i * 2, column=0, rowspan=2)
        labelc2 = list(range(Pamount * 2))
        for i in range(Pamount):
            labelc2[0 + 2 * i] = tk.Label(frame, text='分频最大振级/dB')
            labelc2[0 + 2 * i].grid(row=1 + i * 2 + 0, column=1)
            labelc2[1 + 2 * i] = tk.Label(frame, text='分频最大振级对应中心频率/Hz')
            labelc2[1 + 2 * i].grid(row=1 + i * 2 + 1, column=1)
        labelc3 = list(range(Pamount * 2))
        for i in range(Pamount):
            labelc3[0 + 2 * i] = tk.Label(frame, text=vl_max_w_am[i])
            labelc3[0 + 2 * i].grid(row=1 + i * 2 + 0, column=2)
            labelc3[1 + 2 * i] = tk.Label(frame, text=vl_max_w_fre[i])
            labelc3[1 + 2 * i].grid(row=1 + i * 2 + 1, column=2)
        labelc4 = list(range(Pamount))
        img = list(range(Pamount))
        photo = list(range(Pamount))
        for i in range(Pamount):
            labelc4[i] = tk.Label(frame)
            img[i] = Image.open("cache\\" + dirname + "\\4-3分幅计权分频振级-通道{}".format(i + 1) + ".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            labelc4[i].config(image=photo[i])
            labelc4[i].image = photo[i]
            labelc4[i].grid(row=1 + i * 2, column=3, rowspan=2, columnspan=3)
    if mode==2:
        mid_fre = [4, 5, 6.3, 8, 10, 12.5, 16, 20, 25, 31.5, 40, 50, 63, 80, 100, 125, 160, 200]
        mid_frep = ['4', '5', '6.3', '8', '10', '12.5', '16', '20', '25', '31.5', '40', '50', '63', '80', '100', '125',
                    '160', '200']
        vl_max_w = list(range(len(numname)))
        vl_max_w_am = list(range(len(numname)))
        vl_max_w_fre = list(range(len(numname)))
        for i in range(len(numname)):
            vl_max_w[i] = list(range(Pamount[i]))
            vl_max_w_am[i] = list(range(Pamount[i]))
            vl_max_w_fre[i] = list(range(Pamount[i]))
        for i in range(len(numname)):
            for j in range(Pamount[i]):
                vl_max_w[i][j] = fft_max_w(data[i][j])
                vl_max_w_am[i][j] = max(vl_max_w[i][j])
                for k in range(len(mid_fre)):
                    if vl_max_w_am[i][j] == vl_max_w[i][j][k]:
                        vl_max_w_fre[i][j] = mid_fre[k]
        for i in range(len(numname)):
            for j in range(Pamount[i]):
                plt.figure(figsize=(8, 0.9))
                plt.plot(mid_frep, vl_max_w[i][j])
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.05, right=0.95, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname + "\\" + numname[i] + "\\4-3分幅计权分频振级-通道{}".format(j + 1) + ".png")
                plt.close()
        for i in range(len(numname)):
            for j in range(Pamount[i]):
                plt.figure()
                plt.plot(mid_fre, vl_max_w[i][j])
                plt.xlabel("中心频率/Hz", fontsize=14)
                plt.ylabel("分频振级/dB", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname + "\\" + numname[i] + "\\4-3分幅计权分频振级-通道{}".format(j + 1) + ".png")
                plt.close()
        frame = tk.Frame(window, width=1920, height=1080)
        frame.grid(row=0, column=0)
        frame.grid_propagate(0)
        def show(event):
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            labelc3[0 + 2 * i].config( text=vl_max_w_am[pre_num][i])
            labelc3[1 + 2 * i].config(text=vl_max_w_fre[pre_num][i])
            for i in range(Pamount[0]):
                labelc4[i] = tk.Label(frame)
                img[i] = Image.open(
                    "cache\\" + dirname + "\\" + numname[pre_num] + "\\4-3分幅计权分频振级-通道{}".format(i + 1) + ".png")  # 打开图片
                photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                labelc4[i].config(image=photo[i])
                labelc4[i].image = photo[i]
        def show_image72():
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            plt.figure(num=numname[pre_num])
            for i in range(Pamount[pre_num]):
                plt.subplot(Pamount[pre_num], 1, i + 1)
                plt.plot(mid_frep, vl_max_w[pre_num][i])
                plt.ylabel("分频振级/dB", fontsize=14)
                if i == Pamount[pre_num] - 1:
                    plt.xlabel("中心频率", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
            plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = numname
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
        label0 = tk.Label(frame, text='分频振级（加窗计权）分析结果')
        label0.grid(row=0, column=1, columnspan=2)
        label1 = tk.Label(frame, text='分频振级图')
        label1.grid(row=0, column=4)
        buttun2 = tk.Button(frame, text='显示详图', command=show_image72)
        buttun2.grid(row=0, column=5)
        labelc1 = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc1[i] = tk.Label(frame, text='通道{}'.format(i + 1))
            labelc1[i].grid(row=1 + i * 2, column=0, rowspan=2)
        labelc2 = list(range(Pamount[0] * 2))
        for i in range(Pamount[0]):
            labelc2[0 + 2 * i] = tk.Label(frame, text='分频最大振级/dB')
            labelc2[0 + 2 * i].grid(row=1 + i * 2 + 0, column=1)
            labelc2[1 + 2 * i] = tk.Label(frame, text='分频最大振级对应中心频率/Hz')
            labelc2[1 + 2 * i].grid(row=1 + i * 2 + 1, column=1)
        labelc3 = list(range(Pamount[0] * 2))
        for i in range(Pamount[0]):
            labelc3[0 + 2 * i] = tk.Label(frame, text=vl_max_w_am[0][i])
            labelc3[0 + 2 * i].grid(row=1 + i * 2 + 0, column=2)
            labelc3[1 + 2 * i] = tk.Label(frame, text=vl_max_w_fre[0][i])
            labelc3[1 + 2 * i].grid(row=1 + i * 2 + 1, column=2)
        labelc4 = list(range(Pamount[0]))
        img = list(range(Pamount[0]))
        photo = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc4[i] = tk.Label(frame)
            img[i] = Image.open("cache\\" + dirname + "\\" + numname[0] + "\\4-3分幅计权分频振级-通道{}".format(i + 1) + ".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            labelc4[i].config(image=photo[i])
            labelc4[i].image = photo[i]
            labelc4[i].grid(row=1 + i * 2, column=3, rowspan=2, columnspan=3)
    if mode==3:
        global  vl_max_w_am_tj, vl_max_w_fre_tj
        mid_fre = [4, 5, 6.3, 8, 10, 12.5, 16, 20, 25, 31.5, 40, 50, 63, 80, 100, 125, 160,200]
        mid_frep = [ '4', '5', '6.3', '8', '10', '12.5','16', '20', '25', '31.5', '40', '50', '63', '80', '100', '125', '160','200']
        vl_max_w = list(range(len(numname) - 1))
        vl_max_w_am = list(range(len(numname) - 1))
        vl_max_w_fre = list(range(len(numname) - 1))
        vl_max_w_am_tj=list(range(Pamount[0]))
        vl_max_w_fre_tj=list(range(Pamount[0]))
        amp=list(range(Pamount[0]))
        frep=list(range(Pamount[0]))
        for i in range(len(numname)-1):
            vl_max_w[i] = list(range(Pamount[i]))
            vl_max_w_am[i] = list(range(Pamount[i]))
            vl_max_w_fre[i] = list(range(Pamount[i]))
        for i in range(len(numname)-1):
            for j in range(Pamount[i]):
                vl_max_w[i][j] = fft_max_w(data[i][j])
                vl_max_w_am[i][j] = max(vl_max_w[i][j])
                for k in range(len(mid_fre)):
                    if vl_max_w_am[i][j] == vl_max_w[i][j][k]:
                        vl_max_w_fre[i][j] = mid_fre[k]
        for i in range(len(numname)-1):
            for j in range(Pamount[i]):
                plt.figure(figsize=(8, 0.9))
                plt.plot(mid_frep, vl_max_w[i][j])
                plt.xticks(fontsize=9)
                plt.yticks(fontsize=9)
                plt.subplots_adjust(left=0.05, right=0.95, bottom=0.25, top=0.9)
                plt.savefig("cache\\" + dirname + "\\" + numname[i] + "\\4-3分幅计权分频振级-通道{}".format(j + 1) + ".png")
                plt.close()
        for i in range(len(numname)-1):
            for j in range(Pamount[i]):
                plt.figure()
                plt.plot(mid_fre, vl_max_w[i][j])
                plt.xlabel("中心频率/Hz", fontsize=14)
                plt.ylabel("分频振级/dB", fontsize=14)
                plt.xticks(fontsize=14)
                plt.yticks(fontsize=14)
                plt.savefig("result\\" + dirname  + "\\" + numname[i] + "\\4-3分幅计权分频振级-通道{}".format(j + 1) + ".png")
                plt.close()
        for i in range(Pamount[0]):
            vl_max_w_am_tj[i]=list(range(3))
            vl_max_w_fre_tj[i]=list(range(3))
            amp[i]=[]
            frep[i]=[]
            for j in range(len(numname) - 1):
                amp[i].append(vl_max_w_am[j][i])
                frep[i].append(vl_max_w_fre[j][i])
            vl_max_w_am_tj[i][0]=min(vl_max_w_am[0][i], vl_max_w_am[1][i], vl_max_w_am[2][i], vl_max_w_am[3][i], vl_max_w_am[4][i])
            sum1 = 0
            for j in range(len(numname) - 1):
                sum1 += vl_max_w_am[j][i]
            vl_max_w_am_tj[i][1]=round(sum1 / (len(numname) - 1), 0)
            vl_max_w_am_tj[i][2]=max(vl_max_w_am[0][i], vl_max_w_am[1][i], vl_max_w_am[2][i], vl_max_w_am[3][i], vl_max_w_am[4][i])
            vl_max_w_fre_tj[i][0]=min(vl_max_w_fre[0][i], vl_max_w_fre[1][i], vl_max_w_fre[2][i], vl_max_w_fre[3][i], vl_max_w_fre[4][i])
            sum2 = 0
            for j in range(len(numname) - 1):
                sum2 += vl_max_w_fre[j][i]
            vl_max_w_fre_tj[i][1]=round(sum2 / (len(numname) - 1), 0)
            vl_max_w_fre_tj[i][2]=max(vl_max_w_fre[0][i], vl_max_w_fre[1][i], vl_max_w_fre[2][i], vl_max_w_fre[3][i], vl_max_w_fre[4][i])
            plt.figure(figsize=(8, 0.9))
            plt.subplot(1, 2, 1)
            plt.title('分频最大振级变化')
            plt.plot(numname[0:5], amp[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplot(1, 2, 2)
            plt.title('分频振级对应中心频率变化')
            plt.plot(numname[0:5], frep[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.27, top=0.75, wspace=0.25)
            plt.savefig("cache\\" + dirname + "\\4-3分幅计权分频振级对比分析-通道{}".format(i + 1) + ".png")
            plt.close()
            plt.figure(figsize=(8, 0.9))
            plt.subplot(1, 2, 1)
            plt.title('分频最大振级变化')
            plt.plot(numname[0:5], amp[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplot(1, 2, 2)
            plt.title('分频振级对应中心频率变化')
            plt.plot(numname[0:5], frep[i])
            plt.xticks(fontsize=9)
            plt.yticks(fontsize=9)
            plt.subplots_adjust(left=0.1, right=0.9, bottom=0.27, top=0.75, wspace=0.25)
            plt.savefig("result\\" + dirname + "\\4-3分幅计权分频振级对比分析-通道{}".format(i + 1) + ".png")
            plt.close()
        frame = tk.Frame(window, width=1920, height=1080)
        frame.grid(row=0, column=0)
        frame.grid_propagate(0)
        def show(event):
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            if pre_num<=len(numname)-2:
                label1.config(text='分频振级图')
                labelc2[0 + 2 * i].config(text='分频最大振级/dB')
                labelc2[1 + 2 * i].config(text='分频最大振级对应中心频率/Hz')
                labelc3[0 + 2 * i].config(text=vl_max_w_am[pre_num][i])
                labelc3[1 + 2 * i].config(text=vl_max_w_fre[pre_num][i])
                for i in range(Pamount[0]):
                    img[i] = Image.open(
                        "cache\\" + dirname + "\\" + numname[pre_num] + "\\4-3分幅计权分频振级-通道{}".format(
                            i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
            else:
                label1.config(text='对比分析结果')
                for i in range(Pamount[0]):
                    labelc2[0 + 2 * i].config(text='分频最大振级(最小值、均值、最大值）/dB')
                    labelc2[1 + 2 * i].config(text='分频最大振级对应中心频率(最小值、均值、最大值）/Hz')
                    labelc3[0 + 2 * i].config(text=vl_max_w_am_tj[i])
                    labelc3[1 + 2 * i].config(text=vl_max_w_fre_tj[i])
                    img[i] = Image.open("cache\\" + dirname + "\\4-3分幅计权分频振级对比分析-通道{}".format(i + 1) + ".png")  # 打开图片
                    photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
                    labelc4[i].config(image=photo[i])
                    labelc4[i].image = photo[i]
        def show_image73():
            a = var.get()
            for i in range(len(numname)):
                if a == numname[i]:
                    pre_num = i
            if pre_num <= len(numname) - 2:
                plt.figure(num=numname[pre_num])
                for i in range(Pamount[pre_num]):
                    plt.subplot(Pamount[pre_num], 1, i + 1)
                    plt.plot(mid_frep, vl_max_w[pre_num][i])
                    plt.ylabel("分频振级/dB", fontsize=14)
                    if i == Pamount[pre_num] - 1:
                        plt.xlabel("中心频率/Hz", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                plt.show()
            else:
                plt.figure(num=numname[pre_num])
                for i in range(Pamount[0]):
                    plt.subplot(Pamount[0], 2, i * 2 + 1)
                    plt.plot(numname[0:5], amp[i])
                    plt.ylabel("分频振级/dB", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                    plt.subplot(Pamount[0], 2, i * 2 + 2)
                    plt.plot(numname[0:5], frep[i])
                    plt.ylabel("中心频率/Hz", fontsize=14)
                    plt.xticks(fontsize=14)
                    plt.yticks(fontsize=14)
                plt.show()
        var = tkinter.StringVar()
        combobox = tkinter.ttk.Combobox(frame, textvariable=var, width=9)
        combobox['value'] = numname
        combobox.current(0)
        combobox.bind('<<ComboboxSelected>>', show)
        combobox.grid(row=0, column=0)
        label0 = tk.Label(frame, text='分频振级（加窗计权）分析结果')
        label0.grid(row=0, column=1, columnspan=2)
        label1 = tk.Label(frame, text='分频振级图')
        label1.grid(row=0, column=4)
        buttun2 = tk.Button(frame, text='显示详图', command=show_image73)
        buttun2.grid(row=0, column=5)
        labelc1 = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc1[i] = tk.Label(frame, text='通道{}'.format(i + 1))
            labelc1[i].grid(row=1 + i * 2, column=0, rowspan=2)
        labelc2 = list(range(Pamount[0] * 2))
        for i in range(Pamount[0]):
            labelc2[0 + 2 * i] = tk.Label(frame, text='分频最大振级/dB')
            labelc2[0 + 2 * i].grid(row=1 + i * 2 + 0, column=1)
            labelc2[1 + 2 * i] = tk.Label(frame, text='分频最大振级对应中心频率/Hz')
            labelc2[1 + 2 * i].grid(row=1 + i * 2 + 1, column=1)
        labelc3 = list(range(Pamount[0] * 2))
        for i in range(Pamount[0]):
            labelc3[0 + 2 * i] = tk.Label(frame, text=vl_max_w_am[0][i])
            labelc3[0 + 2 * i].grid(row=1 + i * 2 + 0, column=2)
            labelc3[1 + 2 * i] = tk.Label(frame, text=vl_max_w_fre[0][i])
            labelc3[1 + 2 * i].grid(row=1 + i * 2 + 1, column=2)
        labelc4 = list(range(Pamount[0]))
        img = list(range(Pamount[0]))
        photo = list(range(Pamount[0]))
        for i in range(Pamount[0]):
            labelc4[i] = tk.Label(frame)
            img[i] = Image.open(
                "cache\\" + dirname + "\\" + numname[0] + "\\4-3分幅计权分频振级-通道{}".format(i + 1) + ".png")  # 打开图片
            photo[i] = ImageTk.PhotoImage(img[i])  # 用PIL模块的PhotoImage打开
            labelc4[i].config(image=photo[i])
            labelc4[i].image = photo[i]
            labelc4[i].grid(row=1 + i * 2, column=3, rowspan=2, columnspan=3)
def outfile():
    if mode==1:
        doc = Document()
        tit = doc.add_paragraph('轨道环境振动分析报告')
        tit.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_heading('1.振动测试原始数据概览')
        doc.add_paragraph('各个通道统计数据见表1-1')
        par1_1 = doc.add_paragraph('表1-1')
        par1_1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table1 = doc.add_table(Pamount + 1, 4)
        column = ['通道编号', '采样时间/s', '采样点数', '采样频率/Hz']
        for i in range(4):
            table1.cell(0, i).text = column[i]
        for i in range(Pamount):
            table1.cell(i + 1, 0).text = str(i + 1)
            table1.cell(i + 1, 1).text = str(time_last[i])
            table1.cell(i + 1, 2).text = str(num_save[i])
            table1.cell(i + 1, 3).text = str(fr_save[i])
        doc.add_paragraph('利用结果绘制时程图，见下图')
        par1_2 = list(range(Pamount))
        for i in range(Pamount):
            doc.add_picture('result\\' + dirname + '\\0时程图-通道{}'.format(i + 1) + '.png', height=Cm(5))
            par1_2[i] = doc.add_paragraph('图1-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + '时程图')
            par1_2[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_heading('2.时域分析结果')
        doc.add_paragraph('各个通道统计数据时域分析结果见表2-1')
        par2_1 = doc.add_paragraph('表2-1')
        par2_1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        column = ['通道编号', '零点偏移量', '时间趋势项大小', '异常值个数', '加速度最值/ms-2', '加速度有效值/ms-2', 'VAL/dB']
        table2 = doc.add_table(Pamount + 1, 7)
        for i in range(7):
            table2.cell(0, i).text = column[i]
        for i in range(Pamount):
            table2.cell(i + 1, 0).text = str(i + 1)
            table2.cell(i + 1, 1).text = str(round(x[i], 3))
            table2.cell(i + 1, 2).text = str(round(y[i], 3))
            table2.cell(i + 1, 3).text = str(0)
            table2.cell(i + 1, 4).text = str(round(vlmax[i], 3))
            table2.cell(i + 1, 5).text = str(round(vlarms[i], 3))
            table2.cell(i + 1, 6).text = str(round(val[i]))
        doc.add_paragraph('经过数据处理后绘制时程图，见下图')
        par2_2 = list(range(Pamount))
        for i in range(Pamount):
            doc.add_picture("result\\" + dirname + "\\1处理后时程图-通道{}".format(i + 1) + ".png", height=Cm(5))
            par2_2[i] = doc.add_paragraph('图2-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + '处理后时程图')
            par2_2[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_heading('3.傅里叶变换（显著频率）')
        doc.add_paragraph('将时域数据进行频域变换，转换到频域进行分析，计算了各通道振动的显著频率及其对应的振幅大小，结果见表3-1')
        par3_1 = doc.add_paragraph('表3-1')
        par3_1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        column = ['通道编号', '振动显著频率/Hz', '显著频率对应振幅值/ms-2', '功率最值/w']
        table3 = doc.add_table(Pamount + 1, 4)
        for i in range(4):
            table3.cell(0, i).text = column[i]
        for i in range(Pamount):
            table3.cell(i + 1, 0).text = str(i + 1)
            table3.cell(i + 1, 1).text = str(obs_fre[i])
            table3.cell(i + 1, 2).text = str(max_am[i])
            table3.cell(i + 1, 3).text = str(round(max_power[i], 3))
        doc.add_paragraph('经过频域分析后绘制频谱与功率谱，见下图')
        par3_2 = list(range(Pamount))
        table3_1 = list(range(Pamount))
        for i in range(Pamount):
            table3_1[i] = doc.add_table(1, 2)
            cell = table3_1[i].cell(0, 0)
            ph = cell.paragraphs[0]
            run = ph.add_run()
            run.add_picture("result\\" + dirname + "\\2-1频谱-通道{}".format(i + 1) + ".png", height=Cm(5))
            cell = table3_1[i].cell(0, 1)
            ph = cell.paragraphs[0]
            run = ph.add_run()
            run.add_picture("result\\" + dirname + "\\2-2功率谱-通道{}".format(i + 1) + ".png", height=Cm(5))
            par3_2[i] = doc.add_paragraph('图3-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + '频谱及功率谱')
            par3_2[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_heading('4.最大Z振级')
        doc.add_paragraph('利用频域数据进行最大Z振级分析，结果见表4-1')
        par4_1 = doc.add_paragraph('表4-1')
        par4_1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        column = ['通道编号', '最大Z振级（不计权）/dB', '最大Z振级（85版）/dB', '最大Z振级（97版）/dB']
        table4 = doc.add_table(Pamount + 1, 4)
        for i in range(4):
            table4.cell(0, i).text = column[i]
        for i in range(Pamount):
            table4.cell(i + 1, 0).text = str(i + 1)
            table4.cell(i + 1, 1).text = str(vl_maxz[i])
            table4.cell(i + 1, 2).text = str(vl_maxzw[i])
            table4.cell(i + 1, 3).text = str(vl_maxzwk[i])
        doc.add_paragraph('经过频域分析后绘制频谱与功率谱，见下图')
        par4_2 = list(range(Pamount))
        table4_1 = list(range(Pamount))
        for i in range(Pamount):
            table4_1[i] = doc.add_table(1, 3)
            cell = table4_1[i].cell(0, 0)
            ph = cell.paragraphs[0]
            run = ph.add_run()
            run.add_picture("result\\" + dirname + "\\3-1Z振级-不计权-通道{}".format(i + 1) + ".png", height=Cm(4))
            cell = table4_1[i].cell(0, 1)
            ph = cell.paragraphs[0]
            run = ph.add_run()
            run.add_picture("result\\" + dirname + "\\3-2Z振级-85版-通道{}".format(i + 1) + ".png", height=Cm(4))
            cell = table4_1[i].cell(0, 2)
            ph = cell.paragraphs[0]
            run = ph.add_run()
            run.add_picture("result\\" + dirname + "\\3-3Z振级-97版-通道{}".format(i + 1) + ".png", height=Cm(4))
            par4_2[i] = doc.add_paragraph('图4-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + 'Z振级曲线')
            par4_2[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_heading('5.分频最大振级')
        doc.add_paragraph('利用频域数据进行分频最大振级分析，未分幅未计权时结果见表5-1')
        par5_1 = doc.add_paragraph('表5-1')
        par5_1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table5 = doc.add_table(33, Pamount + 1)
        for i in range(Pamount + 1):
            if i == 0:
                table5.cell(0, 0).text = '中心频率/通道编号'
            else:
                table5.cell(0, i).text = str(i)
        mid_fre0 = [1, 1.25, 1.6, 2, 2.5, 3.15, 4, 5, 6.3, 8, 10, 12.5, 16, 20, 25, 31.5, 40, 50, 63, 80, 100, 125, 160,
                    200,250, 315, 400, 500, 630, 800, 1000, 1200]
        for i in range(len(mid_fre0)):
            table5.cell(i + 1, 0).text = str(mid_fre0[i])
        for i in range(len(mid_fre0)):
            for j in range(Pamount):
                table5.cell(i + 1, j + 1).text = str(vl_max_0[j][i])
        table5a=doc.add_table(Pamount + 1, 3)
        table5a.cell(0,0).text='通道编号'
        table5a.cell(0, 1).text = '分频最大振级/dB'
        table5a.cell(0, 2).text = '分频最大振级对应中心频率/Hz'
        for i in range(Pamount):
            table5a.cell(i+1,0).text=str(i+1)
            table5a.cell(i + 1, 1).text = str(vl_max_0_am[i])
            table5a.cell(i + 1, 2).text = str(vl_max_0_fre[i])
        doc.add_paragraph('绘制未分幅未计权分频最大振级图，见下图')
        par5_2 = list(range(Pamount))
        for i in range(Pamount):
            doc.add_picture("result\\" + dirname + "\\4-1未分幅未计权分频振级-通道{}".format(i + 1) + ".png", height=Cm(5))
            par5_2[i] = doc.add_paragraph('图5-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + '未分幅未计权分频最大振级')
            par5_2[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph('利用频域数据进行分频最大振级分析，分幅未计权时结果见表5-2')
        par5_3 = doc.add_paragraph('表5-2')
        par5_3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table51 = doc.add_table(33, Pamount + 1)
        for i in range(Pamount + 1):
            if i == 0:
                table51.cell(0, 0).text = '中心频率/通道编号'
            else:
                table51.cell(0, i).text = str(i)
        mid_fre1 = [4, 5, 6.3, 8, 10, 12.5, 16, 20, 25, 31.5, 40, 50, 63, 80, 100, 125, 160, 200]
        for i in range(32):
            table51.cell(i + 1, 0).text = str(mid_fre0[i])
        for i in range(32):
            for j in range(Pamount):
                table51.cell(i + 1, j + 1).text = str(vl_max_1[j][i])
        table5b = doc.add_table(Pamount + 1, 3)
        table5b.cell(0, 0).text = '通道编号'
        table5b.cell(0, 1).text = '分频最大振级/dB'
        table5b.cell(0, 2).text = '分频最大振级对应中心频率/Hz'
        for i in range(Pamount):
            table5b.cell(i + 1, 0).text = str(i + 1)
            table5b.cell(i + 1, 1).text = str(vl_max_1_am[i])
            table5b.cell(i + 1, 2).text = str(vl_max_1_fre[i])
        doc.add_paragraph('绘制分幅未计权分频最大振级图，见下图')
        par5_4 = list(range(Pamount))
        for i in range(Pamount):
            doc.add_picture("result\\" + dirname + "\\4-2分幅未计权分频振级-通道{}".format(i + 1) + ".png", height=Cm(5))
            par5_4[i] = doc.add_paragraph('图5-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + '加窗计权分频最大振级')
            par5_4[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph('利用频域数据进行分频最大振级分析，分幅计权时结果见表5-3')
        par5_4 = doc.add_paragraph('表5-3')
        par5_4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table53 = doc.add_table(19, Pamount + 1)
        for i in range(Pamount + 1):
            if i == 0:
                table53.cell(0, 0).text = '中心频率/通道编号'
            else:
                table53.cell(0, i).text = str(i)
        for i in range(18):
            table53.cell(i + 1, 0).text = str(mid_fre1[i])
        for i in range(18):
            for j in range(Pamount):
                table53.cell(i + 1, j + 1).text = str(vl_max_w[j][i])
        table5c = doc.add_table(Pamount + 1, 3)
        table5c.cell(0, 0).text = '通道编号'
        table5c.cell(0, 1).text = '分频最大振级/dB'
        table5c.cell(0, 2).text = '分频最大振级对应中心频率/Hz'
        for i in range(Pamount):
            table5c.cell(i + 1, 0).text = str(i + 1)
            table5c.cell(i + 1, 1).text = str(vl_max_w_am[i])
            table5c.cell(i + 1, 2).text = str(vl_max_w_fre[i])
        doc.add_paragraph('绘制分幅未计权分频最大振级图，见下图')
        par5_5 = list(range(Pamount))
        for i in range(Pamount):
            doc.add_picture("result\\" + dirname + "\\4-3分幅计权分频振级-通道{}".format(i + 1) + ".png", height=Cm(5))
            par5_5[i] = doc.add_paragraph('图5-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + '加窗计权分频最大振级')
            par5_5[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.save("result\\" + dirname + "\\环境振动数据分析报告.docx")
    else:
        if mode==2:
            numname0=len(numname)
        else:
            numname0=len(numname)-1
        for k in range(numname0):
            doc = Document()
            tit = doc.add_paragraph('轨道环境振动分析报告')
            tit.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_heading('1.振动测试原始数据概览')
            doc.add_paragraph('各个通道统计数据见表1-1')
            par1_1 = doc.add_paragraph('表1-1')
            par1_1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table1 = doc.add_table(Pamount[k] + 1, 4)
            column = ['通道编号', '采样时间/s', '采样点数', '采样频率/Hz']
            for i in range(4):
                table1.cell(0, i).text = column[i]
            for i in range(Pamount[k]):
                table1.cell(i + 1, 0).text = str(i + 1)
                table1.cell(i + 1, 1).text = str(time_last[k][i])
                table1.cell(i + 1, 2).text = str(num_save[k][i])
                table1.cell(i + 1, 3).text = str(fr_save[k][i])
            doc.add_paragraph('利用结果绘制时程图，见下图')
            par1_2 = list(range(Pamount[k]))
            for i in range(Pamount[k]):
                doc.add_picture('result\\' + dirname +'\\'+numname[k]+ '\\0时程图-通道{}'.format(i + 1) + '.png', height=Cm(5))
                par1_2[i] = doc.add_paragraph('图1-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + '时程图')
                par1_2[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_heading('2.时域分析结果')
            doc.add_paragraph('各个通道统计数据时域分析结果见表2-1')
            par2_1 = doc.add_paragraph('表2-1')
            par2_1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            column = ['通道编号', '零点偏移量', '时间趋势项大小', '异常值个数', '加速度最值/ms-2', '加速度有效值/ms-2', 'VAL/dB']
            table2 = doc.add_table(Pamount[k] + 1, 7)
            for i in range(7):
                table2.cell(0, i).text = column[i]
            for i in range(Pamount[k]):
                table2.cell(i + 1, 0).text = str(i + 1)
                table2.cell(i + 1, 1).text = str(round(x[k][i], 3))
                table2.cell(i + 1, 2).text = str(round(y[k][i], 3))
                table2.cell(i + 1, 3).text = str(0)
                table2.cell(i + 1, 4).text = str(round(vlmax[k][i], 3))
                table2.cell(i + 1, 5).text = str(round(vlarms[k][i], 3))
                table2.cell(i + 1, 6).text = str(round(val[k][i]))
            doc.add_paragraph('经过数据处理后绘制时程图，见下图')
            par2_2 = list(range(Pamount[k]))
            for i in range(Pamount[k]):
                doc.add_picture("result\\" + dirname + '\\'+numname[k]+"\\1处理后时程图-通道{}".format(i + 1) + ".png", height=Cm(5))
                par2_2[i] = doc.add_paragraph('图2-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + '处理后时程图')
                par2_2[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_heading('3.傅里叶变换（显著频率）')
            doc.add_paragraph('将时域数据进行频域变换，转换到频域进行分析，计算了各通道振动的显著频率及其对应的振幅大小，结果见表3-1')
            par3_1 = doc.add_paragraph('表3-1')
            par3_1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            column = ['通道编号', '振动显著频率/Hz', '显著频率对应振幅值/ms-2', '功率最值/w']
            table3 = doc.add_table(Pamount[k] + 1, 4)
            for i in range(4):
                table3.cell(0, i).text = column[i]
            for i in range(Pamount[k]):
                table3.cell(i + 1, 0).text = str(i + 1)
                table3.cell(i + 1, 1).text = str(obs_fre[k][i])
                table3.cell(i + 1, 2).text = str(max_am[k][i])
                table3.cell(i + 1, 3).text = str(round(max_power[k][i], 3))
            doc.add_paragraph('经过频域分析后绘制频谱与功率谱，见下图')
            par3_2 = list(range(Pamount[k]))
            table3_1 = list(range(Pamount[k]))
            for i in range(Pamount[k]):
                table3_1[i] = doc.add_table(1, 2)
                cell = table3_1[i].cell(0, 0)
                ph = cell.paragraphs[0]
                run = ph.add_run()
                run.add_picture("result\\" + dirname + '\\'+numname[k]+"\\2-1频谱-通道{}".format(i + 1) + ".png", height=Cm(5))
                cell = table3_1[i].cell(0, 1)
                ph = cell.paragraphs[0]
                run = ph.add_run()
                run.add_picture("result\\" + dirname +'\\'+numname[k]+ "\\2-2功率谱-通道{}".format(i + 1) + ".png", height=Cm(5))
                par3_2[i] = doc.add_paragraph('图3-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + '频谱及功率谱')
                par3_2[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_heading('4.最大Z振级')
            doc.add_paragraph('利用频域数据进行最大Z振级分析，结果见表4-1')
            par4_1 = doc.add_paragraph('表4-1')
            par4_1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            column = ['通道编号', '最大Z振级（不计权）/dB', '最大Z振级（85版）/dB', '最大Z振级（97版）/dB']
            table4 = doc.add_table(Pamount[k] + 1, 4)
            for i in range(4):
                table4.cell(0, i).text = column[i]
            for i in range(Pamount[k]):
                table4.cell(i + 1, 0).text = str(i + 1)
                table4.cell(i + 1, 1).text = str(vl_maxz[k][i])
                table4.cell(i + 1, 2).text = str(vl_maxzw[k][i])
                table4.cell(i + 1, 3).text = str(vl_maxzwk[k][i])
            doc.add_paragraph('经过频域分析后绘制频谱与功率谱，见下图')
            par4_2 = list(range(Pamount[k]))
            table4_1 = list(range(Pamount[k]))
            for i in range(Pamount[k]):
                table4_1[i] = doc.add_table(1, 3)
                cell = table4_1[i].cell(0, 0)
                ph = cell.paragraphs[0]
                run = ph.add_run()
                run.add_picture("result\\" + dirname + '\\'+numname[k]+"\\3-1Z振级-不计权-通道{}".format(i + 1) + ".png", height=Cm(4))
                cell = table4_1[i].cell(0, 1)
                ph = cell.paragraphs[0]
                run = ph.add_run()
                run.add_picture("result\\" + dirname + '\\'+numname[k]+"\\3-2Z振级-85版-通道{}".format(i + 1) + ".png", height=Cm(4))
                cell = table4_1[i].cell(0, 2)
                ph = cell.paragraphs[0]
                run = ph.add_run()
                run.add_picture("result\\" + dirname + '\\'+numname[k]+"\\3-3Z振级-97版-通道{}".format(i + 1) + ".png", height=Cm(4))
                par4_2[i] = doc.add_paragraph('图4-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + 'Z振级曲线')
                par4_2[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_heading('5.分频最大振级')
            doc.add_paragraph('利用频域数据进行分频最大振级分析，未分幅未计权时结果见表5-1')
            par5_1 = doc.add_paragraph('表5-1')
            par5_1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table5 = doc.add_table(33, Pamount[k] + 1)
            for i in range(Pamount[k] + 1):
                if i == 0:
                    table5.cell(0, 0).text = '中心频率/通道编号'
                else:
                    table5.cell(0, i).text = str(i)
            mid_fre0 = [1, 1.25, 1.6, 2, 2.5, 3.15, 4, 5, 6.3, 8, 10, 12.5, 16, 20, 25, 31.5, 40, 50, 63, 80, 100, 125,
                        160,
                        200, 250, 315, 400, 500, 630, 800, 1000, 1200]
            for i in range(len(mid_fre0)):
                table5.cell(i + 1, 0).text = str(mid_fre0[i])
            for i in range(len(mid_fre0)):
                for j in range(Pamount[k]):
                    table5.cell(i + 1, j + 1).text = str(vl_max_0[k][j][i])
            table5a = doc.add_table(Pamount[k] + 1, 3)
            table5a.cell(0, 0).text = '通道编号'
            table5a.cell(0, 1).text = '分频最大振级/dB'
            table5a.cell(0, 2).text = '分频最大振级对应中心频率/Hz'
            for i in range(Pamount[k]):
                table5a.cell(i + 1, 0).text = str(i + 1)
                table5a.cell(i + 1, 1).text = str(vl_max_0_am[k][i])
                table5a.cell(i + 1, 2).text = str(vl_max_0_fre[k][i])
            doc.add_paragraph('绘制未分幅未计权分频最大振级图，见下图')
            par5_2 = list(range(Pamount[k]))
            for i in range(Pamount[k]):
                doc.add_picture("result\\" + dirname + '\\'+numname[k]+"\\4-1未分幅未计权分频振级-通道{}".format(i + 1) + ".png", height=Cm(5))
                par5_2[i] = doc.add_paragraph('图5-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + '未分幅未计权分频最大振级')
                par5_2[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph('利用频域数据进行分频最大振级分析，分幅未计权时结果见表5-2')
            par5_3 = doc.add_paragraph('表5-2')
            par5_3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table51 = doc.add_table(33, Pamount[k] + 1)
            for i in range(Pamount[k] + 1):
                if i == 0:
                    table51.cell(0, 0).text = '中心频率/通道编号'
                else:
                    table51.cell(0, i).text = str(i)
            mid_fre1 = [4, 5, 6.3, 8, 10, 12.5, 16, 20, 25, 31.5, 40, 50, 63, 80, 100, 125, 160, 200]
            for i in range(32):
                table51.cell(i + 1, 0).text = str(mid_fre0[i])
            for i in range(32):
                for j in range(Pamount[k]):
                    table51.cell(i + 1, j + 1).text = str(vl_max_1[k][j][i])
            table5b = doc.add_table(Pamount[k] + 1, 3)
            table5b.cell(0, 0).text = '通道编号'
            table5b.cell(0, 1).text = '分频最大振级/dB'
            table5b.cell(0, 2).text = '分频最大振级对应中心频率/Hz'
            for i in range(Pamount[k]):
                table5b.cell(i + 1, 0).text = str(i + 1)
                table5b.cell(i + 1, 1).text = str(vl_max_1_am[k][i])
                table5b.cell(i + 1, 2).text = str(vl_max_1_fre[k][i])
            doc.add_paragraph('绘制分幅未计权分频最大振级图，见下图')
            par5_4 = list(range(Pamount[k]))
            for i in range(Pamount[k]):
                doc.add_picture("result\\" + dirname +'\\'+numname[k]+ "\\4-2分幅未计权分频振级-通道{}".format(i + 1) + ".png", height=Cm(5))
                par5_4[i] = doc.add_paragraph('图5-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + '加窗计权分频最大振级')
                par5_4[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph('利用频域数据进行分频最大振级分析，分幅计权时结果见表5-3')
            par5_4 = doc.add_paragraph('表5-3')
            par5_4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table53 = doc.add_table(19, Pamount[k] + 1)
            for i in range(Pamount[k] + 1):
                if i == 0:
                    table53.cell(0, 0).text = '中心频率/通道编号'
                else:
                    table53.cell(0, i).text = str(i)
            for i in range(18):
                table53.cell(i + 1, 0).text = str(mid_fre1[i])
            for i in range(18):
                for j in range(Pamount[k]):
                    table53.cell(i + 1, j + 1).text = str(vl_max_w[k][j][i])
            table5c = doc.add_table(Pamount[k] + 1, 3)
            table5c.cell(0, 0).text = '通道编号'
            table5c.cell(0, 1).text = '分频最大振级/dB'
            table5c.cell(0, 2).text = '分频最大振级对应中心频率/Hz'
            for i in range(Pamount[k]):
                table5c.cell(i + 1, 0).text = str(i + 1)
                table5c.cell(i + 1, 1).text = str(vl_max_w_am[k][i])
                table5c.cell(i + 1, 2).text = str(vl_max_w_fre[k][i])
            doc.add_paragraph('绘制分幅未计权分频最大振级图，见下图')
            par5_5 = list(range(Pamount[k]))
            for i in range(Pamount[k]):
                doc.add_picture("result\\" + dirname +'\\'+numname[k]+ "\\4-3分幅计权分频振级-通道{}".format(i + 1) + ".png", height=Cm(5))
                par5_5[i] = doc.add_paragraph('图5-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + '加窗计权分频最大振级')
                par5_5[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.save("result\\" + dirname +'\\'+numname[k]+ "\\环境振动数据分析报告.docx")
        if mode==3:
            k=0
            doc = Document()
            tit = doc.add_paragraph('轨道环境振动对比分析报告')
            tit.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_heading('1.预处理')
            doc.add_paragraph('各个通道统计数据见表1-1')
            par1_1 = doc.add_paragraph('表1-1')
            par1_1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table1 = doc.add_table(Pamount[k] + 1, 11)
            column = ['通道编号', '零点偏移量', '零点偏移量', '零点偏移量', '零点偏移量', '零点偏移量', '时间趋势项大小', '时间趋势项大小', '时间趋势项大小', '时间趋势项大小', '时间趋势项大小' ]
            for i in range(11):
                table1.cell(0, i).text = column[i]
            for i in range(Pamount[k]):
                table1.cell(i + 1, 0).text = str(i + 1)
                table1.cell(i + 1, 1).text = str(x[0][i])
                table1.cell(i + 1, 2).text = str(x[1][i])
                table1.cell(i + 1, 3).text = str(x[2][i])
                table1.cell(i + 1, 4).text = str(x[3][i])
                table1.cell(i + 1, 5).text = str(x[4][i])
                table1.cell(i + 1, 6).text = str(y[0][i])
                table1.cell(i + 1, 7).text = str(y[1][i])
                table1.cell(i + 1, 8).text = str(y[2][i])
                table1.cell(i + 1, 9).text = str(y[3][i])
                table1.cell(i + 1, 10).text = str(y[4][i])
            table2 = doc.add_table(Pamount[k] + 1, 3)
            column1 = ['通道编号', '零点偏移量', '时间趋势项']
            for i in range(3):
                table2.cell(0, i).text = column1[i]
            for i in range(Pamount[k]):
                table2.cell(i + 1, 0).text = str(i + 1)
                table2.cell(i + 1, 1).text = str(x_tj[i])
                table2.cell(i + 1, 2).text = str(y_tj[i])
            table3 = doc.add_table(Pamount[k] + 1, 16)
            column3 = ['通道编号', '1', '1', '1', '1', '1', '2', '2', '2', '2','2','1', '1', '1', '1', '1']
            for i in range(16):
                table3.cell(0, i).text = column3[i]
            for i in range(Pamount[k]):
                table3.cell(i + 1, 0).text = str(i + 1)
                table3.cell(i + 1, 1).text = str(vlmax[0][i])
                table3.cell(i + 1, 2).text = str(vlmax[1][i])
                table3.cell(i + 1, 3).text = str(vlmax[2][i])
                table3.cell(i + 1, 4).text = str(vlmax[3][i])
                table3.cell(i + 1, 5).text = str(vlmax[4][i])
                table3.cell(i + 1, 6).text = str(vlarms[0][i])
                table3.cell(i + 1, 7).text = str(vlarms[1][i])
                table3.cell(i + 1, 8).text = str(vlarms[2][i])
                table3.cell(i + 1, 9).text = str(vlarms[3][i])
                table3.cell(i + 1, 10).text = str(vlarms[4][i])
                table3.cell(i + 1, 11).text = str(val[0][i])
                table3.cell(i + 1, 12).text = str(val[1][i])
                table3.cell(i + 1, 13).text = str(val[2][i])
                table3.cell(i + 1, 14).text = str(val[3][i])
                table3.cell(i + 1, 15).text = str(val[4][i])
            table4 = doc.add_table(Pamount[k] + 1, 4)
            column4 = ['通道编号', 'vlmax', 'vlarms','val']
            for i in range(4):
                table4.cell(0, i).text = column4[i]
            for i in range(Pamount[k]):
                table4.cell(i + 1, 0).text = str(i + 1)
                table4.cell(i + 1, 1).text = str(vlmax_tj[i])
                table4.cell(i + 1, 2).text = str(vlarms_tj[i])
                table4.cell(i + 1, 3).text = str(val_tj[i])
            doc.add_paragraph('利用结果绘制时程图，见下图')
            par4 = list(range(Pamount[k]))
            for i in range(Pamount[k]):
                doc.add_picture("result\\"+dirname+"\\1时域结果对比分析-通道{}".format(i+1)+".png",height=Cm(5))
                par4[i] = doc.add_paragraph('图1-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + '时域结果对比分析图')
                par4[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table5 = doc.add_table(Pamount[k] + 1, 11)
            column5 = ['通道编号', 'am', 'am', 'am', 'am', 'am', 'pow', 'pow', 'pow', 'pow','pow']
            for i in range(11):
                table5.cell(0, i).text = column5[i]
            for i in range(Pamount[k]):
                table5.cell(i + 1, 0).text = str(i + 1)
                table5.cell(i + 1, 1).text = str(max(max_am[0][i]))
                table5.cell(i + 1, 2).text = str(max(max_am[1][i]))
                table5.cell(i + 1, 3).text = str(max(max_am[2][i]))
                table5.cell(i + 1, 4).text = str(max(max_am[3][i]))
                table5.cell(i + 1, 5).text = str(max(max_am[4][i]))
                table5.cell(i + 1, 6).text = str(max_power[0][i])
                table5.cell(i + 1, 7).text = str(max_power[1][i])
                table5.cell(i + 1, 8).text = str(max_power[2][i])
                table5.cell(i + 1, 9).text = str(max_power[3][i])
                table5.cell(i + 1, 10).text = str(max_power[4][i])
            table6 = doc.add_table(Pamount[k] + 1, 3)
            column6 = ['通道编号', '最大振幅', '最大功率']
            for i in range(3):
                table6.cell(0, i).text = column6[i]
            for i in range(Pamount[k]):
                table6.cell(i + 1, 0).text = str(i + 1)
                table6.cell(i + 1, 1).text = str(am_tj[i])
                table6.cell(i + 1, 2).text = str(power_tj[i])
            par6 = list(range(Pamount[k]))
            for i in range(Pamount[k]):
                doc.add_picture("result\\"+dirname+"\\2频域结果对比分析-通道{}".format(i+1)+".png", height=Cm(5))
                par6[i] = doc.add_paragraph('图2-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + '频域结果对比分析图')
                par6[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table7 = doc.add_table(Pamount[k] + 1, 16)
            column7 = ['通道编号', '1', '1', '1', '1', '1', '2', '2', '2', '2', '2', '1', '1', '1', '1', '1']
            for i in range(16):
                table7.cell(0, i).text = column7[i]
            for i in range(Pamount[k]):
                table7.cell(i + 1, 0).text = str(i + 1)
                table7.cell(i + 1, 1).text = str(vl_maxz[0][i])
                table7.cell(i + 1, 2).text = str(vl_maxz[1][i])
                table7.cell(i + 1, 3).text = str(vl_maxz[2][i])
                table7.cell(i + 1, 4).text = str(vl_maxz[3][i])
                table7.cell(i + 1, 5).text = str(vl_maxz[4][i])
                table7.cell(i + 1, 6).text = str(vl_maxzw[0][i])
                table7.cell(i + 1, 7).text = str(vl_maxzw[1][i])
                table7.cell(i + 1, 8).text = str(vl_maxzw[2][i])
                table7.cell(i + 1, 9).text = str(vl_maxzw[3][i])
                table7.cell(i + 1, 10).text = str(vl_maxzw[4][i])
                table7.cell(i + 1, 11).text = str(vl_maxzwk[0][i])
                table7.cell(i + 1, 12).text = str(vl_maxzwk[1][i])
                table7.cell(i + 1, 13).text = str(vl_maxzwk[2][i])
                table7.cell(i + 1, 14).text = str(vl_maxzwk[3][i])
                table7.cell(i + 1, 15).text = str(vl_maxzwk[4][i])
            table8 = doc.add_table(Pamount[k] + 1, 4)
            column8 = ['通道编号', 'vlmaxz', 'vlmaxzw', 'vlmaxzwk']
            for i in range(4):
                table8.cell(0, i).text = column8[i]
            for i in range(Pamount[k]):
                table8.cell(i + 1, 0).text = str(i + 1)
                table8.cell(i + 1, 1).text = str(vl_maxz_tj[i])
                table8.cell(i + 1, 2).text = str(vl_maxzw_tj[i])
                table8.cell(i + 1, 3).text = str(vl_maxzwk_tj[i])
            par81 = list(range(Pamount[k]))
            for i in range(Pamount[k]):
                doc.add_picture("result\\"+dirname+"\\3-1Z振级-不计权对比分析-通道{}".format(i+1)+".png", height=Cm(5))
                par81[i] = doc.add_paragraph('图3-1-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + 'Z振级-不计权对比分析图')
                par81[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            par82 = list(range(Pamount[k]))
            for i in range(Pamount[k]):
                doc.add_picture("result\\"+dirname+"\\3-2Z振级-85版对比分析-通道{}".format(i+1)+".png", height=Cm(5))
                par82[i] = doc.add_paragraph('图3-2-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + 'Z振级-85版对比分析图')
                par82[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            par83 = list(range(Pamount[k]))
            for i in range(Pamount[k]):
                doc.add_picture("result\\"+dirname+"\\3-3Z振级-97版对比分析-通道{}".format(i+1)+".png", height=Cm(5))
                par83[i] = doc.add_paragraph('图3-3-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + 'Z振级-97版对比分析图')
                par83[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table9 = doc.add_table(Pamount[k] + 1, 11)
            column9 = ['通道编号', 'am', 'am', 'am', 'am', 'am', 'pow', 'pow', 'pow', 'pow', 'pow']
            for i in range(11):
                table9.cell(0, i).text = column9[i]
            for i in range(Pamount[k]):
                table9.cell(i + 1, 0).text = str(i + 1)
                table9.cell(i + 1, 1).text = str(vl_max_0_am[0][i])
                table9.cell(i + 1, 2).text = str(vl_max_0_am[1][i])
                table9.cell(i + 1, 3).text = str(vl_max_0_am[2][i])
                table9.cell(i + 1, 4).text = str(vl_max_0_am[3][i])
                table9.cell(i + 1, 5).text = str(vl_max_0_am[4][i])
                table9.cell(i + 1, 6).text = str(vl_max_0_fre[0][i])
                table9.cell(i + 1, 7).text = str(vl_max_0_fre[1][i])
                table9.cell(i + 1, 8).text = str(vl_max_0_fre[2][i])
                table9.cell(i + 1, 9).text = str(vl_max_0_fre[3][i])
                table9.cell(i + 1, 10).text = str(vl_max_0_fre[4][i])
            table10 = doc.add_table(Pamount[k] + 1, 3)
            column10 = ['通道编号', '分频最大振级', '分频最大振级对应频率']
            for i in range(3):
                table10.cell(0, i).text = column10[i]
            for i in range(Pamount[k]):
                table10.cell(i + 1, 0).text = str(i + 1)
                table10.cell(i + 1, 1).text = str(vl_max_0_am_tj[i])
                table10.cell(i + 1, 2).text = str(vl_max_0_fre_tj[i])
            par10 = list(range(Pamount[k]))
            for i in range(Pamount[k]):
                doc.add_picture("result\\"+dirname+"\\4-1未分幅未计权分频振级对比分析-通道{}".format(i+1)+".png", height=Cm(5))
                par10[i] = doc.add_paragraph('图4-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + '未分幅未计权分频振级对比分析图')
                par10[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table11 = doc.add_table(Pamount[k] + 1, 11)
            column11 = ['通道编号', 'am', 'am', 'am', 'am', 'am', 'pow', 'pow', 'pow', 'pow', 'pow']
            for i in range(11):
                table11.cell(0, i).text = column11[i]
            for i in range(Pamount[k]):
                table11.cell(i + 1, 0).text = str(i + 1)
                table11.cell(i + 1, 1).text = str(vl_max_1_am[0][i])
                table11.cell(i + 1, 2).text = str(vl_max_1_am[1][i])
                table11.cell(i + 1, 3).text = str(vl_max_1_am[2][i])
                table11.cell(i + 1, 4).text = str(vl_max_1_am[3][i])
                table11.cell(i + 1, 5).text = str(vl_max_1_am[4][i])
                table11.cell(i + 1, 6).text = str(vl_max_1_fre[0][i])
                table11.cell(i + 1, 7).text = str(vl_max_1_fre[1][i])
                table11.cell(i + 1, 8).text = str(vl_max_1_fre[2][i])
                table11.cell(i + 1, 9).text = str(vl_max_1_fre[3][i])
                table11.cell(i + 1, 10).text = str(vl_max_1_fre[4][i])
            table12 = doc.add_table(Pamount[k] + 1, 3)
            column12 = ['通道编号', '分频最大振级', '分频最大振级对应频率']
            for i in range(3):
                table12.cell(0, i).text = column12[i]
            for i in range(Pamount[k]):
                table12.cell(i + 1, 0).text = str(i + 1)
                table12.cell(i + 1, 1).text = str(vl_max_1_am_tj[i])
                table12.cell(i + 1, 2).text = str(vl_max_1_fre_tj[i])
            par12 = list(range(Pamount[k]))
            for i in range(Pamount[k]):
                doc.add_picture("result\\"+dirname+"\\4-2分幅未计权分频振级对比分析-通道{}".format(i+1)+".png", height=Cm(5))
                par12[i] = doc.add_paragraph('图5-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + '分幅未计权分频振级对比分析图')
                par12[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table13 = doc.add_table(Pamount[k] + 1, 11)
            column13 = ['通道编号', 'am', 'am', 'am', 'am', 'am', 'pow', 'pow', 'pow', 'pow', 'pow']
            for i in range(11):
                table13.cell(0, i).text = column13[i]
            for i in range(Pamount[k]):
                table13.cell(i + 1, 0).text = str(i + 1)
                table13.cell(i + 1, 1).text = str(vl_max_w_am[0][i])
                table13.cell(i + 1, 2).text = str(vl_max_w_am[1][i])
                table13.cell(i + 1, 3).text = str(vl_max_w_am[2][i])
                table13.cell(i + 1, 4).text = str(vl_max_w_am[3][i])
                table13.cell(i + 1, 5).text = str(vl_max_w_am[4][i])
                table13.cell(i + 1, 6).text = str(vl_max_w_fre[0][i])
                table13.cell(i + 1, 7).text = str(vl_max_w_fre[1][i])
                table13.cell(i + 1, 8).text = str(vl_max_w_fre[2][i])
                table13.cell(i + 1, 9).text = str(vl_max_w_fre[3][i])
                table13.cell(i + 1, 10).text = str(vl_max_w_fre[4][i])
            table14 = doc.add_table(Pamount[k] + 1, 3)
            column14 = ['通道编号', '分频最大振级', '分频最大振级对应频率']
            for i in range(3):
                table14.cell(0, i).text = column14[i]
            for i in range(Pamount[k]):
                table14.cell(i + 1, 0).text = str(i + 1)
                table14.cell(i + 1, 1).text = str(vl_max_w_am_tj[i])
                table14.cell(i + 1, 2).text = str(vl_max_w_fre_tj[i])
            par14 = list(range(Pamount[k]))
            for i in range(Pamount[k]):
                doc.add_picture("result\\"+dirname+"\\4-3分幅计权分频振级对比分析-通道{}".format(i+1)+".png", height=Cm(5))
                par14[i] = doc.add_paragraph('图6-{}'.format(i + 1) + ' 通道{}'.format(i + 1) + '分幅计权分频振级对比分析图')
                par14[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.save("result\\" + dirname + "\\环境振动数据分析报告.docx")
window = tk.Window()
window.title("振动数据分析（BJTU）")
window.geometry('1300x850')
if os.path.exists("result")==True:
    pass
else:
    os.makedirs("result")
if os.path.exists("cache")==True:
    pass
else:
    os.makedirs("cache")
menubar=tk.Menu(window)
modeselection_menu=tk.Menu(menubar,tearoff=0)
menubar.add_cascade(label="分析模式选择",menu=modeselection_menu)
modeselection_menu.add_command(label='单独分析',command=mode1)
modeselection_menu.add_command(label='批量分析',command=mode2)
modeselection_menu.add_command(label='对比分析',command=mode3)
menubar.add_command(label='导入数据',command=readfile)
prepare_menu=tk.Menu(menubar,tearoff=0)
menubar.add_cascade(label="预处理",menu=prepare_menu)
prepare_menu.add_command(label='预处理（不截取数据）',command=prepare_1)
prepare_menu.add_command(label='预处理（截取数据）',command=prepare_0)
menubar.add_command(label='时域分析',command=vltime)
fre_analsys_menu=tk.Menu(menubar,tearoff=0)
menubar.add_cascade(label='频域分析',menu=fre_analsys_menu)
fre_analsys_menu.add_command(label='傅里叶变换(显著频率)',command=fre)
fre_analsys_menu.add_command(label='最大Z振级分析',command=vlzmax)
vl_max_menu=tk.Menu(fre_analsys_menu,tearoff=0)
fre_analsys_menu.add_cascade(label='分频最大振级',menu=vl_max_menu)
vl_max_menu.add_command(label='未分幅未计权',command=vlmax0)
vl_max_menu.add_command(label='分幅未计权',command=vlmax1)
vl_max_menu.add_command(label='分幅计权',command=vlmaxw)
menubar.add_command(label='导出结果',command=outfile)
menubar.add_command(label='退出',command=window.quit)
window.config(menu=menubar)
window.mainloop()